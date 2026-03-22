from google import genai
from google.genai import types
import os
import glob
from datetime import datetime

client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))

def generate_lecture_material(lecture_num, level_from, level_to):
    today = datetime.now().strftime("%Y年%m月%d日")

    research_dir = os.path.expanduser(
        "~/Desktop/AI経営本部/北の株式投資大学/資料/デイリーリサーチ"
    )
    files = sorted(glob.glob(f"{research_dir}/*.md"), reverse=True)[:30]
    research_content = ""
    for f in files[:5]:
        with open(f, "r") as fp:
            research_content += fp.read()[:1000] + "\n\n"

    prompt_part1 = f"""
あなたは真田孔明のAI活用講義の講師AIです。
第{lecture_num}回の前半パートの講義資料を作成してください。

【作成日】{today}

【前半パート①：今月のAI業界最新情報（30分）】
以下の直近リサーチデータを元に、
今月最も重要なAI業界の最新情報を5〜7つ厳選して解説してください。

リサーチデータ：
{research_content}

出力形式：
# 第{lecture_num}回 前半①：今月のAI業界最新情報

## 今月の重要トピックTop5

### 1. 【トピック名】
・概要：
・なぜ重要か（WHY）：
・受講生のビジネスへの影響：

（以下同様に5〜7つ）

## 今月の一言まとめ
（初心者でも理解できる一文で今月のAIトレンドをまとめる）

【前半パート②：初心者向け解説（30分）】
上記の最新情報の中から最も初心者に役立つものを1つ選んで
Level 1の受講生でも今日から使える具体的な方法を解説してください。

出力形式：
# 第{lecture_num}回 前半②：今日から使える！初心者向け解説

## 今月のピックアップ：【テーマ名】

## なぜこれを選んだか

## 実際にやってみよう（ステップバイステップ）
Step 1：
Step 2：
Step 3：

## よくある質問・つまずきポイント
Q：
A：
"""

    prompt_part2 = f"""
あなたは真田孔明のAI活用講義の講師AIです。
第{lecture_num}回の後半パートの講義資料を作成してください。

【後半パート（講義内容）の文章ルール】
・文学スキル（俳句・村上春樹・Seth Godin）は使わない
・全ての主張に客観的な根拠・データを付ける
・結論→根拠→事実→応用の順で書く
・設計図・フレームワーク・ステップ形式を使う
・感情訴求は使わない
・論理性・根拠・データで受講生を動かす
・真田孔明の実体験は必ず具体的な数字で語る

【後半パート：Level {level_from}→{level_to} 実践カリキュラム（60分）】

以下の内容でワーク形式の講義資料を作成してください。

出力形式：
# 第{lecture_num}回 後半：Level {level_from}→{level_to} 実践

## 今回のゴール
この講義が終わったら何ができるようになるか

## Level {level_from}とLevel {level_to}の違い
・Level {level_from}の人は今：
・Level {level_to}の人は：
・その差を生むのは：

## 実践ワーク（手を動かす）

### ワーク1（15分）：
手順：
1.
2.
3.
確認ポイント：

### ワーク2（20分）：
手順：
1.
2.
3.
確認ポイント：

### ワーク3（20分）：
手順：
1.
2.
3.
確認ポイント：

## 本日の注意事項・よくある間違い
このLevel・テーマに関連するAI活用の注意点を3〜5個記載すること。
各注意点に：
・具体的な内容
・根拠・データ
・正しい対処法
※参照：AI活用講義/講義資料/第01回/AIと向き合う前に知っておくべきこと.md

## 本日のまとめ
・できるようになったこと：
・次回までの宿題：
・質問・サポートは：
"""

    save_dir = os.path.expanduser(
        f"~/Desktop/AI経営本部/AI活用講義/講義資料/第{lecture_num:02d}回"
    )
    os.makedirs(save_dir, exist_ok=True)

    print(f"第{lecture_num}回 前半パートを生成中...")
    try:
        response1 = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt_part1,
            config=types.GenerateContentConfig(
                tools=[types.Tool(google_search=types.GoogleSearch())]
            )
        )
    except Exception:
        response1 = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt_part1
        )

    with open(f"{save_dir}/前半_最新AI情報＋初心者解説.md", "w") as f:
        f.write(response1.text)
    print(f"✅ 前半パート保存完了")

    print(f"第{lecture_num}回 後半パートを生成中...")
    response2 = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt_part2
    )

    with open(f"{save_dir}/後半_Level{level_from}→{level_to}_実践.md", "w") as f:
        f.write(response2.text)
    print(f"✅ 後半パート保存完了")
    print(f"✅ 第{lecture_num}回 全資料完成：{save_dir}")

    # Word版の生成
    create_lecture_word(lecture_num, level_from, level_to,
                        response1.text, response2.text, save_dir)


def create_lecture_word(lecture_num, level_from, level_to,
                        part1_text, part2_text, save_dir):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    NAVY = RGBColor(0x1B, 0x4F, 0x72)

    # 表紙
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"第{lecture_num:02d}回 AI活用講義")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"Level {level_from}→{level_to} 実践カリキュラム")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x88, 0x87, 0x80)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"真田孔明 AI活用講座\n{datetime.now().strftime('%Y年%m月%d日')}")
    run3.font.size = Pt(11)

    doc.add_page_break()

    # 前半パート
    doc.add_heading("前半：AI最新情報＋初心者向け解説", level=1)
    doc.add_paragraph("[画像挿入箇所：今月のAI業界トピックマップ]").runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    for line in part1_text.split("\n"):
        if line.strip():
            if line.startswith("# ") or line.startswith("## "):
                doc.add_heading(line.lstrip("#").strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line.lstrip("#").strip(), level=3)
            else:
                doc.add_paragraph(line)

    doc.add_page_break()

    # 後半パート
    doc.add_heading(f"後半：Level {level_from}→{level_to} 実践", level=1)
    doc.add_paragraph("[画像挿入箇所：Level構造図]").runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    for line in part2_text.split("\n"):
        if line.strip():
            if line.startswith("# ") or line.startswith("## "):
                doc.add_heading(line.lstrip("#").strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line.lstrip("#").strip(), level=3)
            else:
                doc.add_paragraph(line)

    word_file = os.path.join(save_dir, f"第{lecture_num:02d}回_講義資料.docx")
    doc.save(word_file)
    print(f"✅ Wordファイル保存：{word_file}")

    # 画像プロンプト一覧を別ファイルで保存
    image_prompts = [
        {
            "section": "今月のAI業界トピックマップ",
            "prompt": f"ナノバナナプロで以下の内容を図解してください。\n"
                      f"第{lecture_num}回AI活用講義の前半パートで扱う今月のAI業界最新トピック5〜7つを\n"
                      f"中心から放射状に広がるマインドマップ形式で整理。\n"
                      f"スタイル：霞が関のポンチ絵のように情報を整理。サイズ：横長16:9",
            "location": "前半パート冒頭"
        },
        {
            "section": "Level構造図",
            "prompt": f"ナノバナナプロで以下のレベル構造を図解してください。\n"
                      f"Level 1（質問・回答のみ）→Level 2（プロンプト工夫）→Level 3（API活用）→"
                      f"Level 4（複数AI連携）→Level 5（AI組織経営）の5段階の階段図。\n"
                      f"Level {level_from}からLevel {level_to}の部分をハイライト。\n"
                      f"各Levelの説明を簡潔に記載。サイズ：横長16:9",
            "location": "後半パート冒頭"
        },
        {
            "section": "実践ワークフロー図",
            "prompt": f"ナノバナナプロで以下のフローを図解してください。\n"
                      f"第{lecture_num}回の実践ワーク（Level {level_from}→{level_to}）の\n"
                      f"ステップ1→ステップ2→ステップ3の流れを横方向のフローチャートで表現。\n"
                      f"各ステップにアイコンを追加。\n"
                      f"スタイル：シンプルで見やすいフローチャート。背景は白。サイズ：横長16:9",
            "location": "後半パート・実践ワークセクション"
        },
    ]

    prompt_file = os.path.join(save_dir, f"第{lecture_num:02d}回_画像プロンプト一覧.txt")
    with open(prompt_file, "w") as f:
        f.write(f"=== 第{lecture_num:02d}回 画像プロンプト一覧 ===\n")
        f.write("以下をGemini（ナノバナナプロ）に貼り付けて画像を生成してください。\n\n")
        for i, p in enumerate(image_prompts, 1):
            f.write(f"【画像{i}：{p['section']}用】\n")
            f.write(f"ナノバナナプロへの指示：\n{p['prompt']}\n")
            f.write(f"挿入箇所：{p['location']}\n\n")
        f.write("=== 画像プロンプト終わり ===\n")

    print(f"✅ 画像プロンプト一覧：{prompt_file}")


CURRICULUM = {
    1:  (1, 2),
    2:  (1, 2),
    3:  (2, 3),
    4:  (2, 3),
    5:  (2, 3),
    6:  (3, 4),
    7:  (3, 4),
    8:  (3, 4),
    9:  (4, 5),
    10: (4, 5),
    11: (4, 5),
    12: (5, 5),
}

if __name__ == "__main__":
    import sys
    lecture_num = int(sys.argv[1]) if len(sys.argv) > 1 else 1
    level_from, level_to = CURRICULUM.get(lecture_num, (1, 2))
    generate_lecture_material(lecture_num, level_from, level_to)
