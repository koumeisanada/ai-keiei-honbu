import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SAVE_DIR = os.path.expanduser(
    "~/Desktop/AI経営本部/AI活用講義"
)

def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '1A1A2E')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    return p

def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f'■ {text}')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def add_subsection_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f'◆ {text}')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    indent = '　' * level
    run = p.add_run(f'{indent}・{text}')
    run.font.size = Pt(10.5)
    return p

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A1A2E')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def create_report():
    today = datetime.now().strftime("%Y年%m月%d日")
    doc = Document()

    # ページ設定
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # =============================
    # 表紙
    # =============================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    add_title(doc, '新規プロジェクト企画報告書')
    doc.add_paragraph()
    add_subtitle(doc, '最新AI活用講座プロジェクト')
    add_subtitle(doc, '〜 エンジニア経験ゼロでも1年でAI組織を構築する年間プログラム 〜')
    doc.add_paragraph()
    doc.add_paragraph()

    # 報告書情報表
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    set_table_border(table)

    info = [
        ('報告日', today),
        ('報告者', '真田孔明'),
        ('プロジェクト名', '最新AI活用講座（正式名称決定予定）'),
        ('報告先', 'メンター・上長'),
    ]
    for i, (label, value) in enumerate(info):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_bg(row.cells[0], 'E8E8F0')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    if cell == row.cells[0]:
                        run.font.bold = True

    doc.add_page_break()

    # =============================
    # 第1章：プロジェクト概要
    # =============================
    add_section_title(doc, '第1章：プロジェクト概要')
    add_hr(doc)
    doc.add_paragraph()

    add_subsection_title(doc, 'プロジェクトの目的')
    add_body(doc,
        '日本のビジネスパーソン・個人事業主・経営者を対象に、'
        'AIを「道具として使う」段階から「組織として経営する」段階へと導く'
        '年間講義プログラムを提供する。\n\n'
        'エンジニア経験ゼロの講師自身が48時間でAI組織を構築した実証済みの'
        'ノウハウを体系化し、受講生が1年間でAI活用Level 5に到達できる'
        '再現性のあるカリキュラムを提供する。'
    )
    doc.add_paragraph()

    add_subsection_title(doc, '事業規模・目標数値')

    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Table Grid'
    set_table_border(table2)

    kpi_data = [
        ('目標年商', '1億円（最大3億円）'),
        ('目標受講生数', '300名（最大1,000名）'),
        ('価格設定', '月額3万円 または 年間一括30万円'),
        ('講義頻度', '月2回・年間24回'),
        ('講義形式', 'ZOOMオンライン ＋ オフライン（選択制）'),
        ('完成目標', '2026年4月19日（セールスレター完成・告知開始）'),
    ]
    for i, (label, value) in enumerate(kpi_data):
        row = table2.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_bg(row.cells[0], 'E8E8F0')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    if cell == row.cells[0]:
                        run.font.bold = True
    doc.add_paragraph()

    add_subsection_title(doc, '収益計画（逆算）')

    table3 = doc.add_table(rows=4, cols=3)
    table3.style = 'Table Grid'
    set_table_border(table3)

    headers = ['シナリオ', '受講生数', '年間売上']
    header_row = table3.rows[0]
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        set_cell_bg(header_row.cells[i], '1A1A2E')
        for para in header_row.cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)

    revenue_data = [
        ('最低目標', '300名', '約9,000万円'),
        ('標準目標', '500名', '約1億5,000万円'),
        ('最大目標', '1,000名', '約3億円'),
    ]
    for i, (scenario, members, revenue) in enumerate(revenue_data):
        row = table3.rows[i + 1]
        row.cells[0].text = scenario
        row.cells[1].text = members
        row.cells[2].text = revenue
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)

    doc.add_page_break()

    # =============================
    # 第2章：ターゲット
    # =============================
    add_section_title(doc, '第2章：メインターゲット')
    add_hr(doc)
    doc.add_paragraph()

    add_subsection_title(doc, 'メインターゲット像')

    table4 = doc.add_table(rows=7, cols=2)
    table4.style = 'Table Grid'
    set_table_border(table4)

    target_data = [
        ('年齢層', '40代（アラフィフ）が中心・30〜50代'),
        ('職業', 'サラリーマン・個人事業主・中小企業経営者'),
        ('AI経験', 'Level 1〜2（ChatGPTを使ったことがある程度）'),
        ('技術背景', 'エンジニア経験なし・プログラミング未経験'),
        ('抱える課題', 'AIを使いたいが何から始めればいいかわからない'),
        ('求めていること', 'ビジネスの効率化・時間の創出・収益向上'),
        ('既存メンバー', '北の株式投資大学・地下ソサエティ・ワンチームのメンバー'),
    ]
    for i, (label, value) in enumerate(target_data):
        row = table4.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_bg(row.cells[0], 'E8E8F0')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    if cell == row.cells[0]:
                        run.font.bold = True
    doc.add_paragraph()

    add_subsection_title(doc, 'AI活用Level定義（共通基準）')

    table5 = doc.add_table(rows=6, cols=3)
    table5.style = 'Table Grid'
    set_table_border(table5)

    level_headers = ['Level', '定義', '人口割合']
    header_row5 = table5.rows[0]
    for i, h in enumerate(level_headers):
        header_row5.cells[i].text = h
        set_cell_bg(header_row5.cells[i], '1A1A2E')
        for para in header_row5.cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)

    level_data = [
        ('Level 1', '質問・回答のみ（ChatGPTに聞くだけ）', '約80%'),
        ('Level 2', 'プロンプト工夫・業務への組み込み', '約15%'),
        ('Level 3', 'API活用・自動化スクリプトの実行', '約4%'),
        ('Level 4', '複数AI連携・エージェント設計', '約0.9%'),
        ('Level 5', 'AI組織経営・完全自動化（目標）', '約0.1%'),
    ]
    for i, (level, definition, ratio) in enumerate(level_data):
        row = table5.rows[i + 1]
        row.cells[0].text = level
        row.cells[1].text = definition
        row.cells[2].text = ratio
        if level == 'Level 5':
            set_cell_bg(row.cells[0], '534AB7')
            set_cell_bg(row.cells[1], '534AB7')
            set_cell_bg(row.cells[2], '534AB7')
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
                        run.font.size = Pt(10.5)
        else:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10.5)

    doc.add_paragraph()
    add_body(doc, '※ターゲットはLevel 1〜2にいる方。1年間でLevel 5を目指す。')

    doc.add_page_break()

    # =============================
    # 第3章：講座の差別化ポイント
    # =============================
    add_section_title(doc, '第3章：講座の差別化ポイント')
    add_hr(doc)
    doc.add_paragraph()

    add_subsection_title(doc, '他社との根本的な違い')

    table6 = doc.add_table(rows=7, cols=3)
    table6.style = 'Table Grid'
    set_table_border(table6)

    diff_headers = ['比較項目', '一般的なAI講座', '真田孔明のAI講座']
    header_row6 = table6.rows[0]
    for i, h in enumerate(diff_headers):
        header_row6.cells[i].text = h
        set_cell_bg(header_row6.cells[i], '1A1A2E')
        for para in header_row6.cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)

    diff_data = [
        ('ゴールLevel', 'Level 2〜3', 'Level 5'),
        ('講師の背景', 'エンジニア・研究者', 'エンジニア経験ゼロの営業出身'),
        ('講義頻度', '月1回が多い', '月2回・年間24回'),
        ('形式', 'オンラインのみ', 'ZOOM＋オフライン選択制'),
        ('補足内容', '自社カリキュラムのみ', '世界45校の最新ノウハウを補足提供'),
        ('再現性の証明', 'なし・または理論のみ', '48時間でLevel 5を実証済み'),
    ]
    for i, (item, general, sanada) in enumerate(diff_data):
        row = table6.rows[i + 1]
        row.cells[0].text = item
        row.cells[1].text = general
        row.cells[2].text = sanada
        set_cell_bg(row.cells[0], 'E8E8F0')
        set_cell_bg(row.cells[2], 'F0F0FF')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    if cell == row.cells[0]:
                        run.font.bold = True
    doc.add_paragraph()

    add_subsection_title(doc, '講師プロフィール・権威性')
    bullets = [
        '2004年〜メルマガ発信継続（20年以上・言語化力の蓄積）',
        'Amazonベストセラー著者',
        '北の株式投資大学・地下ソサエティ・ワンチーム主宰（12年以上）',
        'NVIDIA・Apple・Microsoft・Netflix等の米国株投資家',
        'エンジニア経験ゼロで48時間でAI Level 5を実現（再現性の証明）',
        '3大AI（Claude・Gemini・ChatGPT）横断活用システムを構築・実稼働中',
        'AI活用世界上位0.1%（Level 5）の実証済み',
    ]
    for b in bullets:
        add_bullet(doc, b)

    doc.add_page_break()

    # =============================
    # 第4章：年間カリキュラム
    # =============================
    add_section_title(doc, '第4章：年間カリキュラム（全24回）')
    add_hr(doc)
    doc.add_paragraph()

    add_subsection_title(doc, '毎回の講義構成')

    table7 = doc.add_table(rows=3, cols=2)
    table7.style = 'Table Grid'
    set_table_border(table7)

    structure_data = [
        ('第1回のみ', 'Level 1→5の全体像・注意事項・Level 1→2への第一歩\n（AI最新情報パートなし）'),
        ('前半60分\n（第2回以降）', 'AI米国最前線情報＋具体的な活用方法\n・新AIモデルのリリース情報\n・NVIDIA・Apple・Microsoft等の最新AI戦略\n・日本・EU・米国のAI規制の動向\n・今月登場した話題のAIツール\n・今日から使える具体的な活用方法の実演'),
        ('後半60分\n（全回共通）', 'Level 1→5 年間プログラム\n・体系的なカリキュラムに沿った実践\n・永続的に使える資産となる資料'),
    ]
    for i, (label, content) in enumerate(structure_data):
        row = table7.rows[i]
        row.cells[0].text = label
        row.cells[1].text = content
        set_cell_bg(row.cells[0], 'E8E8F0')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if cell == row.cells[0]:
                        run.font.bold = True
    doc.add_paragraph()

    add_subsection_title(doc, '全24回 カリキュラム一覧')

    curriculum = [
        ('第1クール', 'Level 1→2\n基礎固め', [
            ('第1回', 'AI経営本部の全体像と始め方',
             'Level 1〜5の定義・年間ロードマップ・注意事項・推奨デバイス'),
            ('第2回', '3大AIの正しい使い分け',
             'Claude・Gemini・ChatGPTの違い・強みと弱み・用途別使い分け'),
            ('第3回', 'プロンプトの精度を上げる',
             'ビジネス別プロンプトの作り方・実践ワーク'),
            ('第4回', '業務への組み込み方',
             '時間削減の設計図・Level 2到達確認'),
        ]),
        ('第2クール', 'Level 2→3\n自動化の入口', [
            ('第5回', 'APIキーの取得と接続',
             '3大AIのAPI取得・環境変数への設定・接続テスト'),
            ('第6回', '最初の自動化を作る',
             '最初のスクリプト実行・エラーの対処法'),
            ('第7回', 'パイプラインの設計',
             'テーマ入力→全自動生成の仕組み・自分用パイプライン設計'),
            ('第8回', 'パイプラインの完成と実運用',
             'デイリーパイプラインの完成・自動保存・Level 3到達確認'),
        ]),
        ('第3クール', 'Level 3→4\nAI組織の構築', [
            ('第9回', 'AI社長の設計',
             'CLAUDE.mdとは何か・AI社長の役割と設定'),
            ('第10回', 'AIマネージャーの配置',
             'プロジェクト別マネージャー設定・役割分担の設計図'),
            ('第11回', 'SKILLファイルで記憶を管理する',
             'AIの記憶をファイルで管理・コスト最小化の設計'),
            ('第12回', 'AI組織の完成と実運用',
             '3大AI横断活用の実装・品質チェックの自動化・Level 4到達確認'),
        ]),
        ('第4クール', 'Level 4→5\n完全自動化', [
            ('第13回', 'GitHubでクラウド管理',
             'AI経営本部のクラウド化・バックアップと履歴管理'),
            ('第14回', 'iPhoneからの遠隔操作',
             'Claude Code Webの設定・スマホ1台で経営する仕組み'),
            ('第15回', '競合調査・リサーチの自動化',
             'ライバル調査の自動化・情報収集コストをゼロにする'),
            ('第16回', '完全自動化の完成',
             '毎朝6時の自動実行設定・全業務の自動化チェックリスト・Level 5到達確認'),
        ]),
        ('第5クール', 'Level 5\n実装と深化', [
            ('第17回', 'ビジネス別AI実装①サラリーマン・会社員',
             '業務効率化・社内評価向上・副業への応用'),
            ('第18回', 'ビジネス別AI実装②物販ビジネス',
             '商品リサーチ・出品文章・顧客対応の自動化'),
            ('第19回', 'ビジネス別AI実装③融資・法人経営',
             '銀行融資資料のAI自動作成・法人経営効率化'),
            ('第20回', 'ビジネス別AI実装④投資リサーチ',
             '米国株9社デイリー分析・投資判断サポートの自動化'),
        ]),
        ('第6クール', 'Level 5\n維持と発展', [
            ('第21回', 'セールス・集客へのAI応用',
             'メルマガ・LINE・SNSの自動生成・セールスレターのAI制作'),
            ('第22回', '世界45校のAI講座補足ノウハウ',
             'Make・Zapier・n8nの活用・AIエージェントの最新動向'),
            ('第23回', 'AI組織の改善と進化',
             'スキルファイルの更新・新しいAIツールへの対応方法'),
            ('第24回', '次年度の設計と総復習',
             '1年間の振り返り・Level 5維持のロードマップ・次年度の目標設定'),
        ]),
    ]

    for cool_name, cool_sub, lectures in curriculum:
        table_c = doc.add_table(rows=1 + len(lectures), cols=3)
        table_c.style = 'Table Grid'
        set_table_border(table_c)

        header_row_c = table_c.rows[0]
        header_row_c.cells[0].text = f'{cool_name}：{cool_sub}'
        header_row_c.cells[1].text = '講義タイトル'
        header_row_c.cells[2].text = '内容・習得スキル'
        set_cell_bg(header_row_c.cells[0], '534AB7')
        set_cell_bg(header_row_c.cells[1], '534AB7')
        set_cell_bg(header_row_c.cells[2], '534AB7')
        for cell in header_row_c.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)

        for i, (num, title, content) in enumerate(lectures):
            row = table_c.rows[i + 1]
            row.cells[0].text = num
            row.cells[1].text = title
            row.cells[2].text = content
            set_cell_bg(row.cells[0], 'F0F0FF')
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        if cell == row.cells[0]:
                            run.font.bold = True

        table_c.columns[0].width = Cm(2)
        table_c.columns[1].width = Cm(6)
        table_c.columns[2].width = Cm(8.5)

        doc.add_paragraph()

    doc.add_page_break()

    # =============================
    # 第5章：マーケティング計画
    # =============================
    add_section_title(doc, '第5章：マーケティング計画')
    add_hr(doc)
    doc.add_paragraph()

    add_subsection_title(doc, '集客チャネルと優先順位')

    table8 = doc.add_table(rows=6, cols=3)
    table8.style = 'Table Grid'
    set_table_border(table8)

    mkt_headers = ['優先度', 'チャネル', '戦略・施策']
    header_row8 = table8.rows[0]
    for i, h in enumerate(mkt_headers):
        header_row8.cells[i].text = h
        set_cell_bg(header_row8.cells[i], '1A1A2E')
        for para in header_row8.cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)

    mkt_data = [
        ('1位', '既存メルマガ読者', '20年・数万人の信頼関係・転換率最高'),
        ('2位', '既存コミュニティ\nメンバー', '北の株式投資大学・地下ソサエティ・ワンチーム\nバックエンド販売として自然な流れ'),
        ('3位', 'Instagram\nゼータアカウント', 'Level 1〜5シリーズリール投稿\n俳句×村上春樹スタイルの差別化コンテンツ'),
        ('4位', '口コミ・紹介', '受講生が「Level 5になった」と話す\nコスト最小・信頼性最高'),
        ('5位', 'AmazonKindle本', 'AI講座内容の書籍化\n「本を読んで受講したい」という流入'),
    ]
    for i, row_data in enumerate(mkt_data):
        row = table8.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        set_cell_bg(row.cells[0], 'F0F0FF')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

    add_subsection_title(doc, 'スケジュール')

    table9 = doc.add_table(rows=5, cols=2)
    table9.style = 'Table Grid'
    set_table_border(table9)

    schedule_data = [
        ('2026年3月', 'セールスレター制作中・既存メルマガでAI話題の発信開始'),
        ('2026年4月19日', 'セールスレター完成・公開・初回募集開始'),
        ('2026年4〜6月', '初回受講生50〜100名獲得・無料ZOOMミニ講義開催'),
        ('2026年7〜9月', '受講生の変化を発信・既存コミュニティへのクロスセル強化'),
        ('2026年10〜12月', '年間一括オファー・紹介制度・累計300名目標'),
    ]
    for i, (period, action) in enumerate(schedule_data):
        row = table9.rows[i]
        row.cells[0].text = period
        row.cells[1].text = action
        set_cell_bg(row.cells[0], 'E8E8F0')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    if cell == row.cells[0]:
                        run.font.bold = True

    doc.add_page_break()

    # =============================
    # 第6章：講座の特典・サポート体制
    # =============================
    add_section_title(doc, '第6章：特典・サポート体制')
    add_hr(doc)
    doc.add_paragraph()

    add_subsection_title(doc, '受講生への特典')

    table10 = doc.add_table(rows=4, cols=2)
    table10.style = 'Table Grid'
    set_table_border(table10)

    benefit_data = [
        ('特典①', 'AI用語集（Word版）\n・80語以上収録\n・40代の非エンジニアでも100%わかる解説\n・3大AI横断制作で正確性を保証'),
        ('特典②', '講義資料（Word版・毎回）\n・前半：AI最新情報資料（毎回更新）\n・後半：Level講義資料（永続的に使える資産）'),
        ('特典③', 'AI経営本部構築マニュアル\n・4日間で自分のAI経営本部を構築できる手順書\n・自分のビジネスに合わせたカスタマイズガイド付き'),
    ]
    for i, (label, content) in enumerate(benefit_data):
        row = table10.rows[i]
        row.cells[0].text = label
        row.cells[1].text = content
        set_cell_bg(row.cells[0], '534AB7')
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10.5)

    doc.add_paragraph()

    add_subsection_title(doc, '必要デバイス・環境')

    table11 = doc.add_table(rows=3, cols=3)
    table11.style = 'Table Grid'
    set_table_border(table11)

    device_headers = ['デバイス', '推奨スペック', '最低スペック']
    header_row11 = table11.rows[0]
    for i, h in enumerate(device_headers):
        header_row11.cells[i].text = h
        set_cell_bg(header_row11.cells[i], '1A1A2E')
        for para in header_row11.cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)

    device_data = [
        ('Mac', 'MacBook Air M2以降\nメモリ16GB以上\nmacOS Ventura以降',
         'Intel Mac 2018年以降\nメモリ8GB\nmacOS Monterey以降'),
        ('Windows', 'Core i5以降\nメモリ16GB以上\nWindows 11',
         'Core i3以降\nメモリ8GB\nWindows 10'),
    ]
    for i, (device, recommend, minimum) in enumerate(device_data):
        row = table11.rows[i + 1]
        row.cells[0].text = device
        row.cells[1].text = recommend
        row.cells[2].text = minimum
        set_cell_bg(row.cells[0], 'E8E8F0')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if cell == row.cells[0]:
                        run.font.bold = True

    doc.add_paragraph()
    add_body(doc, '月額API費用目安：17,000〜20,000円（Gemini・OpenAI・Claude API）')
    add_body(doc, '※全てのAPIは無料枠から始められます')

    doc.add_page_break()

    # =============================
    # 第7章：まとめ・所見
    # =============================
    add_section_title(doc, '第7章：まとめ・所見')
    add_hr(doc)
    doc.add_paragraph()

    add_subsection_title(doc, 'このプロジェクトの本質的な価値')
    add_body(doc,
        '本プロジェクトは単なる「AIツールの使い方講座」ではない。\n\n'
        'エンジニア経験ゼロの営業出身のアラフィフが、48時間でAI組織を構築した。'
        'この事実そのものが、最大の差別化であり、最強の再現性の証明である。\n\n'
        'AIを「使う」人を育てるのではなく、AIを「経営する」人を育てる。\n'
        'Level 1（約80%）から始まり、Level 5（約0.1%）を目指す。\n'
        'この1年間の旅が、受講生のビジネスと人生を根本的に変える。'
    )
    doc.add_paragraph()

    add_subsection_title(doc, '今後のアクション')
    actions = [
        '2026年3月中：プロジェクト正式名称の決定',
        '2026年4月19日：セールスレター完成・公開',
        '2026年4月：初回募集開始・無料ZOOMミニ講義開催',
        '2026年6月：第1期受講生スタート',
        '2026年12月：累計300名達成・年商1億円目標',
    ]
    for action in actions:
        add_bullet(doc, action)

    doc.add_paragraph()
    doc.add_paragraph()

    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sign = p_sign.add_run(f'報告者：真田孔明\n作成日：{today}')
    run_sign.font.size = Pt(11)

    # 保存
    os.makedirs(SAVE_DIR, exist_ok=True)
    word_file = os.path.join(
        SAVE_DIR,
        f'【報告書】最新AI活用講座プロジェクト_{today}.docx'
    )
    doc.save(word_file)
    print(f'✅ 報告書完成：{word_file}')
    return word_file

def main():
    today = datetime.now().strftime("%Y年%m月%d日")
    print("=" * 60)
    print("プロジェクト報告書 制作開始")
    print("=" * 60)

    word_file = create_report()

    import subprocess
    subprocess.run(["bash", "-c",
        "cd ~/Desktop/AI経営本部 && "
        "git add AI活用講義/ && "
        f"git commit -m 'AI講座プロジェクト報告書完成（{today}）' && "
        "git push origin salesletter-development"
    ])

    print("\n✅ GitHubプッシュ完了")
    print(f"\n完成ファイル：{word_file}")

if __name__ == "__main__":
    main()
