import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SAVE_DIR = os.path.expanduser(
    "~/Desktop/AI経営本部/AI活用講義/講義資料/第01回"
)

def set_table_border(table, color='1A1A2E'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

def add_hr(doc, color='534AB7'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f'■ {text}')
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.font.name = 'メイリオ'
    add_hr(doc)
    doc.add_paragraph()

def add_subsection(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f'◆ {text}')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    run.font.name = 'メイリオ'

def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = 'メイリオ'
    return p

def add_point_box(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_border(table, '534AB7')
    set_cell_bg(table.rows[0].cells[0], 'F0F0FF')
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].clear()
    p1 = cell.add_paragraph()
    r1 = p1.add_run(f'💡 {title}')
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    r1.font.name = 'メイリオ'
    p2 = cell.add_paragraph(text)
    for r in p2.runs:
        r.font.size = Pt(10.5)
        r.font.name = 'メイリオ'
    doc.add_paragraph()

def add_warning_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_border(table, 'CC0000')
    set_cell_bg(table.rows[0].cells[0], 'FFF0F0')
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].clear()
    p = cell.add_paragraph()
    r = p.add_run(f'⚠️ 注意　{text}')
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    r.font.name = 'メイリオ'
    doc.add_paragraph()

def add_step_table(doc, steps):
    table = doc.add_table(rows=len(steps)+1, cols=3)
    set_table_border(table)
    for i, h in enumerate(['手順', '操作内容', '確認ポイント']):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, '1A1A2E')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)
                run.font.name = 'メイリオ'
    for i, (step, content, check) in enumerate(steps):
        row = table.rows[i+1]
        row.cells[0].text = step
        row.cells[1].text = content
        row.cells[2].text = check
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, 'F0F0FF')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'メイリオ'
    table.columns[0].width = Cm(2)
    table.columns[1].width = Cm(10)
    table.columns[2].width = Cm(4.5)
    doc.add_paragraph()

def add_prompt_box(doc, title, prompt_text):
    p = doc.add_paragraph()
    r = p.add_run(f'📋 {title}')
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    r.font.name = 'メイリオ'
    table = doc.add_table(rows=1, cols=1)
    set_table_border(table, '1A1A2E')
    set_cell_bg(table.rows[0].cells[0], 'F5F5F5')
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].clear()
    for line in prompt_text.strip().split('\n'):
        p2 = cell.add_paragraph(line)
        for r2 in p2.runs:
            r2.font.size = Pt(10)
            r2.font.name = 'Courier New'
            r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()

def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        r = p.add_run(f'☐  {item}')
        r.font.size = Pt(10.5)
        r.font.name = 'メイリオ'

def create_manual():
    today = datetime.now().strftime('%Y年%m月%d日')
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    add_page_number(doc)

    # ========== 表紙 ==========
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Claudeだけで作る')
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    r.font.name = 'メイリオ'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('AI経営本部 構築マニュアル')
    r2.font.size = Pt(28)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    r2.font.name = 'メイリオ'

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('〜 エンジニア経験ゼロでも・プログラミング不要・\nClaude一つで自分のAI組織が動き出す 〜')
    r3.font.size = Pt(13)
    r3.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    r3.font.name = 'メイリオ'

    doc.add_paragraph()

    tbl = doc.add_table(rows=3, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl)
    for i, (label, value) in enumerate([
        ('発行日', today),
        ('対象', '真田孔明 AI活用講座 受講生'),
        ('所要時間', '全工程 約4時間で完成'),
    ]):
        row = tbl.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_bg(row.cells[0], 'E8E8F0')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = 'メイリオ'
                    if cell == row.cells[0]:
                        run.font.bold = True

    doc.add_page_break()

    # ========== はじめに ==========
    add_section_title(doc, 'はじめに')
    add_body(doc,
        'このマニュアルは、Claude一つだけを使って\n'
        'AI経営本部を構築するための手順書です。\n\n'
        'Gemini・ChatGPTなど他のAIは一切使いません。\n'
        'Claudeのチャット画面と、Claude Codeだけで完結します。\n\n'
        'エンジニア経験ゼロでも大丈夫です。\n'
        'コードを書く必要もありません。\n'
        '日本語で話しかけるだけで、全てが動き出します。'
    )
    doc.add_paragraph()
    add_point_box(doc, 'このマニュアルで完成すること',
        '✅ AI社長（CLAUDE.md）の設定\n'
        '✅ ビジネス別AIマネージャーの配置\n'
        '✅ SKILLファイルによる記憶の管理\n'
        '✅ GitHubでのクラウド保存\n'
        '✅ iPhoneからの遠隔操作\n'
        '✅ 毎日のルーティン作業の自動化'
    )
    doc.add_page_break()

    # ========== 第1章：準備 ==========
    add_section_title(doc, '第1章：準備するもの')
    add_subsection(doc, '1-1. 必要なものは2つだけ')
    doc.add_paragraph()

    tbl2 = doc.add_table(rows=3, cols=3)
    set_table_border(tbl2)
    for i, h in enumerate(['必要なもの', '用途', '費用']):
        tbl2.rows[0].cells[i].text = h
        set_cell_bg(tbl2.rows[0].cells[i], '1A1A2E')
        for para in tbl2.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)
                run.font.name = 'メイリオ'
    for i, (item, use, cost) in enumerate([
        ('Claudeアカウント', 'AIチャット・Claude Code', 'Maxプラン 月額約15,000円'),
        ('GitHubアカウント', 'データのクラウド保存', '無料'),
    ]):
        row = tbl2.rows[i+1]
        row.cells[0].text = item
        row.cells[1].text = use
        row.cells[2].text = cost
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, 'F0F0FF')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = 'メイリオ'
    doc.add_paragraph()

    add_subsection(doc, '1-2. Claudeのアカウント作成')
    add_step_table(doc, [
        ('①', 'https://claude.ai をブラウザで開く', 'サインアップ画面が表示される'),
        ('②', 'メールアドレスで登録する', '確認メールが届く'),
        ('③', 'Maxプランにアップグレードする', '月額約15,000円'),
        ('④', 'claude.aiにログインして\nチャット画面が表示されることを確認', 'チャット画面が開けばOK'),
    ])

    add_subsection(doc, '1-3. Claude Codeのインストール')
    add_body(doc, 'Claude Codeとは、パソコンのターミナル（黒い画面）から\nClaudeを操作してファイル作成・保存ができるツールです。')
    doc.add_paragraph()
    add_step_table(doc, [
        ('①', 'Macの場合：Launchpadで「ターミナル」と検索して起動\nWindowsの場合：「コマンドプロンプト」を起動', 'ターミナルが開く'),
        ('②', '以下を入力してEnter：\nnpm install -g @anthropic-ai/claude-code', 'インストールが始まる'),
        ('③', 'インストール完了後に以下を入力：\nclaude login', 'ブラウザが開く'),
        ('④', 'Claudeのアカウントでログインする', '「認証完了」と表示される'),
        ('⑤', 'ターミナルに戻って以下を入力：\nclaude', 'Claude Codeが起動する'),
    ])

    add_warning_box(doc, 'Node.jsが未インストールの場合はhttps://nodejs.org からインストールしてください。インストール後に上記手順②から再開してください。')
    doc.add_page_break()

    # ========== 第2章：フォルダ構造の作成 ==========
    add_section_title(doc, '第2章：AI経営本部のフォルダを作る')
    add_subsection(doc, '2-1. Claude Codeに話しかけるだけで完成する')
    add_body(doc, 'Claude Codeを起動したら、以下のプロンプトをコピペして送信してください。\nこれだけで全フォルダが自動で作成されます。')
    doc.add_paragraph()

    add_prompt_box(doc, 'プロンプト①：フォルダ構造の自動作成',
        'デスクトップに「AI経営本部」フォルダを作成して\n'
        '以下の構造を全て自動で作ってください。\n\n'
        'AI経営本部/\n'
        '├── CLAUDE.md（AI社長）\n'
        '├── 人財育成/\n'
        '│   ├── CLAUDE.md\n'
        '│   └── 資料/\n'
        '├── 物販ビジネス/\n'
        '│   ├── CLAUDE.md\n'
        '│   └── 資料/\n'
        '├── 銀行融資/\n'
        '│   ├── CLAUDE.md\n'
        '│   └── 資料/\n'
        '├── 会社法人/\n'
        '│   ├── CLAUDE.md\n'
        '│   └── 資料/\n'
        '├── 米国成長株/\n'
        '│   ├── CLAUDE.md\n'
        '│   └── 資料/\n'
        '└── 成果物/'
    )

    add_point_box(doc, 'なぜフォルダ分けが重要か',
        'フォルダ＝部署です。\n'
        'AI経営本部フォルダを開いてClaude Codeを起動すると、\n'
        'Claudeはそのフォルダのルール（CLAUDE.md）を自動で読み込みます。\n'
        '物販ビジネスフォルダで起動すれば物販の専門家として動き、\n'
        '米国成長株フォルダで起動すれば株の専門家として動きます。'
    )
    doc.add_page_break()

    # ========== 第3章：AI社長の設定 ==========
    add_section_title(doc, '第3章：AI社長を設定する')
    add_subsection(doc, '3-1. CLAUDE.mdとは何か')
    add_body(doc,
        'CLAUDE.mdとは、Claudeへの指示書です。\n'
        '会社で言えば「就業規則」のようなものです。\n\n'
        'AI経営本部のルートフォルダに置くCLAUDE.mdが「AI社長」です。\n'
        'このファイルを読み込んだClaudeは、\n'
        'あなたのビジネス全体を把握した社長として動きます。'
    )
    doc.add_paragraph()

    add_prompt_box(doc, 'プロンプト②：AI社長CLAUDE.mdの作成',
        'Desktop/AI経営本部/CLAUDE.mdに\n'
        '以下の内容を書き込んでください。\n\n'
        '# AI社長\n\n'
        'あなたは[あなたの名前]のAI経営本部を統括するAI社長です。\n\n'
        '## 担当ビジネス\n'
        '・人財育成\n'
        '・物販ビジネス\n'
        '・銀行融資\n'
        '・会社法人\n'
        '・米国成長株\n\n'
        '## 行動原則\n'
        '・全ての成果物は日本語で作成する\n'
        '・架空の事例は使わない\n'
        '・必ず根拠・データを付ける\n'
        '・品質チェックを必ず実施する\n\n'
        '## 出力フォーマット\n'
        '・文章：です・ます調\n'
        '・数字は具体的に記載する'
    )

    add_subsection(doc, '3-2. 各AIマネージャーの設定')
    add_body(doc, '各フォルダのCLAUDE.mdに役割を設定します。\n以下のプロンプトをClaude Codeに送信してください。')
    doc.add_paragraph()

    add_prompt_box(doc, 'プロンプト③：5つのAIマネージャーを一括設定',
        '以下の5つのAIマネージャーを設定してください。\n'
        '各フォルダのCLAUDE.mdに役割を書き込んでください。\n\n'
        '■人財育成マネージャー\n'
        'Desktop/AI経営本部/人財育成/CLAUDE.md\n'
        '役割：チームマネジメント・人材育成の資料作成とアドバイスを行う\n\n'
        '■物販ビジネスマネージャー\n'
        'Desktop/AI経営本部/物販ビジネス/CLAUDE.md\n'
        '役割：商品リサーチ・出品文章作成・顧客対応・売上分析を行う\n\n'
        '■銀行融資マネージャー\n'
        'Desktop/AI経営本部/銀行融資/CLAUDE.md\n'
        '役割：融資申請書・事業計画書・収支計画書の作成サポートを行う\n\n'
        '■会社法人マネージャー\n'
        'Desktop/AI経営本部/会社法人/CLAUDE.md\n'
        '役割：議事録・契約書ドラフト・法人経営に関する資料作成を行う\n\n'
        '■米国成長株マネージャー\n'
        'Desktop/AI経営本部/米国成長株/CLAUDE.md\n'
        '役割：NVIDIA・Apple・Microsoft等の企業リサーチと分析を行う\n\n'
        '全て完了したら確認してください。'
    )
    doc.add_page_break()

    # ========== 第4章：SKILLファイル ==========
    add_section_title(doc, '第4章：SKILLファイルで記憶を管理する')
    add_subsection(doc, '4-1. SKILLファイルとは何か')
    add_body(doc,
        'SKILLファイルとは、Claudeへの取扱説明書です。\n\n'
        '例えば「私の文章はです・ます調で書いてください」\n'
        '「架空の事例は絶対に使わないでください」\n'
        'というルールをファイルに書いておくことで、\n'
        '毎回同じ品質のアウトプットが得られます。\n\n'
        '会話履歴に頼らないため、\n'
        'コストを大幅に削減できるのも大きなメリットです。'
    )
    doc.add_paragraph()

    add_prompt_box(doc, 'プロンプト④：基本SKILLファイルの作成',
        'Desktop/AI経営本部/に\n'
        '以下の3つのSKILLファイルを作成してください。\n\n'
        '【SKILL_自己紹介.md】\n'
        '# 私のプロフィール\n'
        '・名前：[あなたの名前]\n'
        '・職業：[あなたの職業]\n'
        '・ビジネスの目標：[目標]\n'
        '・得意なこと：[得意なこと]\n'
        '・苦手なこと：[苦手なこと]\n\n'
        '【SKILL_文章スタイル.md】\n'
        '# 文章のルール\n'
        '・文体：です・ます調\n'
        '・架空の事例は使わない\n'
        '・数字は具体的に書く\n'
        '・改行を多用して読みやすくする\n\n'
        '【SKILL_品質チェック.md】\n'
        '# 品質チェック基準\n'
        '・事実と異なる内容は含まれていないか\n'
        '・根拠・データは明記されているか\n'
        '・日本語として自然か\n'
        '・100点満点で採点して報告すること'
    )

    add_point_box(doc, 'SKILLファイルの使い方',
        'Claude Codeで作業を始める前に以下を入力するだけです：\n\n'
        'cat ~/Desktop/AI経営本部/SKILL_自己紹介.md\n\n'
        'これだけでClaudeがあなたのことを全て把握した上で作業を開始します。\n'
        'このコマンドをCLAUDE.mdに書いておけば自動で読み込まれます。'
    )
    doc.add_page_break()

    # ========== 第5章：GitHubでクラウド管理 ==========
    add_section_title(doc, '第5章：GitHubでクラウド管理する')
    add_subsection(doc, '5-1. なぜGitHubで管理するか')

    tbl3 = doc.add_table(rows=4, cols=2)
    set_table_border(tbl3)
    for i, h in enumerate(['メリット', '内容']):
        tbl3.rows[0].cells[i].text = h
        set_cell_bg(tbl3.rows[0].cells[i], '1A1A2E')
        for para in tbl3.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)
                run.font.name = 'メイリオ'
    for i, (merit, content) in enumerate([
        ('バックアップ', 'パソコンが壊れても全データが安全に保管される'),
        ('履歴管理', 'いつ・何を変更したか全て記録。間違えても元に戻せる'),
        ('どこからでもアクセス', 'iPhone・別のパソコンからでも操作できる'),
    ]):
        row = tbl3.rows[i+1]
        row.cells[0].text = merit
        row.cells[1].text = content
        if i % 2 == 0:
            set_cell_bg(row.cells[0], 'F0F0FF')
            set_cell_bg(row.cells[1], 'F0F0FF')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = 'メイリオ'
    doc.add_paragraph()

    add_prompt_box(doc, 'プロンプト⑤：GitHubと連携する',
        'GitHubと連携してAI経営本部を\n'
        'クラウドで管理できるようにしてください。\n\n'
        'GitHubのユーザー名：[あなたのGitHubユーザー名]\n'
        'リポジトリ名：ai-keiei-honbu\n\n'
        '以下を全て実行してください。\n'
        '・Gitの初期設定\n'
        '・リポジトリの作成\n'
        '・AI経営本部フォルダをGitHubに接続\n'
        '・最初のコミット＆プッシュ\n'
        '・.gitignoreの設定（APIキーが漏れないよう）\n\n'
        '全て完了したらGitHubのURLを教えてください。'
    )

    add_subsection(doc, '5-2. iPhoneから操作する方法')
    add_step_table(doc, [
        ('①', 'iPhoneでSafariを開く', ''),
        ('②', 'https://claude.ai を開いてログインする', 'チャット画面が開く'),
        ('③', 'Claude Codeと同じようにプロンプトを入力する', 'そのまま作業できる'),
    ])
    add_point_box(doc, 'iPhoneから使う時の注意',
        'iPhoneのClaude.aiでは、ローカルのファイル操作はできません。\n'
        'GitHubに保存したファイルの確認・指示出し・アイデア整理に使いましょう。\n'
        '実際のファイル操作はMacのClaude Codeで行います。'
    )
    doc.add_page_break()

    # ========== 第6章：毎日の使い方 ==========
    add_section_title(doc, '第6章：毎日の使い方')
    add_subsection(doc, '6-1. 朝のルーティン（所要時間：約5分）')
    add_body(doc, 'Claude Codeを起動して以下のプロンプトを送信するだけです。')
    doc.add_paragraph()

    add_prompt_box(doc, '朝のルーティンプロンプト（毎日使う）',
        'おはようございます。\n'
        '今日の作業を始めます。\n\n'
        'まず以下を読み込んでください。\n'
        'cat ~/Desktop/AI経営本部/CLAUDE.md\n'
        'cat ~/Desktop/AI経営本部/SKILL_自己紹介.md\n\n'
        '読み込んだら今日やるべきことを\n'
        '優先順位をつけて提案してください。'
    )

    add_subsection(doc, '6-2. ビジネス別の使い方')
    tbl4 = doc.add_table(rows=6, cols=3)
    set_table_border(tbl4)
    for i, h in enumerate(['やりたいこと', 'フォルダ', '入力するプロンプト例']):
        tbl4.rows[0].cells[i].text = h
        set_cell_bg(tbl4.rows[0].cells[i], '1A1A2E')
        for para in tbl4.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)
                run.font.name = 'メイリオ'
    usage_data = [
        ('株式のリサーチをしたい', '米国成長株/', 'NVIDIAの最新決算を分析して\n投資判断のサポート情報をまとめてください'),
        ('融資資料を作りたい', '銀行融資/', '以下の事業内容で融資申請書の\nドラフトを作成してください\n事業内容：[内容を入力]'),
        ('商品の出品文章を作りたい', '物販ビジネス/', '以下の商品の出品文章を\nメルカリ向けに作成してください\n商品：[商品名と状態]'),
        ('議事録を作りたい', '会社法人/', '以下の会議内容を\n議事録形式でまとめてください\n[会議内容を貼り付け]'),
        ('スタッフ育成資料を作りたい', '人財育成/', '新入スタッフ向けの\n[テーマ]の研修資料を作成してください'),
    ]
    for i, (want, folder, prompt_ex) in enumerate(usage_data):
        row = tbl4.rows[i+1]
        row.cells[0].text = want
        row.cells[1].text = folder
        row.cells[2].text = prompt_ex
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, 'F0F0FF')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'メイリオ'
    doc.add_paragraph()
    doc.add_page_break()

    # ========== 第7章：便利なコマンド ==========
    add_section_title(doc, '第7章：覚えておくと便利なコマンド')
    add_subsection(doc, '7-1. Claude Codeの便利コマンド')
    doc.add_paragraph()

    tbl5 = doc.add_table(rows=5, cols=3)
    set_table_border(tbl5)
    for i, h in enumerate(['コマンド', '意味', '使うタイミング']):
        tbl5.rows[0].cells[i].text = h
        set_cell_bg(tbl5.rows[0].cells[i], '1A1A2E')
        for para in tbl5.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)
                run.font.name = 'メイリオ'
    cmd_data = [
        ('/compact', '会話を要約・軽量化する', '長い作業の後・動きが遅くなった時'),
        ('/clear', '会話をリセットする', '新しい作業を始める前'),
        ('claude --version', 'バージョン確認', '正常に動いているか確認する時'),
        ('claude login', 'ログインし直す', 'ログインが切れた時'),
    ]
    for i, (cmd, meaning, timing) in enumerate(cmd_data):
        row = tbl5.rows[i+1]
        row.cells[0].text = cmd
        row.cells[1].text = meaning
        row.cells[2].text = timing
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, 'F0F0FF')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'メイリオ'
    doc.add_paragraph()

    add_subsection(doc, '7-2. よく使うGitHubコマンド')
    add_prompt_box(doc, '作業内容を保存する（毎日の習慣にする）',
        'cd ~/Desktop/AI経営本部\n'
        'git add .\n'
        'git commit -m "今日の作業内容をひと言で"\n'
        'git push origin main'
    )
    add_warning_box(doc, 'gitコマンドは必ずAI経営本部フォルダで実行してください。他のフォルダで実行しないよう注意してください。')
    doc.add_page_break()

    # ========== 第8章：チェックリスト ==========
    add_section_title(doc, '第8章：完成確認チェックリスト')
    add_body(doc, '以下の全項目にチェックが入ったら、AI経営本部の完成です。')
    doc.add_paragraph()

    add_subsection(doc, '8-1. 環境構築チェック')
    add_checklist(doc, [
        'Claudeアカウントを作成した（Maxプラン推奨）',
        'GitHubアカウントを作成した',
        'Claude Codeをインストールした',
        'claude loginでログインした',
        'claudeと入力してClaude Codeが起動した',
    ])
    doc.add_paragraph()

    add_subsection(doc, '8-2. AI経営本部構築チェック')
    add_checklist(doc, [
        'デスクトップにAI経営本部フォルダが作成された',
        'AI社長CLAUDE.mdが設定された',
        '5つのAIマネージャーが設定された',
        'SKILL_自己紹介.mdが作成された',
        'SKILL_文章スタイル.mdが作成された',
        'SKILL_品質チェック.mdが作成された',
        'GitHubと連携できた',
        'iPhoneからclaude.aiにアクセスできた',
    ])
    doc.add_paragraph()

    add_subsection(doc, '8-3. 実際に動作確認')
    add_checklist(doc, [
        '物販ビジネスマネージャーに商品文章を作らせた',
        '米国成長株マネージャーにリサーチをさせた',
        '作成したファイルがGitHubに保存された',
        '/compactと/clearコマンドを試した',
        '毎日のルーティンプロンプトを実行した',
    ])
    doc.add_paragraph()

    add_subsection(doc, '8-4. よくあるトラブルと解決方法')
    tbl6 = doc.add_table(rows=5, cols=3)
    set_table_border(tbl6)
    for i, h in enumerate(['症状', '原因', '解決方法']):
        tbl6.rows[0].cells[i].text = h
        set_cell_bg(tbl6.rows[0].cells[i], '1A1A2E')
        for para in tbl6.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)
                run.font.name = 'メイリオ'
    trouble_data = [
        ('claudeと入力しても起動しない', 'Node.js未インストール', 'https://nodejs.org からNode.jsをインストールして再試行'),
        ('動きが遅くなった', 'コンテキストが溜まりすぎ', '/compactを入力して会話を軽量化する'),
        ('ログインが切れた', '認証の期限切れ', 'claude loginを再実行する'),
        ('Gitにプッシュできない', 'GitHub認証エラー', '「GitHubの認証エラーが出ました。解決してください」とClaudeに入力'),
    ]
    for i, (symptom, cause, solution) in enumerate(trouble_data):
        row = tbl6.rows[i+1]
        row.cells[0].text = symptom
        row.cells[1].text = cause
        row.cells[2].text = solution
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, 'F0F0FF')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'メイリオ'

    doc.add_page_break()

    # ========== 付録 ==========
    add_section_title(doc, '付録：コピペ用プロンプト集')
    add_body(doc, '以下のプロンプトは全てコピペしてClaude Codeに貼り付けるだけで使えます。')
    doc.add_paragraph()

    prompts = [
        ('毎朝の起動プロンプト',
         'おはようございます。今日の作業を始めます。\n'
         'まず以下を読み込んでください。\n'
         'cat ~/Desktop/AI経営本部/CLAUDE.md\n'
         'cat ~/Desktop/AI経営本部/SKILL_自己紹介.md\n'
         '読み込んだら今日やるべきことを優先順位をつけて提案してください。'),
        ('エラーが出た時のプロンプト',
         'エラーが出ました。以下がエラーの内容です。\n'
         '[エラーメッセージをここに貼り付け]\n'
         '原因と解決方法を教えてください。'),
        ('作業が途中で止まった時のプロンプト',
         '続きをお願いします。'),
        ('GitHubに保存するプロンプト',
         '今日の作業内容をGitHubに保存してください。\n'
         'cd ~/Desktop/AI経営本部\n'
         'git add .\n'
         'git commit -m "[今日の作業内容をひと言で]"\n'
         'git push origin main'),
        ('新しいSKILLファイルを作るプロンプト',
         'Desktop/AI経営本部/SKILL_[名前].mdを作成してください。\n'
         '内容は以下です。\n'
         '[SKILLファイルに書きたい内容を入力]'),
    ]
    for title, prompt_text in prompts:
        add_prompt_box(doc, title, prompt_text)

    os.makedirs(SAVE_DIR, exist_ok=True)
    filepath = os.path.join(
        SAVE_DIR,
        '【受講生用】Claudeだけで作るAI経営本部_構築マニュアル.docx'
    )
    doc.save(filepath)
    print(f'✅ マニュアル完成：{filepath}')
    return filepath

def main():
    print('=' * 60)
    print('Claude単体版 AI経営本部構築マニュアル 生成開始')
    print('=' * 60)
    filepath = create_manual()
    import subprocess
    subprocess.run(['bash', '-c',
        'cd ~/Desktop/AI経営本部 && '
        'git add AI活用講義/講義資料/ && '
        'git commit -m "Claude単体版AI経営本部構築マニュアル完成" && '
        'git push origin salesletter-development'
    ])
    print('✅ GitHubプッシュ完了')
    print(f'\n完成ファイル：{filepath}')

if __name__ == '__main__':
    main()
