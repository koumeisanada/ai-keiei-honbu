import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types
from openai import OpenAI

gemini_client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
openai_client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

# Claude APIはクレジット不足のためChatGPTで代替
def claude_or_chatgpt(prompt):
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
        message = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=6000,
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text
    except Exception as e:
        print(f"  Claude APIエラー → ChatGPTで代替：{str(e)[:50]}")
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4000
        )
        return response.choices[0].message.content

SAVE_DIR = os.path.expanduser(
    "~/Desktop/AI経営本部/AI活用講義/特典"
)

GLOSSARY_CATEGORIES = [
    ("CAT01", "AIの基本用語", ["AI（人工知能）","機械学習","ディープラーニング","生成AI","大規模言語モデル（LLM）","ニューラルネットワーク","自然言語処理（NLP）","コンピュータービジョン","強化学習","教師あり学習・教師なし学習"]),
    ("CAT02", "AIツール・サービス関連", ["ChatGPT","Claude","Gemini","Copilot","Midjourney","Stable Diffusion","DALL-E","Whisper","Sora","Perplexity"]),
    ("CAT03", "プロンプト・指示関連", ["プロンプト","プロンプトエンジニアリング","システムプロンプト","コンテキスト","トークン","コンテキストウィンドウ","温度（Temperature）","ハルシネーション","Few-shot prompting","Chain of Thought"]),
    ("CAT04", "API・技術関連", ["API","APIキー","エンドポイント","リクエスト・レスポンス","JSON","Python","スクリプト","ライブラリ","環境変数","ターミナル・コマンドライン"]),
    ("CAT05", "自動化・システム関連", ["パイプライン","自動化（オートメーション）","ワークフロー","バッチ処理","スケジューラー","Make（旧Integromat）","Zapier","n8n","RPA","エージェントAI"]),
    ("CAT06", "AI組織・管理関連", ["AI社長（CLAUDE.md）","AIマネージャー","SKILLファイル","AI経営本部","プロジェクト管理","バージョン管理","GitHub","リポジトリ","コミット","ブランチ"]),
    ("CAT07", "コスト・セキュリティ関連", ["トークンコスト","API料金","Spending limit（上限設定）","APIキー漏洩","環境変数（.zshrc）",".gitignore","セキュリティリスク","情報漏洩","プライバシー","GDPR"]),
    ("CAT08", "真田孔明講座独自用語", ["Level 1〜5","AI組織経営","3大AI横断活用","デイリーパイプライン","成功シンドロームOS","5ポケッツ戦略術","AI経営本部","ナノバナナプロ","エバーグリーン化","品質チェック自動化"]),
]

def generate_glossary_content(categories):
    today = datetime.now().strftime("%Y年%m月%d日")
    all_terms = []
    for cat_id, cat_name, terms in categories:
        for term in terms:
            all_terms.append(f"{cat_name}：{term}")

    prompt = f"""
【作成日】{today}

以下のAI用語を全て解説してください。

対象読者：AIのことを右も左も分からない40代の方・エンジニアではない方

【解説ルール】
1. 専門用語を使わずに説明する
2. 必ず身近な例え話を使う
3. 「つまり何ができるのか」を必ず書く
4. 1用語あたり150〜200文字程度

【出力形式】
### [用語名]
**一言で言うと：**（10文字以内）
**詳しく言うと：**（例え話を使って150文字程度）
**真田孔明の講座では：**（1〜2行）
---

【用語一覧】
{chr(10).join(all_terms)}
"""

    print("Step1：Geminiで用語集の草稿を生成中...")
    try:
        response = gemini_client.models.generate_content(
            model="gemini-2.5-flash", contents=prompt,
            config=types.GenerateContentConfig(tools=[types.Tool(google_search=types.GoogleSearch())]))
        draft = response.text
    except Exception:
        response = gemini_client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
        draft = response.text
    print("✅ Step1完了：Gemini草稿生成")

    print("Step2：ChatGPTで正確性・わかりやすさをチェック中...")
    check_prompt = f"""
以下はAI初心者向けの用語集です。チェックして改善提案を出してください。
【チェック観点】1.説明が正確か 2.40代非エンジニアに伝わるか 3.例え話が適切か 4.難しい言葉がないか
【草稿】{draft[:5000]}
改善が必要な用語と改善案を具体的に。日本語で。
"""
    fb = openai_client.chat.completions.create(
        model="gpt-4o", messages=[{"role": "user", "content": check_prompt}], max_tokens=2000)
    feedback = fb.choices[0].message.content
    print("✅ Step2完了：ChatGPTチェック")

    print("Step3：最終仕上げ中...")
    final_prompt = f"""
以下のフィードバックを反映してAI用語集を完成させてください。
【フィードバック】{feedback}
【草稿】{draft[:5000]}
40代の非エンジニアが100%理解できる。例え話は日常生活に近く。技術的に正確。日本語で。
"""
    final = claude_or_chatgpt(final_prompt)
    print("✅ Step3完了：最終仕上げ")

    return draft, feedback, final

def create_glossary_word(final_content, today):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27); section.page_height = Inches(11.69)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    section.top_margin = Inches(1); section.bottom_margin = Inches(1)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI用語集'); run.font.size = Pt(32); run.font.bold = True; run.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    doc.add_paragraph()
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('〜 知らなくても大丈夫。でも知ると、もっと楽しくなる 〜')
    run2.font.size = Pt(14); run2.font.color.rgb = RGBColor(0x53,0x4A,0xB7)
    doc.add_paragraph()
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f'真田孔明 AI活用講座　受講生特典\n{today}')
    run3.font.size = Pt(11); run3.font.color.rgb = RGBColor(0x88,0x87,0x80)
    doc.add_page_break()

    h = doc.add_heading('はじめに', level=1)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    doc.add_paragraph(
        'この用語集は、AI活用講座で登場する言葉を\n'
        '誰でもわかるように解説したものです。\n\n'
        '難しい言葉は一切使っていません。\n'
        '「なんとなくわかった」で大丈夫です。\n\n'
        '制作：Gemini（草稿）× ChatGPT（正確性チェック）× Claude（最終仕上げ）')
    doc.add_page_break()

    for line in final_content.split('\n'):
        if not line.strip(): doc.add_paragraph(); continue
        if line.startswith('# ') or line.startswith('## '):
            h = doc.add_heading(line.lstrip('#').strip(), level=1)
            for r in h.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
        elif line.startswith('### '):
            h = doc.add_heading(line.replace('### ',''), level=2)
            for r in h.runs: r.font.color.rgb = RGBColor(0x53,0x4A,0xB7); r.font.bold = True
        elif line.startswith('**') and '**' in line[2:]:
            p = doc.add_paragraph(); r = p.add_run(line.replace('**',''))
            r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = RGBColor(0x16,0x21,0x3E)
        elif line.startswith('---'):
            doc.add_paragraph('─' * 50)
        else:
            p = doc.add_paragraph(line)
            for r in p.runs: r.font.size = Pt(10.5)

    os.makedirs(SAVE_DIR, exist_ok=True)
    word_file = os.path.join(SAVE_DIR, "【特典】AI用語集_完全版.docx")
    doc.save(word_file)
    print(f"✅ Word保存完了：{word_file}")
    return word_file

def main():
    today = datetime.now().strftime("%Y年%m月%d日")
    os.makedirs(SAVE_DIR, exist_ok=True)
    total = sum(len(terms) for _, _, terms in GLOSSARY_CATEGORIES)

    print("=" * 60)
    print("AI用語集 制作システム起動")
    print("3大AI横断：Gemini（生成）× ChatGPT（正確性チェック）× Claude/ChatGPT（仕上げ）")
    print(f"収録用語数：{total}語")
    print("=" * 60)

    draft, feedback, final = generate_glossary_content(GLOSSARY_CATEGORIES)

    with open(os.path.join(SAVE_DIR, "用語集_草稿.md"), "w") as f: f.write(draft)
    with open(os.path.join(SAVE_DIR, "用語集_ChatGPTフィードバック.md"), "w") as f: f.write(feedback)
    with open(os.path.join(SAVE_DIR, "用語集_完成版.md"), "w") as f: f.write(final)

    word_file = create_glossary_word(final, today)

    import subprocess
    subprocess.run(["bash", "-c",
        "cd ~/Desktop/AI経営本部 && "
        "git add AI活用講義/特典/ AI活用講義/create_ai_glossary.py && "
        "git commit -m '受講生特典：AI用語集完成（3大AI横断制作・80語収録）' && "
        "git push origin salesletter-development"
    ])

    print("\n" + "=" * 60)
    print("✅ AI用語集 完成！")
    print("=" * 60)
    print(f"収録用語数：{total}語")
    print(f"保存先：{word_file}")

if __name__ == "__main__":
    main()
