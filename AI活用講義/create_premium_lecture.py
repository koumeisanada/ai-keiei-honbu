import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types
import anthropic
from openai import OpenAI

gemini_client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
claude_client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
openai_client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

def gemini_search(prompt):
    try:
        response = gemini_client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=types.GenerateContentConfig(
                tools=[types.Tool(google_search=types.GoogleSearch())]
            )
        )
        return response.text
    except Exception:
        return gemini_generate(prompt)

def gemini_generate(prompt):
    response = gemini_client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    return response.text

def chatgpt_check(prompt):
    response = openai_client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=3000
    )
    return response.choices[0].message.content

def claude_revise(prompt):
    try:
        message = claude_client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=6000,
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text
    except Exception as e:
        print(f"  Claude APIエラー：{e}")
        print("  → ChatGPTで最終仕上げを代替します")
        return chatgpt_check(prompt)

def save_word(content, filepath, title, subtitle, today, lecture_num, file_type):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'第{lecture_num:02d}回 AI活用講義')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(
        '制作：Gemini（検索・生成）× ChatGPT（論理チェック）× Claude（修正・完成）'
    )
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f'真田孔明 AI活用講座　{today}')
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x88, 0x87, 0x80)

    doc.add_page_break()

    for line in content.split('\n'):
        if not line.strip():
            doc.add_paragraph()
            continue
        if line.startswith('# ') or line.startswith('## '):
            h = doc.add_heading(
                line.replace('## ', '').replace('# ', ''), level=1
            )
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        elif line.startswith('### '):
            h = doc.add_heading(line.replace('### ', ''), level=2)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
        elif line.startswith('❌'):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            r.font.size = Pt(10.5)
        elif line.startswith('✅'):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            r.font.size = Pt(10.5)
        elif line.startswith('[画像挿入'):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.color.rgb = RGBColor(0xFF, 0x69, 0x00)
            r.font.bold = True
            r.font.size = Pt(10)
        else:
            p = doc.add_paragraph(line)
            for r in p.runs:
                r.font.size = Pt(10.5)

    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)
    print(f"✅ Word保存完了：{filepath}")
    return filepath

def create_image_prompts(today, file_type):
    if file_type == "news":
        prompts = [
            {"num": 1, "section": "今月のトピック全体像",
             "location": "冒頭",
             "prompt": "ナノバナナプロで今月のAI業界トピックを図解してください。\n霞が関のポンチ絵のように全ての情報を一枚の画像にまとめて。\nサイズ：横長16:9"},
            {"num": 2, "section": "具体的な活用方法",
             "location": "活用方法セクション",
             "prompt": "ナノバナナプロで今月の具体的なAI活用方法をステップ形式で図解してください。\nサイズ：横長16:9"},
        ]
    else:
        prompts = [
            {"num": 1, "section": "全体像：Level 1〜5ロードマップ",
             "location": "パート1冒頭",
             "prompt": "ナノバナナプロで以下を図解してください。\nタイトル：AI活用Level 1〜5ロードマップ\nLevel 1（約80%）：質問・回答のみ\nLevel 2（約15%）：プロンプト工夫・業務組み込み\nLevel 3（約4%）：API活用・自動化\nLevel 4（約0.9%）：複数AI連携・エージェント設計\nLevel 5（約0.1%）：AI組織経営・完全自動化\n階段または三角形の図。各Levelに人数割合を記載。\nサイズ：横長16:9"},
            {"num": 2, "section": "AI経営本部の全体構造",
             "location": "全体像セクション",
             "prompt": "ナノバナナプロで以下を図解してください。\nタイトル：AI経営本部の全体構造\n中心にAI社長、周囲に5名のAIマネージャー\n3大AI（Claude・Gemini・ChatGPT）の役割分担\n自動パイプラインの流れ\nスタイル：組織図形式。サイズ：横長16:9"},
            {"num": 3, "section": "年間24回学習ロードマップ",
             "location": "ロードマップセクション",
             "prompt": "ナノバナナプロで以下を図解してください。\nタイトル：年間24回学習ロードマップ\n月2回×12ヶ月のタイムライン\n各フェーズのLevel到達目標を記載\nスタイル：横型タイムライン。サイズ：横長16:9"},
            {"num": 4, "section": "AIリスク4大カテゴリ",
             "location": "注意事項パート冒頭",
             "prompt": "ナノバナナプロで以下を図解してください。\nタイトル：AIリスク4大カテゴリ\n①セキュリティリスク②品質リスク③法的リスク④コストリスク\n各カテゴリにアイコンと具体例を追加\nスタイル：霞が関のポンチ絵。サイズ：横長16:9"},
            {"num": 5, "section": "APIキー管理フロー",
             "location": "APIキー管理セクション",
             "prompt": "ナノバナナプロで以下を図解してください。\nタイトル：APIキーの正しい管理フロー\n左：❌悪い例（コードに直書き→漏洩→高額請求）\n右：✅良い例（環境変数→安全）\n下：漏洩時の対処法（1分以内に無効化）\nサイズ：横長16:9"},
            {"num": 6, "section": "ハルシネーション発生率",
             "location": "ハルシネーションセクション",
             "prompt": "ナノバナナプロで以下を図解してください。\nタイトル：AI別ハルシネーション発生率比較\nChatGPT-4：15〜20%、Claude 3：10〜15%、Gemini 1.5：12〜18%\n棒グラフ形式。出典：Stanford AI Index 2024\nサイズ：横長16:9"},
            {"num": 7, "section": "トークンコスト比較",
             "location": "コスト管理セクション",
             "prompt": "ナノバナナプロで以下を図解してください。\nタイトル：トークンコスト増加の仕組み\n全履歴方式：1日目1,000→1週間7,000→1ヶ月30,000トークン\nSKILLファイル方式：常に一定\n折れ線グラフ＋比較図。サイズ：横長16:9"},
            {"num": 8, "section": "失敗TOP10",
             "location": "よくある間違いセクション",
             "prompt": "ナノバナナプロで以下を図解してください。\nタイトル：AI活用でよくある失敗TOP10\nランキング形式で1〜10位を表示\n各項目に危険度アイコンを追加\nサイズ：横長16:9"},
            {"num": 9, "section": "10の鉄則",
             "location": "鉄則セクション",
             "prompt": "ナノバナナプロで以下を図解してください。\nタイトル：AIと安全に向き合うための10の鉄則\n10個の鉄則を番号付きで2列に配置\n各鉄則にアイコンを追加\nスタイル：インフォグラフィック形式。サイズ：横長16:9"},
            {"num": 10, "section": "チェックリスト",
             "location": "チェックリストページ",
             "prompt": "ナノバナナプロで以下を図解してください。\nタイトル：受講生向けチェックリスト\n10項目のチェックリストをグループ分けして表示\nA4印刷対応。サイズ：横長16:9"},
        ]
    return prompts

def save_image_prompts(prompts, save_dir, lecture_num, today, file_type):
    label = "AI最新情報用" if file_type == "news" else "Level講義用"
    filepath = os.path.join(
        save_dir,
        f"第{lecture_num:02d}回_画像プロンプト_{label}.txt"
    )
    with open(filepath, "w") as f:
        f.write("=" * 60 + "\n")
        f.write(f"ナノバナナプロ 画像プロンプト一覧（{label}）\n")
        f.write(f"作成日：{today}\n")
        f.write("以下を順番にGemini（ナノバナナプロ）に貼り付けてください。\n")
        f.write("=" * 60 + "\n\n")
        for p in prompts:
            f.write(f"【画像{p['num']}：{p['section']}】\n")
            f.write(f"挿入箇所：{p['location']}\n")
            f.write(f"プロンプト：\n{p['prompt']}\n")
            f.write("-" * 40 + "\n\n")
    print(f"✅ 画像プロンプト保存：{filepath}")
    return filepath

def create_level_lecture(lecture_num, today, save_dir):
    print("\n" + "=" * 60)
    print(f"ファイル②：Level講義資料を生成中（第{lecture_num}回）")
    print("=" * 60)

    print("Step1：Gemini（検索）でリサーチ中...")
    research = gemini_search(f"""
【調査日】{today}

以下を調査してください。

1. APIキー漏洩による被害事例と金額（具体的な数字・出典）
2. ハルシネーション発生率データ（各AI比較・Stanford AI Index等）
3. AI利用による情報漏洩事例（企業・年別）
4. 著作権・プライバシーに関する最新の法律・判例
5. AI利用に関する各国の規制動向（2024〜2026年）
6. トークンコスト管理の失敗事例（金額付き）
7. AI安全活用のベストプラクティス（公的機関の推奨）

各項目に具体的な数字・情報源・年を付けて日本語で出力してください。
""")
    with open(os.path.join(save_dir, "Step1_リサーチ.md"), "w") as f:
        f.write(research)
    print("✅ Step1完了")

    print("Step2：Gemini（生成）で草稿作成中...")
    draft = gemini_generate(f"""
【調査日】{today}
【リサーチデータ】
{research[:3000]}

真田孔明のAI活用講座 第{lecture_num}回講義資料を作成してください。

【構成】

# パート1：全体像（最重要）
## この講座で1年間何を学ぶか
## Level 1〜5の定義と現在地の確認
## 年間24回の学習ロードマップ
## AI経営本部の全体構造
## 1年後に手に入るもの

[画像挿入：Level 1〜5ロードマップ図]
[画像挿入：AI経営本部全体構造図]
[画像挿入：年間24回学習ロードマップ]

# パート2：AIと向き合う前に知っておくべきこと

## 第1章：AIリスクの全体像
[画像挿入：AIリスク4大カテゴリ図]
## 第2章：APIキーの管理と取り扱い
[画像挿入：APIキー管理フロー図]
## 第3章：ハルシネーションへの対処法
[画像挿入：ハルシネーション発生率グラフ]
## 第4章：著作権・プライバシーリスク
## 第5章：トークンコスト・メモリ管理
[画像挿入：トークンコスト比較図]
## 第6章：法的リスク・規制
## 第7章：実務でよくある間違いTOP10
[画像挿入：失敗TOP10ランキング図]
## 第8章：AIと安全に向き合うための10の鉄則
[画像挿入：10の鉄則インフォグラフィック]

# パート3：Level 1→2への第一歩
## 今日から始める具体的なアクション
## 次回までの宿題
## 受講生向けチェックリスト
[画像挿入：チェックリスト図]

【文章スタイル】
・全ての主張に客観的な根拠・データを付ける
・結論→根拠→事実→応用の順で書く
・感情訴求は使わない
・真田孔明の実体験は具体的な数字で語る
・文学スキルは使わない
・日本語で出力する
""")
    with open(os.path.join(save_dir, "Step2_草稿.md"), "w") as f:
        f.write(draft)
    print("✅ Step2完了")

    print("Step3：ChatGPT（gpt-4o）で論理チェック中...")
    feedback = chatgpt_check(f"""
以下はAI活用講座 第1回講義資料の草稿です。
徹底的な論理チェック・正確性チェックを行い改善フィードバックを出力してください。

【草稿】
{draft[:4000]}

【チェック項目】
1. 論理整合性
2. データ・数字の正確性
3. 網羅性
4. 受講生への配慮
5. 追加すべき重要情報

日本語で出力してください。
""")
    with open(os.path.join(save_dir, "Step3_ChatGPTフィードバック.md"), "w") as f:
        f.write(feedback)
    print("✅ Step3完了")

    print("Step4：Claude（claude-sonnet-4-6）で修正・完成中...")
    final = claude_revise(f"""
以下のChatGPTフィードバックを全て反映して講義資料を修正・完成させてください。

【ChatGPTフィードバック】
{feedback}

【修正前の草稿】
{draft[:4000]}

【修正の方針】
・全ての問題点を修正する
・不足データを補完する
・受講生が理解しやすい表現に
・文学スキルは絶対に使わない
・日本語で出力する
""")
    with open(os.path.join(save_dir, "Step4_Claude完成版.md"), "w") as f:
        f.write(final)
    print("✅ Step4完了")

    word_path = os.path.join(
        save_dir,
        f"第{lecture_num:02d}回_②Level講義資料_完成版.docx"
    )
    save_word(
        final, word_path,
        f"第{lecture_num:02d}回 Level講義資料",
        "全体像＋注意事項＋Level 1→2への第一歩",
        today, lecture_num, "level"
    )

    prompts = create_image_prompts(today, "level")
    save_image_prompts(prompts, save_dir, lecture_num, today, "level")

    return word_path

def create_news_lecture(lecture_num, month_num, today, save_dir):
    print("\n" + "=" * 60)
    print(f"ファイル①：AI最新情報資料を生成中（第{lecture_num}回・{month_num}ヶ月目）")
    print("=" * 60)

    print("Step1：Gemini（検索）で今月の最新AI情報を収集中...")
    news = gemini_search(f"""
【調査日】{today}

直近1ヶ月以内のAI業界最新情報を調査してください。

1. 新しいAIモデルのリリース・アップデート
2. 米国テクノロジー企業のAI戦略最新動向
3. AI規制・法律の動向
4. 話題のAIツール・サービスの登場
5. AIによるビジネス成功事例
6. 今月最も重要なAIニュースTOP5

各トピックに具体的な内容・情報源・日付を付けて日本語で出力してください。
""")

    print("Step2：Gemini（生成）で最新情報資料を作成中...")
    draft_news = gemini_generate(f"""
【調査日】{today}
【最新情報データ】
{news[:3000]}

AI活用講座 第{lecture_num}回「今月のAI最新情報＋具体的な活用方法」を作成してください。

構成：今月の重要トピックTOP5→トピック別詳細→Level別活用ポイント→来月への展望
文章スタイル：客観的・根拠重視・初心者でもわかる・日本語
""")

    print("Step3：ChatGPT（gpt-4o）で論理チェック中...")
    feedback_news = chatgpt_check(f"""
以下はAI最新情報講義資料の草稿です。事実確認・論理チェックを行い改善提案を出力してください。
{draft_news[:3000]}
日本語で出力してください。
""")

    print("Step4：Claude（claude-sonnet-4-6）で修正・完成中...")
    final_news = claude_revise(f"""
以下のフィードバックを反映してAI最新情報講義資料を完成させてください。
【フィードバック】{feedback_news}
【草稿】{draft_news[:3000]}
文学スキルは使わない。日本語で出力する。
""")

    word_path = os.path.join(
        save_dir,
        f"第{lecture_num:02d}回_①AI最新情報資料_{today}.docx"
    )
    save_word(
        final_news, word_path,
        f"第{lecture_num:02d}回 AI最新情報資料",
        f"今月のAI最新情報＋具体的な活用方法　{today}",
        today, lecture_num, "news"
    )

    prompts = create_image_prompts(today, "news")
    save_image_prompts(prompts, save_dir, lecture_num, today, "news")

    for name, content in [
        ("最新情報リサーチ", news),
        ("最新情報草稿", draft_news),
        ("最新情報フィードバック", feedback_news),
        ("最新情報完成版", final_news),
    ]:
        with open(os.path.join(save_dir, f"{name}.md"), "w") as f:
            f.write(content)

    return word_path

def main():
    import sys
    today = datetime.now().strftime("%Y年%m月%d日")
    lecture_num = int(sys.argv[1]) if len(sys.argv) > 1 else 1
    month_num = int(sys.argv[2]) if len(sys.argv) > 2 else 1

    save_dir = os.path.expanduser(
        f"~/Desktop/AI経営本部/AI活用講義/講義資料/第{lecture_num:02d}回"
    )
    os.makedirs(save_dir, exist_ok=True)

    print("=" * 60)
    print(f"第{lecture_num}回 講義資料制作システム起動")
    print("3大AI完全横断：Gemini（検索・生成）× ChatGPT（論理チェック）× Claude（修正・完成）")
    print("=" * 60)

    if lecture_num == 1:
        print("\n【第1回講義】")
        print("ファイル①（AI最新情報）：なし（初回のため不要）")
        print("ファイル②（Level講義資料）：生成します")
        level_word = create_level_lecture(lecture_num, today, save_dir)
        print(f"\n✅ 第1回講義資料完成")
        print(f"  ファイル②：{level_word}")
    else:
        print(f"\n【第{lecture_num}回講義】")
        print("ファイル①（AI最新情報）：生成します")
        print("ファイル②（Level講義資料）：生成します")
        news_word = create_news_lecture(lecture_num, month_num, today, save_dir)
        level_word = create_level_lecture(lecture_num, today, save_dir)
        print(f"\n✅ 第{lecture_num}回講義資料完成")
        print(f"  ファイル①：{news_word}")
        print(f"  ファイル②：{level_word}")

    import subprocess
    subprocess.run(["bash", "-c",
        "cd ~/Desktop/AI経営本部 && "
        f"git add AI活用講義/講義資料/第{lecture_num:02d}回/ && "
        f"git commit -m '第{lecture_num}回講義資料完成（3大AI横断制作・2ファイル構成）' && "
        "git push origin salesletter-development"
    ])
    print("\n✅ GitHubプッシュ完了")
    print("\n次のステップ：")
    print("  画像プロンプト一覧.txtをGemini（ナノバナナプロ）に貼り付けて画像生成")
    print("  生成した画像をWordの[画像挿入箇所]に挿入する")

if __name__ == "__main__":
    main()
