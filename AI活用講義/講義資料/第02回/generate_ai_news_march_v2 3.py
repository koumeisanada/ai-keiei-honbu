#!/usr/bin/env python3
"""第02回 ①AI最新情報 2026年3月まとめ Word生成 v2
create_project_report.py / create_ai_glossary.py と同一パターン"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SAVE_DIR = os.path.expanduser(
    "~/Desktop/AI経営本部/AI活用講義/講義資料/第02回"
)

# ============================================================
# ヘルパー関数（create_project_report.py と同一パターン）
# ============================================================

def set_table_border(table, color='1A1A2E'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_font(cell, size=Pt(10), bold=False, color=None):
    for para in cell.paragraphs:
        para.paragraph_format.space_after = Pt(2)
        for run in para.runs:
            run.font.name = 'Meiryo'
            run.font.size = size
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')


def add_cell_text(cell, text, size=Pt(10), bold=False, color=None, align=None):
    """セルにテキストを設定（runs経由で確実にフォント制御）"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Meiryo'
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')


def add_title(doc, text, size=Pt(28)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Meiryo'
    run.font.size = size
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')
    return p


def add_subtitle(doc, text, size=Pt(14)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Meiryo'
    run.font.size = size
    run.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')
    return p


def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f'■ {text}')
    run.font.name = 'Meiryo'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')
    return p


def add_subsection_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f'◆ {text}')
    run.font.name = 'Meiryo'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')
    return p


def add_body(doc, text, size=Pt(10.5)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Meiryo'
    run.font.size = size
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    indent = '　' * level
    run = p.add_run(f'{indent}・{text}')
    run.font.name = 'Meiryo'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')
    return p


def add_hr(doc, color='1A1A2E'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_colored_table(doc, headers, rows, header_bg='1A1A2E', border_color='1A1A2E'):
    """ヘッダー色付き・交互行色分けテーブル"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table, border_color)

    # ヘッダー行
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, header_bg)
        add_cell_text(cell, h, size=Pt(9.5), bold=True,
                      color=RGBColor(0xFF, 0xFF, 0xFF),
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    # データ行（交互色分け）
    for ri, row_data in enumerate(rows):
        bg = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            set_cell_bg(cell, bg)
            add_cell_text(cell, val, size=Pt(9))

    doc.add_paragraph()
    return table


def add_highlight_box(doc, title, lines, bg_color='EDE7F6', border_color='534AB7'):
    """強調ボックス"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table, border_color)
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)

    # タイトル
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'◉ {title}')
    run.font.name = 'Meiryo'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')

    # コンテンツ
    for line in lines:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        r = p2.add_run(line)
        r.font.name = 'Meiryo'
        r.font.size = Pt(10)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')

    doc.add_paragraph()


def add_warning_box(doc, title, lines):
    add_highlight_box(doc, title, lines, bg_color='FFEBEE', border_color='C0392B')


def add_success_box(doc, title, lines):
    add_highlight_box(doc, title, lines, bg_color='E8F5E9', border_color='27AE60')


def add_page_number(doc):
    """フッターにページ番号を挿入"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = p.add_run('真田孔明 AI活用講座 第2回  |  ')
    run1.font.name = 'Meiryo'
    run1.font.size = Pt(8)
    run1.font.color.rgb = RGBColor(0x95, 0x95, 0x95)
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run2 = p.add_run()
    run2.font.name = 'Meiryo'
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)


# ============================================================
# メイン生成
# ============================================================
def create_document():
    doc = Document()

    # ページ設定
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # フッター
    add_page_number(doc)

    # ========================================================
    # 表紙
    # ========================================================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # ネイビーバー
    bar = doc.add_table(rows=1, cols=1)
    bar.style = 'Table Grid'
    bar.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(bar, '1A1A2E')
    bc = bar.cell(0, 0)
    set_cell_bg(bc, '1A1A2E')
    add_cell_text(bc, 'SANADA KOMEI AI ACADEMY\n真田孔明のAI活用講座',
                  size=Pt(11), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    add_subtitle(doc, '第2回講義資料　前半パート', size=Pt(13))
    doc.add_paragraph()
    add_title(doc, '2026年3月\nAI最新情報まとめ', size=Pt(30))
    doc.add_paragraph()
    add_subtitle(doc, '〜 個人・ビジネスオーナーが今すぐ使える最新AI動向 〜', size=Pt(13))
    doc.add_paragraph()
    doc.add_paragraph()

    # パープルバー
    bar2 = doc.add_table(rows=1, cols=1)
    bar2.style = 'Table Grid'
    bar2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(bar2, '534AB7')
    bc2 = bar2.cell(0, 0)
    set_cell_bg(bc2, '534AB7')
    add_cell_text(bc2, '作成日：2026年3月23日  |  AI活用講義 月額制サブスクリプション',
                  size=Pt(10), color=RGBColor(0xFF, 0xFF, 0xFF),
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_conf = p_conf.add_run('CONFIDENTIAL  |  For Members Only')
    r_conf.font.name = 'Meiryo'
    r_conf.font.size = Pt(9)
    r_conf.font.italic = True
    r_conf.font.color.rgb = RGBColor(0x95, 0x95, 0x95)
    r_conf.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')

    doc.add_page_break()

    # ========================================================
    # 第1章：重要トピックTOP10
    # ========================================================
    add_section_title(doc, '第1章：3月の重要トピックTOP10')
    add_hr(doc)
    doc.add_paragraph()

    add_body(doc, '2026年3月に起きたAI業界の重要ニュースを、ビジネスインパクト順にランキングしました。')

    add_colored_table(doc,
        ['順位', '重要度', 'トピック名', '概要', 'ビジネスへの影響'],
        [
            ('1', '★★★★★', 'NVIDIA「Vera Rubin」\nプラットフォーム発表',
             'GTC 2026で次世代AIプラットフォーム発表。\n推論コスト最大1/10、学習GPU数1/4に削減。\nGroq 3 LPU統合で推論特化型へ転換',
             'AI商用化コストの劇的低減。\nエージェント型AIの普及加速。\n2027年AI HW需要1兆ドル規模'),
            ('2', '★★★★★', 'EU AI Act\n本格施行',
             '世界初の包括的AI規制が3月に本格施行。\nリスクベース4段階分類。\n域外適用あり（日本企業も対象）',
             '違反時最大3500万ユーロ or 売上7%罰金。\n高リスクAI利用企業は適合性評価必須。\n日本企業も対応が急務'),
            ('3', '★★★★★', 'Claude Opus 4.6\nリリース',
             '100万トークンコンテキスト対応。\nエージェントチーム機能で複数AI協調。\nAdaptive Thinking・音声モード追加',
             '大規模文書分析が一度に可能。\n開発タスクの並列自動化。\nPro $20/月で利用可能'),
            ('4', '★★★★☆', 'Gemini 3\nGmail/Chrome統合',
             'AI InboxでGmail自動管理。\nGemini in Chromeで50言語対応。\nエージェント機能で業務自動化',
             'メール処理時間の大幅削減。\nブラウザ内でリサーチ・要約。\nGWSユーザーは即活用可'),
            ('5', '★★★★☆', 'ChatGPT\nGPT-5.3/5.4統合',
             'GPT-5.1系廃止→5.3/5.4に統合。\nGPT-5.4 mini無料展開。\nインタラクティブ学習機能追加',
             '無料でも最新モデル利用可能。\nハルシネーション低減・日本語向上。\n教育・研修分野の活用拡大'),
            ('6', '★★★★☆', 'Sora 2 vs Runway\n動画生成AI三つ巴',
             'Sora 2は最長25秒+音声同期。\nRunway Gen-4はフレーム精密制御。\nVeo 3.1はYouTube Shorts連携',
             'プロモーション動画の自作が現実的。\nSNS短尺動画制作コスト激減。\n個人でもプロ級映像が可能'),
            ('7', '★★★★☆', '日本「AI推進法」\n成立',
             '国の基本計画策定が法定化。\n1兆円規模の国産AI開発支援。\nソフトバンク・PFN等が新会社設立',
             '国産AIエコシステムの充実。\n政府支援による開発加速。\n2〜3年後に成果が期待'),
            ('8', '★★★☆☆', 'Suno v5 / ACE-Step\n音楽生成AI進化',
             'Suno v5はプロ品質の楽曲生成。\nACE-Step v1.5はOSS・4GB VRAM動作。\nWMGがライセンス契約→合法化',
             'BGM・ジングル制作の革命。\n月額ゼロの音楽生成が可能。\n商用ライセンスが明確化'),
            ('9', '★★★☆☆', 'ElevenLabs v3\n音声生成革命',
             '70言語で感情表現（笑い・ささやき等）。\nScribe v2で150ms低遅延音声認識。\n日本語エラー率5%未満',
             'ナレーション・ポッドキャスト制作。\n多言語動画吹替が個人で可能。\n音声チャットボット構築（ノーコード）'),
            ('10', '★★★☆☆', 'Amazon AIインフラ\n巨額投資',
             'スペインに180億ユーロ追加投資。\nTrainiumチップ開発ラボ開設。\nAI推論コスト削減競争が激化',
             'AWSのAI環境がさらに高性能化。\nクラウドAI利用コスト低下。\n自社チップでNVIDIA依存軽減'),
        ],
        header_bg='1A1A2E'
    )

    doc.add_page_break()

    # ========================================================
    # 第2章：GAFAM+NVIDIA
    # ========================================================
    add_section_title(doc, '第2章：NVIDIA・Apple・Microsoft・Meta・Google 最新動向')
    add_hr(doc)
    doc.add_paragraph()

    add_body(doc, '2026年3月の各社の主要な動きを比較表で整理しました。')

    add_colored_table(doc,
        ['企業', '3月の主要発表', 'AI戦略の方向性', '個人への影響'],
        [
            ('NVIDIA',
             '・GTC 2026「Vera Rubin」発表\n・Groq 3 LPU（推論特化チップ）\n・DLSS 5（ニューラルレンダリング）\n・AI HW需要1兆ドル規模予測',
             'GPU中心→統合型AIインフラへ転換\n推論コスト1/10を目指す\nPrefill/Decode分業型',
             'AI利用コストが今後大幅低下\nゲーム・映像のAI高画質化\n推論効率向上で応答速度改善'),
            ('Apple',
             '・iOS 18.2でSiri×ChatGPT連携\n・Apple Intelligence版Siri準備中\n・デバイス上AI処理の強化',
             'デバイス統合型AI戦略\nプライバシー重視のオンデバイスAI\n外部AI（ChatGPT）との連携',
             'iPhone設定からSiri拡張を有効化\n音声でAIに質問が可能に\n写真分析・創作文章もSiri経由で'),
            ('Microsoft',
             '・Copilot全製品への統合加速\n・自律型AIエージェント機能強化\n・Azure AI推論サービス拡充',
             'AI co-pilot戦略の深化\nOffice×AI統合でエンタープライズ\n7つのAIトレンド提言を発表',
             'Word/Excel/TeamsでAI活用\nCopilot無料版で日常タスク効率化\n会議要約・資料作成の自動化'),
            ('Meta',
             '・Llama 4系モデルの強化継続\n・Meta AI単体アプリ展開\n・ARグラス Gen 2（2026年予定）',
             'オープンソースAI戦略\nSNS×AI統合\nAR/VR×AI融合デバイス開発',
             'Meta AIアプリで無料AI利用\nIG/WhatsApp内でAI活用\nARグラスで次世代体験'),
            ('Google',
             '・Gemini 3でGmail/Chrome大型更新\n・Gemini 3.1 Flash-Liteプレビュー\n・Gemini in Chrome 50言語対応\n・Google Home音声応答40%改善',
             'Google全製品へのAI統合\nエージェント機能の本格展開\nPersonal Intelligence構想',
             'Gmail AI Inboxで自動管理\nChromeサイドパネルでAI活用\n検索・メール・文書をAI強化'),
        ],
        header_bg='534AB7'
    )

    doc.add_page_break()

    # ========================================================
    # 第3章：AI規制・法律
    # ========================================================
    add_section_title(doc, '第3章：AI規制・法律の動向')
    add_hr(doc)
    doc.add_paragraph()

    add_body(doc, '2026年3月は世界的にAI規制が大きく動いた月です。日本・EU・米国の最新状況を整理します。')

    add_colored_table(doc,
        ['国・地域', '法規制名', '施行/発表', '主な内容', '違反時の罰則'],
        [
            ('EU', 'EU AI Act\n（AI規制法）', '2026年3月\n本格施行',
             '・リスクベース4段階分類\n・高リスクAIに厳格な規制\n・域外適用あり（日本企業も対象）\n・透明性・説明責任の義務化',
             '最大3500万ユーロ\nまたは全世界年間\n売上高の7%'),
            ('日本', 'AI推進法', '2026年3月\n成立',
             '・国の基本計画策定を法定化\n・研究開発推進\n・人権侵害リスクへの調査指導\n・1兆円規模の国産AI開発支援',
             '罰則規定なし\n（推進法のため）'),
            ('日本', 'AI開発\nガイドライン', '2026年3月6日\n発表',
             '・倫理基準の強化\n・国際協調の推進\n・Human-in-the-Loop推奨\n・AI事業者の責任明確化',
             '罰則なし\n（ガイドライン）'),
            ('米国', 'AI国家政策\nフレームワーク', '2026年3月20日\n立法提言公表',
             '・州規制乱立の防止\n・連邦統一基準の設定方針\n・イノベーションと規制のバランス',
             '未定\n（立法提言段階）'),
            ('日本', '著作権\nガイドライン', '2025年1月改訂\n（文化庁）',
             '・AI学習は原則許容（30条の4）\n・生成物は通常の著作権法適用\n・「本質的特徴の直接感得」で侵害',
             '通常の著作権法に\n基づく損害賠償'),
        ],
        header_bg='C0392B'
    )

    add_warning_box(doc, 'ビジネスオーナーへの警告', [
        '・EU市場でビジネスを行う場合、域外適用により日本企業も規制対象',
        '・高リスクAIシステムを利用している場合、適合性評価と透明性確保が急務',
        '・AI利用ポリシーの策定・更新は今月中に着手すべき',
        '・「知らなかった」は通用しない。継続的なキャッチアップ体制を構築すること',
    ])

    doc.add_page_break()

    # ========================================================
    # 第4章：注目AIツール
    # ========================================================
    add_section_title(doc, '第4章：今月登場した注目AIツール・サービス')
    add_hr(doc)
    doc.add_paragraph()

    add_colored_table(doc,
        ['ツール名', 'カテゴリ', '特徴', '活用方法', '料金', '難易度'],
        [
            ('GPT-5.4 mini', 'テキスト生成', 'OpenAI最新の軽量版\n無料ユーザーにも展開\nハルシネーション低減', '質問応答・文章作成\nアイデア出し', '無料', 'Level 1'),
            ('Claude Opus 4.6', 'テキスト/分析', '100万トークン対応\nエージェントチーム\nAdaptive Thinking', '長文PDF分析\n契約書比較\nコード開発', '無料枠あり\nPro $20/月', 'Level 1-3'),
            ('Gemini in Chrome', 'ブラウザAI', '50言語対応\n記事要約・下書き\n画像生成対応', 'Webリサーチ効率化\nSNS投稿下書き', '無料', 'Level 1'),
            ('Suno v5', '音楽生成', 'プロ品質の楽曲生成\n12ステム書き出し\n日本語歌詞対応', 'BGM・ジングル制作\nYouTube用音楽', '無料枠あり\n$10/月〜', 'Level 1'),
            ('ACE-Step v1.5', '音楽生成', 'オープンソース\n4GB VRAM動作\nSuno v5相当の品質', 'ローカル音楽生成\n月額コストゼロ', '完全無料', 'Level 3'),
            ('ElevenLabs v3', '音声生成', '70言語で感情表現\n150ms低遅延認識\n日本語エラー率5%未満', 'ナレーション制作\n多言語吹替', '無料枠あり\n$5/月〜', 'Level 2'),
            ('Sora 2', '動画生成', '最長25秒+音声同期\nDisney提携', 'プロモ動画\nSNS短尺動画', 'Plus $20/月〜', 'Level 2'),
            ('Runway Gen-4', '動画生成', 'フレーム精密制御\nAlephエディタ', '映像制作\n広告動画', '無料枠あり', 'Level 2'),
            ('Bing Image Creator', '画像生成', 'DALL-Eベース\n完全無料', 'SNS用画像\nブログ素材', '完全無料', 'Level 1'),
            ('Canva AI', 'デザイン', '動画生成（5回無料）\nテンプレート豊富', 'SNSデザイン\nプレゼン資料', '無料枠あり', 'Level 1'),
        ],
        header_bg='534AB7'
    )

    add_highlight_box(doc, '今すぐ使えるベスト3', [
        '1位：ChatGPT（GPT-5.4 mini）— 無料で最新AI。まず登録すべき1つ',
        '2位：Gemini in Chrome — ブラウザに統合済み。Googleユーザーは設定するだけ',
        '3位：Claude — 長文分析の王者。PDF・契約書の分析なら最強',
    ])

    doc.add_page_break()

    # ========================================================
    # 第5章：AI活用事例
    # ========================================================
    add_section_title(doc, '第5章：AIで業績を上げた企業・個人の事例')
    add_hr(doc)
    doc.add_paragraph()

    add_colored_table(doc,
        ['企業/個人', '活用内容', '成果', '使用ツール'],
        [
            ('IVRy\n（電話自動応答）', 'Geminiベースで電話自動応答\nサービスの基盤を移行', '対応精度向上\n人件費削減', 'Gemini\nエージェント機能'),
            ('イオンリテール', '商品情報登録プロセスの\n半自動化にAIエージェント導入', '登録工数の大幅削減\n入力ミスの低減', 'Gemini Enterprise'),
            ('AI経営本部\n（真田孔明）', 'メルマガ・LINE・リール動画の\n日次自動生成パイプライン構築', '毎日5種類の成果物を自動生成\nAPI費用を最小化', 'Gemini API\nClaude Code'),
            ('個人ブロガー', 'ChatGPT + Canva AIで\nブログ＋SNS画像を同時制作', '記事制作時間が1/3に短縮\nPV数が月間2倍に増加', 'ChatGPT\nCanva AI'),
            ('教育機関', 'ChatGPTのインタラクティブ\n学習機能で数学・科学教育', '理解度テスト平均点向上\n個別学習支援の充実', 'ChatGPT\n（無料版対応）'),
        ],
        header_bg='27AE60'
    )

    doc.add_page_break()

    # ========================================================
    # 第6章：今日から使える活用方法3選
    # ========================================================
    add_section_title(doc, '第6章：今日から使える！具体的な活用方法3選')
    add_hr(doc)
    doc.add_paragraph()

    # 活用法1
    add_subsection_title(doc, '活用法1：3大AIを比較して使い分ける')
    add_success_box(doc, 'Step by Step（所要時間15分 / 無料 / Level 1）', [
        'Step 1：ChatGPT（chatgpt.com）に無料登録する',
        'Step 2：Claude（claude.ai）に無料登録する',
        'Step 3：Gemini（gemini.google.com）に無料登録する',
        'Step 4：同じ質問を3つのAIに投げて、回答を比較する',
        '　例：「40代サラリーマンが副業で月10万円稼ぐ方法を5つ提案して」',
        'Step 5：用途別に使い分けを確立する',
        '　・日常の質問 → ChatGPT（最も汎用的）',
        '　・長文分析・コード → Claude（100万トークン対応）',
        '　・Google連携・検索 → Gemini（Gmail/Chrome統合）',
    ])

    # 活用法2
    add_subsection_title(doc, '活用法2：Gemini in Chromeで情報収集を自動化')
    add_success_box(doc, 'Step by Step（所要時間5分 / 無料 / Level 1）', [
        'Step 1：Google Chromeを最新版に更新する',
        'Step 2：設定 → 実験的機能 → Gemini in Chromeを有効化',
        'Step 3：任意のWebサイトでサイドパネルのGeminiアイコンをクリック',
        'Step 4：「この記事を300文字で要約して」と入力',
        'Step 5：「この記事の要点を箇条書きにして」と追加質問',
        '',
        '応用：',
        '　・競合サイト分析 →「このサービスの強みと弱みを分析して」',
        '　・英語記事の翻訳 →「この記事を日本語で要約して」',
        '　・SNS投稿下書き →「この記事をSNS投稿用にリライトして」',
    ])

    # 活用法3
    add_subsection_title(doc, '活用法3：AIで毎日のメール処理時間を半分にする')
    add_success_box(doc, 'Step by Step（初期設定10分 / 無料 / Level 1）', [
        'Step 1：GmailでGemini機能を有効化する',
        'Step 2：受信メールを開き「返信の下書き」ボタンをクリック',
        'Step 3：Geminiが提案する返信案を確認・修正して送信',
        'Step 4：長文メールは「要約」機能で3行にまとめてもらう',
        'Step 5：定型返信はChatGPTでテンプレートを事前作成',
        '',
        '時短効果：',
        '　・メール1通あたりの処理時間：5分 → 2分',
        '　・1日30通のメール処理：150分 → 60分（90分削減）',
        '　・月間で約30時間の時間創出',
    ])

    doc.add_page_break()

    # ========================================================
    # 第7章：来月の注目トレンド予測
    # ========================================================
    add_section_title(doc, '第7章：来月の注目トレンド予測（2026年4月）')
    add_hr(doc)
    doc.add_paragraph()

    add_colored_table(doc,
        ['No.', '注目トレンド', '予測される動き', 'ビジネスへの影響'],
        [
            ('1', 'NVIDIA Groq LPU\n市場導入本格化', 'Vera Rubinの提供開始時期明確化。\n推論特化チップの性能ベンチマーク公開', 'AI推論コストの価格破壊開始\nAIエージェント導入コスト大幅低下'),
            ('2', 'EU AI Act\n高リスクAI対応事例', '具体的な企業対応事例が続々登場。\nコンプライアンスツール市場が活性化', '対応先行企業の事例が参考に\nAIガバナンスコンサル需要増'),
            ('3', 'Google I/O 2026\n事前発表（5月予定）', 'Gemini 4系の噂。\nAndroid AI機能の大幅強化', '次世代AIモデルの方向性が見える\n開発者向けAPI新機能の予告'),
            ('4', '音楽生成AI\nライセンス整備', 'メジャーレーベルのAI音楽\nライセンス契約が拡大', '商用BGM制作のAI活用加速\n著作権リスクの明確化'),
            ('5', 'ChatGPT教育機能\nの拡充', 'インタラクティブ学習の対象拡大。\n企業研修向け機能の提供開始', '社内研修のAI化が加速\n個人学習の効率が大幅向上'),
        ],
        header_bg='534AB7'
    )

    add_highlight_box(doc, '4月のアクションプラン', [
        '・Google I/O 2026の事前情報に注目 → 次世代AI戦略のヒントを掴む',
        '・EU AI Act対応の先行企業事例を収集 → 自社の対応計画に反映',
        '・AI推論コスト低下のトレンドを活用 → 新規AIプロジェクトの検討開始',
    ])

    doc.add_page_break()

    # ========================================================
    # まとめ：重要ポイント10箇条
    # ========================================================
    add_section_title(doc, 'まとめ：今月の重要ポイント10箇条')
    add_hr(doc)
    doc.add_paragraph()

    add_colored_table(doc,
        ['No.', 'ポイント', '解説'],
        [
            ('1', 'NVIDIA Vera Rubinで推論コスト1/10の時代が到来', 'AIの商用化コストが劇的に低減。エージェント型AIの普及が加速する'),
            ('2', 'EU AI Actが本格施行 — 域外適用で日本企業も対象', '自社のAIリスク評価とコンプライアンス体制構築が急務'),
            ('3', 'Claude Opus 4.6の100万トークンで長文分析の常識が変わった', '書籍数冊分を一度に分析。契約書比較・議事録分析が一瞬で可能'),
            ('4', 'Gemini 3がGmail/Chromeに統合 — 日常業務が静かに革命', '設定するだけで使えるAI。メール処理・Webリサーチが自動化'),
            ('5', 'ChatGPT無料版でもGPT-5.4 miniが使える時代に', 'AIの民主化が進行。「無料でここまでできる」が驚異的なレベル'),
            ('6', '動画生成AI三つ巴 — 個人でもプロ級映像が作れる', 'Sora 2/Runway/Veo 3.1。SNS動画の制作コストがほぼゼロ'),
            ('7', '音楽生成AIが商用ライセンス時代に突入', 'WMGのSuno/Udioライセンス契約で「グレーゾーン」が解消へ'),
            ('8', '日本のAI推進法成立 — 1兆円規模の国産AI支援が始動', '規制よりも推進寄りの法律。国産AIエコシステムの充実に期待'),
            ('9', 'AIエージェントが「ツール」から「チームメンバー」に変化', '目的を与えるだけでAIが自律実行。マルチエージェント時代の幕開け'),
            ('10', '3大AI（ChatGPT/Claude/Gemini）を無料で使い分けるのが正解', '1つに固定するのは損。用途別に使い分ける「3刀流」が最強戦略'),
        ],
        header_bg='1A1A2E'
    )

    add_highlight_box(doc, '今週のアクション', [
        'Q1：今日試せるツールは何か？',
        '→ ChatGPT / Claude / Gemini の3つに無料登録し、同じ質問を投げて比較する',
        '',
        'Q2：あなたの仕事で最初に自動化できるのはどこか？',
        '→ メール返信・データ整理・レポート作成を棚卸しし、1つをAIに委任する',
        '',
        'Q3：来月までにLevel 2に上がるためにやること1つは？',
        '→ Gemini in Chromeを有効化し、毎日1回はAI要約を使う習慣をつける',
    ])

    doc.add_paragraph()

    # 次回予告
    p_next = doc.add_paragraph()
    p_next.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_next = p_next.add_run('次回予告：第2回後半パート「主要AIツールの徹底比較と使い分け」')
    r_next.font.name = 'Meiryo'
    r_next.font.size = Pt(13)
    r_next.font.bold = True
    r_next.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    r_next.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')

    doc.add_paragraph()

    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run('ご清聴ありがとうございました！')
    r_end.font.name = 'Meiryo'
    r_end.font.size = Pt(14)
    r_end.font.bold = True
    r_end.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    r_end.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')

    # ========================================================
    # 保存
    # ========================================================
    os.makedirs(SAVE_DIR, exist_ok=True)
    filepath = os.path.join(SAVE_DIR, '第02回_①AI最新情報_2026年3月まとめ.docx')
    doc.save(filepath)
    return filepath


if __name__ == '__main__':
    filepath = create_document()
    size = os.path.getsize(filepath)
    print(f'✅ 保存完了：{filepath}')
    print(f'📄 ファイルサイズ：{size:,} bytes（{size/1024:.1f} KB）')
