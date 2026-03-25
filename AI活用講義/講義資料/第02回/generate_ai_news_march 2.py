#!/usr/bin/env python3
"""第02回 ①AI最新情報 2026年3月まとめ Word生成"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# カラー定数
C_NAVY = "1A1A2E"
C_PURPLE = "534AB7"
C_WHITE = "FFFFFF"
C_LGRAY = "F5F5F5"
C_MGRAY = "E8E8E8"
C_RED = "C0392B"
C_GREEN = "27AE60"
C_ORANGE = "E67E22"
C_LPURPLE = "EDE7F6"
C_LNAVY = "E8EAF6"
C_LGREEN = "E8F5E9"
C_LRED = "FFEBEE"
C_LORANGE = "FFF3E0"

RC_NAVY = RGBColor(0x1A, 0x1A, 0x2E)
RC_PURPLE = RGBColor(0x53, 0x4A, 0xB7)
RC_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RC_TEXT = RGBColor(0x2C, 0x2C, 0x2C)
RC_RED = RGBColor(0xC0, 0x39, 0x2B)
RC_ORANGE = RGBColor(0xE6, 0x7E, 0x22)
RC_GREEN = RGBColor(0x27, 0xAE, 0x60)
RC_GRAY = RGBColor(0x6A, 0x6A, 0x6A)
RC_LGRAY = RGBColor(0x95, 0x95, 0x95)

FJ = "Meiryo"
FE = "Meiryo"

doc = Document()

# ページ設定
for s in doc.sections:
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

# スタイル
sn = doc.styles['Normal']
sn.font.name = FE
sn.font.size = Pt(10.5)
sn.font.color.rgb = RC_TEXT
sn.paragraph_format.space_after = Pt(6)
sn.paragraph_format.line_spacing = 1.5
sn.element.rPr.rFonts.set(qn('w:eastAsia'), FJ)

for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = FE
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), FJ)
    if lv == 1:
        hs.font.size = Pt(22)
        hs.font.color.rgb = RC_NAVY
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif lv == 2:
        hs.font.size = Pt(16)
        hs.font.color.rgb = RC_PURPLE
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.font.color.rgb = RC_NAVY
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

# ヘルパー関数
def set_shading(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    )

def set_cell(cell, text, bold=False, sz=Pt(10), color=RC_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.font.name = FE
    r.font.size = sz
    r.font.bold = bold
    r.font.color.rgb = color
    r.element.rPr.rFonts.set(qn('w:eastAsia'), FJ)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_run(p, text, bold=False, sz=Pt(10.5), color=RC_TEXT, italic=False):
    r = p.add_run(text)
    r.font.name = FE
    r.font.size = sz
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.element.rPr.rFonts.set(qn('w:eastAsia'), FJ)
    return r

def add_body(text, bold=False, sz=Pt(10.5), color=RC_TEXT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, bold=bold, sz=sz, color=color)
    return p

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    pPr = h._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="8" w:color="{C_PURPLE}"/>'
        f'</w:pBdr>'
    ))
    return h

def add_table_borders(t, color=C_PURPLE):
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="2" w:space="0" w:color="{C_MGRAY}"/>'
        f'<w:insideV w:val="single" w:sz="2" w:space="0" w:color="{C_MGRAY}"/>'
        f'</w:tblBorders>'
    ))

def add_data_table(headers, rows, hdr_color=C_NAVY):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    add_table_borders(t, hdr_color)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_shading(c, hdr_color)
        set_cell(c, h, bold=True, sz=Pt(9.5), color=RC_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, rd in enumerate(rows):
        bg = C_LGRAY if ri % 2 == 0 else C_WHITE
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            set_shading(c, bg)
            set_cell(c, v, sz=Pt(9))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_box(title, lines, bg=C_LPURPLE, hdr_bg=C_PURPLE, hdr_color=RC_WHITE, border=C_PURPLE):
    t = doc.add_table(rows=2, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(t, border)
    tc = t.cell(0, 0)
    set_shading(tc, hdr_bg)
    set_cell(tc, f"  {title}", bold=True, sz=Pt(12), color=hdr_color)
    cc = t.cell(1, 0)
    set_shading(cc, bg)
    cc.paragraphs[0].text = ""
    for ln in lines:
        p = cc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.4
        r = p.add_run(f"  {ln}")
        r.font.name = FE
        r.font.size = Pt(10)
        r.font.color.rgb = RC_TEXT
        r.element.rPr.rFonts.set(qn('w:eastAsia'), FJ)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_warn(title, lines):
    add_box(title, lines, bg=C_LRED, hdr_bg=C_RED, hdr_color=RC_WHITE, border=C_RED)

def add_ok(title, lines):
    add_box(title, lines, bg=C_LGREEN, hdr_bg="1B7A3D", hdr_color=RC_WHITE, border=C_GREEN)

def add_info(title, lines):
    add_box(title, lines, bg=C_LPURPLE, hdr_bg=C_PURPLE, hdr_color=RC_WHITE, border=C_PURPLE)

# フッター
section = doc.sections[0]
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = fp.add_run("真田孔明 AI活用講座 第2回  |  2026年3月 AI最新情報まとめ  |  ")
r1.font.size = Pt(8)
r1.font.color.rgb = RC_LGRAY
r1.font.name = FE
r1.element.rPr.rFonts.set(qn('w:eastAsia'), FJ)
fp._p.append(parse_xml(
    f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
    f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="6A6A6A"/></w:rPr><w:t>1</w:t></w:r>'
    f'</w:fldSimple>'
))
r2 = fp.add_run(" / ")
r2.font.size = Pt(8)
r2.font.color.rgb = RC_LGRAY
fp._p.append(parse_xml(
    f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES \\* MERGEFORMAT ">'
    f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="6A6A6A"/></w:rPr><w:t>1</w:t></w:r>'
    f'</w:fldSimple>'
))

# ================================================================
# 表紙
# ================================================================
for _ in range(3):
    doc.add_paragraph()

# ネイビーバー
tb = doc.add_table(rows=1, cols=1)
tb.alignment = WD_TABLE_ALIGNMENT.CENTER
bc = tb.cell(0, 0)
set_shading(bc, C_NAVY)
bp = bc.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(bp, "SANADA KOMEI AI ACADEMY", bold=True, sz=Pt(12), color=RC_WHITE)
add_run(bp, "\n真田孔明のAI活用講座", sz=Pt(10), color=RGBColor(0xBB, 0xBB, 0xDD))

doc.add_paragraph()
doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(tp, "第2回講義資料  前半パート", sz=Pt(14), color=RC_PURPLE, bold=True)

doc.add_paragraph()

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(tp2, "2026年3月\nAI最新情報まとめ", sz=Pt(32), color=RC_NAVY, bold=True)

doc.add_paragraph()

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(sp, "〜 個人・ビジネスオーナーが今すぐ使える最新AI動向 〜", sz=Pt(13), color=RC_GRAY)

doc.add_paragraph()
doc.add_paragraph()

# パープルバー
tb2 = doc.add_table(rows=1, cols=1)
tb2.alignment = WD_TABLE_ALIGNMENT.CENTER
bc2 = tb2.cell(0, 0)
set_shading(bc2, C_PURPLE)
bp2 = bc2.paragraphs[0]
bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(bp2, "作成日：2026年3月23日  |  AI活用講義 月額制サブスクリプション", sz=Pt(10), color=RC_WHITE)

doc.add_paragraph()

mp = doc.add_paragraph()
mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(mp, "CONFIDENTIAL  |  For Members Only", sz=Pt(9), color=RC_LGRAY, italic=True)

doc.add_page_break()

# ================================================================
# 第1章：3月の重要トピックTOP10
# ================================================================
add_h("第1章：3月の重要トピックTOP10", level=1)

add_body("2026年3月に起きたAI業界の重要ニュースを、ビジネスインパクト順にランキングしました。")

add_data_table(
    ["順位", "重要度", "トピック名", "概要", "ビジネスへの影響"],
    [
        ("1", "★★★★★", "NVIDIA「Vera Rubin」\nプラットフォーム発表", "GTC 2026で次世代AIプラットフォーム発表。\n推論コスト最大1/10、学習GPU数1/4に削減。\nGroq 3 LPU統合で推論特化型アーキテクチャへ転換", "AI商用化コストの劇的低減。\nエージェント型AIの普及を後押し。\n2027年までにAI HW需要1兆ドル規模"),
        ("2", "★★★★★", "EU AI Act\n本格施行", "世界初の包括的AI規制が3月に本格施行。\nリスクベースの4段階分類。\n違反時最大3500万ユーロ or 売上7%の罰金", "EU域外企業にも適用（域外適用）。\n高リスクAI利用企業は適合性評価必須。\n日本企業も対応が急務"),
        ("3", "★★★★★", "Claude Opus 4.6\nリリース", "100万トークンコンテキスト対応。\nエージェントチーム機能で複数AI協調。\nAdaptive Thinking・音声モード追加", "大規模文書分析が一度に可能。\n開発タスクの並列自動化。\nPro $20/月で利用可能"),
        ("4", "★★★★☆", "Gemini 3\nGmail/Chrome統合", "AI InboxでGmail自動管理。\nGemini in Chromeで50言語対応。\nエージェント機能で業務自動化", "メール処理時間の大幅削減。\nブラウザ内でリサーチ・要約・下書き。\nGoogle Workspaceユーザーは即活用可"),
        ("5", "★★★★☆", "ChatGPT GPT-5.3/5.4\nモデル統合", "GPT-5.1系廃止→5.3/5.4に統合。\nGPT-5.4 mini無料展開。\nインタラクティブ学習機能追加", "無料ユーザーも最新モデル利用可能。\nハルシネーション低減・日本語向上。\n教育・研修分野の活用拡大"),
        ("6", "★★★★☆", "Sora 2 vs Runway Gen-4\n動画生成AI三つ巴", "Sora 2は最長25秒+音声同期。\nRunway Gen-4はフレーム精密制御。\nVeo 3.1はYouTube Shorts連携", "プロモーション動画の自作が現実的に。\nSNS短尺動画制作コスト激減。\n個人でもプロ級映像が可能"),
        ("7", "★★★★☆", "日本「AI推進法」\n成立", "国の基本計画策定が法定化。\n1兆円規模の国産AI開発支援。\nソフトバンク・PFN等が新会社設立", "国産AIエコシステムの充実。\n政府支援による開発加速。\n2〜3年後に成果が期待される"),
        ("8", "★★★☆☆", "Suno v5 / ACE-Step\n音楽生成AI進化", "Suno v5はプロ品質の楽曲生成。\nACE-Step v1.5はOSSで4GB VRAM動作。\nWMGがライセンス契約→合法化", "BGM・ジングル制作の革命。\n月額ゼロの音楽生成が可能。\n商用利用のライセンスが明確化"),
        ("9", "★★★☆☆", "ElevenLabs v3\n音声生成革命", "70言語で感情表現（笑い・ささやき等）。\nScribe v2で150ms低遅延音声認識。\n日本語エラー率5%未満", "ナレーション・ポッドキャスト制作。\n多言語動画の吹替が個人で可能。\n音声チャットボット構築（ノーコード）"),
        ("10", "★★★☆☆", "Amazon AIインフラ\n巨額投資", "スペインに180億ユーロ追加投資。\nTrainiumチップ開発ラボ開設。\nAI推論コスト削減競争が激化", "AWSのAI環境がさらに高性能化。\nクラウドAI利用コストの低下。\n自社開発チップでNVIDIA依存軽減"),
    ],
    hdr_color=C_NAVY
)

doc.add_page_break()

# ================================================================
# 第2章：NVIDIA・Apple・Microsoft・Meta・Google
# ================================================================
add_h("第2章：NVIDIA・Apple・Microsoft・Meta・Google 最新動向", level=1)

add_body("2026年3月の各社の主要な動きを比較表で整理しました。", bold=True)

add_data_table(
    ["企業", "3月の主要発表", "AI戦略の方向性", "個人への影響"],
    [
        ("NVIDIA", "・GTC 2026で「Vera Rubin」発表\n・Groq 3 LPU（推論特化チップ）\n・DLSS 5（ニューラルレンダリング）\n・AI HW需要1兆ドル規模予測", "GPU中心→統合型AIインフラへ転換\n推論コスト1/10を目指す\nPrefill/Decode分業型アーキテクチャ", "AI利用コストが今後大幅低下\nゲーム・映像のAI高画質化\n推論効率向上で応答速度改善"),
        ("Apple", "・iOS 18.2でSiri×ChatGPT連携公開\n・Apple Intelligence版Siri準備中\n・デバイス上AI処理の強化", "デバイス統合型AI戦略\nプライバシー重視のオンデバイスAI\n外部AI（ChatGPT）との連携", "iPhone設定からSiri拡張を有効化\n音声でAIに質問が可能に\n写真分析・創作文章もSiri経由で"),
        ("Microsoft", "・Copilot全製品への統合加速\n・自律型AIエージェント機能強化\n・Azure AI推論サービス拡充", "AI co-pilot戦略の深化\nOffice×AI統合でエンタープライズ攻略\n7つのAIトレンド提言を発表", "Word/Excel/TeamsでAI活用\nCopilot無料版で日常タスク効率化\n会議要約・資料作成の自動化"),
        ("Meta", "・Llama 4系モデルの強化継続\n・Meta AI単体アプリ展開\n・ARグラス Gen 2（2026年予定）", "オープンソースAI戦略\nSNS×AI統合（FB/IG/WhatsApp）\nAR/VR×AI融合デバイス開発", "Meta AIアプリで無料AI利用\nInstagram/WhatsApp内でAI活用\nARグラスで次世代体験"),
        ("Google", "・Gemini 3でGmail/Chrome大型更新\n・Gemini 3.1 Flash-Liteプレビュー\n・Gemini in Chrome 50言語対応\n・Google Home音声応答40%改善", "Google全製品へのAI統合\nエージェント機能の本格展開\nPersonal Intelligence構想", "Gmail AI Inboxで受信トレイ自動管理\nChromeサイドパネルでAI活用\n検索・メール・ドキュメントをAI強化"),
    ],
    hdr_color=C_PURPLE
)

doc.add_page_break()

# ================================================================
# 第3章：AI規制・法律の動向
# ================================================================
add_h("第3章：AI規制・法律の動向", level=1)

add_body("2026年3月は世界的にAI規制が大きく動いた月です。日本・EU・米国の最新状況を整理します。")

add_data_table(
    ["国・地域", "法規制名", "施行/発表", "主な内容", "違反時の罰則"],
    [
        ("EU", "EU AI Act\n（AI規制法）", "2026年3月\n本格施行", "・リスクベース4段階分類\n・高リスクAIに厳格な規制\n・域外適用あり（日本企業も対象）\n・透明性・説明責任の義務化", "最大3500万ユーロ\nまたは全世界年間\n売上高の7%"),
        ("日本", "AI推進法", "2026年3月\n成立", "・国の基本計画策定を法定化\n・研究開発推進\n・人権侵害リスクへの調査指導\n・1兆円規模の国産AI開発支援", "罰則規定なし\n（推進法のため）\nただし監査・取引先\n審査で参照開始"),
        ("日本", "AI開発\nガイドライン", "2026年3月6日\n発表", "・倫理基準の強化\n・国際協調の推進\n・Human-in-the-Loop推奨\n・AI事業者の責任明確化", "罰則なし\n（ガイドライン）\n業界標準として\n準拠が求められる"),
        ("米国", "AI国家政策\nフレームワーク", "2026年3月20日\n立法提言公表", "・州規制乱立の防止\n・連邦統一基準の設定方針\n・イノベーションと規制のバランス\n・AI安全性の確保", "未定\n（立法提言段階）\n今後の議会審議で\n具体化見込み"),
        ("日本", "著作権\nガイドライン", "2025年1月改訂\n（文化庁）", "・AI学習は原則許容（30条の4）\n・生成物は通常の著作権法適用\n・「本質的特徴の直接感得」で侵害\n・知財推進計画2026を夏に策定", "通常の著作権法に\n基づく損害賠償\n確定判決はまだなし\n（2026年3月時点）"),
    ],
    hdr_color=C_RED
)

add_warn("ビジネスオーナーへの警告", [
    "・EU市場でビジネスを行う場合、EU AI Actの域外適用により日本企業も規制対象",
    "・高リスクAIシステムを利用している場合、適合性評価と透明性確保が急務",
    "・AI利用ポリシーの策定・更新は今月中に着手すべき",
    "・「知らなかった」は通用しない。AI規制の継続的なキャッチアップ体制を構築すること",
])

doc.add_page_break()

# ================================================================
# 第4章：注目AIツール・サービス
# ================================================================
add_h("第4章：今月登場した注目AIツール・サービス", level=1)

add_data_table(
    ["ツール名", "カテゴリ", "特徴", "活用方法", "料金", "難易度"],
    [
        ("GPT-5.4 mini", "テキスト生成", "OpenAI最新モデルの軽量版\n無料ユーザーにも展開\nハルシネーション低減", "日常の質問応答\n文章作成・要約\nアイデア出し", "無料", "Level 1"),
        ("Claude Opus 4.6", "テキスト生成\n分析", "100万トークン対応\nエージェントチーム機能\nAdaptive Thinking", "長文PDF分析\n契約書比較\n大規模コード開発", "無料枠あり\nPro $20/月", "Level 1-3"),
        ("Gemini in Chrome", "ブラウザAI", "50言語対応\n記事要約・下書き作成\n画像生成対応", "Webリサーチ効率化\nSNS投稿下書き\n翻訳・要約", "無料", "Level 1"),
        ("Gemini 3.1\nFlash-Lite", "テキスト生成", "高速・低コストモデル\nプレビュー版公開\nAPI利用向け", "大量テキスト処理\nチャットボット構築\nデータ分析", "低コスト", "Level 2"),
        ("Suno v5", "音楽生成", "プロ品質の楽曲生成\n12ステム書き出し\n日本語歌詞対応", "BGM・ジングル制作\nYouTube用音楽\nポッドキャスト音楽", "無料枠あり\n$10/月〜", "Level 1"),
        ("ACE-Step v1.5", "音楽生成", "オープンソース（無料）\n4GB VRAM動作\nSuno v5相当の品質", "ローカルで音楽生成\n月額コストゼロ\nカスタマイズ自由", "完全無料\n（OSS）", "Level 3"),
        ("ElevenLabs v3", "音声生成", "70言語で感情表現\n笑い・ささやき対応\n150ms低遅延認識", "ナレーション制作\n多言語吹替\n音声チャットボット", "無料枠あり\n$5/月〜", "Level 2"),
        ("Sora 2", "動画生成", "最長25秒の動画生成\n音声・効果音同期\nDisney提携", "プロモーション動画\nSNS短尺動画\nイメージ映像", "Plus $20/月〜\n（無料枠廃止）", "Level 2"),
        ("Runway Gen-4", "動画生成", "フレーム精密制御\nAlephビデオエディタ\n実写映像AI編集", "映像制作\n広告動画\n実写合成", "無料枠あり\n有料プラン", "Level 2"),
        ("Bing Image\nCreator", "画像生成", "DALL-Eベース\n完全無料\nMicrosoftアカウントで利用", "SNS用画像作成\nブログ素材\nプレゼン資料", "完全無料", "Level 1"),
    ],
    hdr_color=C_PURPLE
)

add_info("今すぐ使えるベスト3", [
    "1位：ChatGPT（GPT-5.4 mini）— 無料で最新AI。まず登録すべき1つ",
    "2位：Gemini in Chrome — ブラウザに統合済み。Googleユーザーは設定するだけ",
    "3位：Claude — 長文分析の王者。PDF・契約書の分析なら最強",
])

doc.add_page_break()

# ================================================================
# 第5章：AIで業績を上げた企業・個人の事例
# ================================================================
add_h("第5章：AIで業績を上げた企業・個人の事例", level=1)

add_data_table(
    ["企業/個人", "活用内容", "成果", "使用ツール"],
    [
        ("IVRy\n（電話自動応答）", "Geminiベースで電話自動応答\nサービスの基盤を移行", "対応精度向上\n人件費削減\nコールセンター業務効率化", "Gemini\nエージェント機能"),
        ("イオンリテール", "商品情報登録プロセスの\n半自動化にAIエージェント導入", "登録工数の大幅削減\n入力ミスの低減\n従業員の創造的業務への集中", "Gemini Enterprise\nパイロット導入"),
        ("AI経営本部\n（真田孔明）", "メルマガ・LINE・リール動画の\n日次自動生成パイプライン構築", "毎日5種類の成果物を自動生成\n集客コンテンツの量産体制確立\nAPI費用を最小化（SKILLファイル方式）", "Gemini API\nClaude Code\nCLAUDE.md"),
        ("個人ブロガー\n（事例報告）", "ChatGPT + Canva AIで\nブログ記事＋SNS画像を同時制作", "記事制作時間が1/3に短縮\nSNS投稿頻度が3倍に\nPV数が月間2倍に増加", "ChatGPT\nCanva AI"),
        ("教育機関", "ChatGPTのインタラクティブ\n学習機能で数学・科学教育", "生徒の理解度テスト平均点向上\n視覚的説明で抽象概念の理解促進\n個別学習支援の充実", "ChatGPT\n（無料版対応）"),
    ],
    hdr_color=C_GREEN
)

doc.add_page_break()

# ================================================================
# 第6章：今日から使える活用方法3選
# ================================================================
add_h("第6章：今日から使える！具体的な活用方法3選", level=1)

# 活用法1
add_h("活用法1：3大AIを比較して使い分ける", level=2)

add_ok("Step by Step", [
    "Step 1：ChatGPT（chatgpt.com）に無料登録する",
    "Step 2：Claude（claude.ai）に無料登録する",
    "Step 3：Gemini（gemini.google.com）に無料登録する",
    "Step 4：同じ質問を3つのAIに投げて、回答を比較する",
    "",
    "例：「40代サラリーマンが副業で月10万円稼ぐ方法を5つ提案して」",
    "",
    "Step 5：用途別に使い分けを確立する",
    "  ・日常の質問 → ChatGPT（最も汎用的）",
    "  ・長文分析・コード → Claude（100万トークン対応）",
    "  ・Google連携・検索 → Gemini（Gmail/Chrome統合）",
    "",
    "所要時間：15分  |  コスト：無料  |  難易度：Level 1",
])

# 活用法2
add_h("活用法2：Gemini in Chromeで情報収集を自動化", level=2)

add_ok("Step by Step", [
    "Step 1：Google Chromeを最新版に更新する",
    "Step 2：Chromeの設定 → 実験的機能 → Gemini in Chromeを有効化",
    "Step 3：任意のWebサイトを開き、サイドパネルのGeminiアイコンをクリック",
    "Step 4：「この記事を300文字で要約して」と入力",
    "Step 5：「この記事の要点を箇条書きにして」と追加質問",
    "",
    "応用：",
    "  ・競合サイトの分析 → 「このサービスの強みと弱みを分析して」",
    "  ・英語記事の翻訳 → 「この記事を日本語で要約して」",
    "  ・SNS投稿の下書き → 「この記事をSNS投稿用にリライトして」",
    "",
    "所要時間：5分  |  コスト：無料  |  難易度：Level 1",
])

# 活用法3
add_h("活用法3：AIで毎日のメール処理時間を半分にする", level=2)

add_ok("Step by Step", [
    "Step 1：GmailでGemini機能を有効化する（設定→全般→スマート機能）",
    "Step 2：受信メールを開き「返信の下書き」ボタンをクリック",
    "Step 3：Geminiが提案する返信案を確認・修正して送信",
    "Step 4：長文メールは「要約」機能で3行にまとめてもらう",
    "Step 5：定型返信はChatGPTでテンプレートを事前作成しておく",
    "",
    "時短効果：",
    "  ・メール1通あたりの処理時間：5分 → 2分",
    "  ・1日30通のメール処理：150分 → 60分（90分削減）",
    "  ・月間で約30時間の時間創出",
    "",
    "所要時間：10分（初期設定）  |  コスト：無料  |  難易度：Level 1",
])

doc.add_page_break()

# ================================================================
# 第7章：来月の注目トレンド予測
# ================================================================
add_h("第7章：来月の注目トレンド予測（2026年4月）", level=1)

add_data_table(
    ["No.", "注目トレンド", "予測される動き", "ビジネスへの影響"],
    [
        ("1", "NVIDIA Groq LPU\n市場導入本格化", "Vera Rubinプラットフォームの\n具体的な提供開始時期が明確化。\n推論特化チップの性能ベンチマーク公開", "AI推論コストの価格破壊が始まる\nAIエージェント導入コストが大幅低下\nクラウドサービス各社の対応が加速"),
        ("2", "EU AI Act\n高リスクAI対応事例", "高リスクAIシステムへの\n具体的な企業対応事例が続々登場。\nコンプライアンスツール市場が活性化", "対応先行企業の事例が参考になる\nAIガバナンスコンサルティング需要増\n対応遅れ企業への警鐘"),
        ("3", "Google I/O 2026\n（5月予定）事前発表", "Google I/O 2026の事前情報リーク。\nGemini 4系の噂。\nAndroid AI機能の大幅強化", "次世代AIモデルの方向性が見える\nGoogle製品のAI活用戦略が明確化\n開発者向けAPI新機能の予告"),
        ("4", "音楽生成AI\nライセンス整備", "WMGに続くメジャーレーベルの\nAI音楽ライセンス契約。\n商用利用ルールの標準化", "商用BGM制作のAI活用が加速\n著作権リスクの明確化\nクリエイターエコノミーへの影響"),
        ("5", "ChatGPT教育機能\nの拡充", "インタラクティブ学習機能の\n対象分野拡大。\n企業研修向け機能の提供開始", "社内研修のAI化が加速\n教育機関でのAI導入事例増加\n個人学習の効率が大幅向上"),
    ],
    hdr_color=C_PURPLE
)

add_info("4月のアクションプラン", [
    "・Google I/O 2026の事前情報に注目 → 次世代AI戦略のヒントを掴む",
    "・EU AI Act対応の先行企業事例を収集 → 自社の対応計画に反映",
    "・AI推論コスト低下のトレンドを活用 → 新規AIプロジェクトの検討開始",
])

doc.add_page_break()

# ================================================================
# まとめ：今月の重要ポイント10箇条
# ================================================================
add_h("まとめ：今月の重要ポイント10箇条", level=1)

points = [
    ("1", "NVIDIA Vera Rubinで推論コスト1/10の時代が到来", "AIの商用化コストが劇的に低減。エージェント型AIの普及が加速する"),
    ("2", "EU AI Actが本格施行 — 域外適用で日本企業も対象", "自社のAIシステムのリスク評価とコンプライアンス体制構築が急務"),
    ("3", "Claude Opus 4.6の100万トークンで長文分析の常識が変わった", "書籍数冊分を一度に分析。契約書比較・議事録分析が一瞬で可能"),
    ("4", "Gemini 3がGmail/Chromeに統合 — 日常業務が静かに革命", "設定するだけで使えるAI。メール処理・Webリサーチが自動化"),
    ("5", "ChatGPT無料版でもGPT-5.4 miniが使える時代に", "AIの民主化が進行。「無料でここまでできる」が驚異的なレベルに"),
    ("6", "動画生成AI三つ巴 — 個人でもプロ級映像が作れる", "Sora 2/Runway/Veo 3.1。SNS短尺動画の制作コストがほぼゼロに"),
    ("7", "音楽生成AIが商用ライセンス時代に突入", "WMGのSuno/Udioライセンス契約で「グレーゾーン」が解消へ"),
    ("8", "日本のAI推進法成立 — 1兆円規模の国産AI支援が始動", "規制よりも推進寄りの法律。国産AIエコシステムの充実に期待"),
    ("9", "AIエージェントが「ツール」から「チームメンバー」に変化", "目的を与えるだけでAIが自律実行。マルチエージェント時代の幕開け"),
    ("10", "3大AI（ChatGPT/Claude/Gemini）を無料で使い分けるのが正解", "1つに固定するのは損。用途別に使い分ける「3刀流」が最強戦略"),
]

add_data_table(
    ["No.", "ポイント", "解説"],
    points,
    hdr_color=C_NAVY
)

add_info("今週のアクション", [
    "Q1：今日試せるツールは何か？",
    "→ ChatGPT / Claude / Gemini の3つに無料登録し、同じ質問を投げて比較する",
    "",
    "Q2：あなたの仕事で最初に自動化できるのはどこか？",
    "→ メール返信・データ整理・レポート作成を棚卸しし、1つをAIに委任する",
    "",
    "Q3：来月までにLevel 2に上がるためにやること1つは？",
    "→ Gemini in Chromeを有効化し、毎日1回はAI要約を使う習慣をつける",
])

doc.add_paragraph()

ep = doc.add_paragraph()
ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(ep, "次回予告：第2回後半パート「主要AIツールの徹底比較と使い分け」", sz=Pt(12), color=RC_PURPLE, bold=True)

doc.add_paragraph()

ep2 = doc.add_paragraph()
ep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(ep2, "ご清聴ありがとうございました！", sz=Pt(14), color=RC_NAVY, bold=True)

# 保存
save_dir = os.path.expanduser("~/Desktop/AI経営本部/AI活用講義/講義資料/第02回")
os.makedirs(save_dir, exist_ok=True)
fp = os.path.join(save_dir, "第02回_①AI最新情報_2026年3月まとめ.docx")
doc.save(fp)
print(f"✅ 保存完了：{fp}")
