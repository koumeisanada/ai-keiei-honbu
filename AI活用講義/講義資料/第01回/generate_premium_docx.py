#!/usr/bin/env python3
"""第01回 講義資料 最高品質版 Word生成スクリプト
McKinsey / 大手コンサル提案書水準のデザイン"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os

# ============================================================
# カラーパレット
# ============================================================
C_NAVY = "1A1A2E"
C_PURPLE = "534AB7"
C_WHITE = "FFFFFF"
C_LIGHT_GRAY = "F5F5F5"
C_MID_GRAY = "E8E8E8"
C_DARK_GRAY = "4A4A4A"
C_TEXT = "2C2C2C"
C_RED = "C0392B"
C_ORANGE = "E67E22"
C_GREEN = "27AE60"
C_LIGHT_PURPLE = "EDE7F6"
C_LIGHT_NAVY = "E8EAF6"
C_LIGHT_GREEN = "E8F5E9"
C_LIGHT_RED = "FFEBEE"
C_LIGHT_ORANGE = "FFF3E0"

RC_NAVY = RGBColor(0x1A, 0x1A, 0x2E)
RC_PURPLE = RGBColor(0x53, 0x4A, 0xB7)
RC_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RC_TEXT = RGBColor(0x2C, 0x2C, 0x2C)
RC_RED = RGBColor(0xC0, 0x39, 0x2B)
RC_ORANGE = RGBColor(0xE6, 0x7E, 0x22)
RC_GREEN = RGBColor(0x27, 0xAE, 0x60)
RC_GRAY = RGBColor(0x4A, 0x4A, 0x4A)
RC_LIGHT_GRAY = RGBColor(0x95, 0x95, 0x95)

FONT_JP = "Yu Gothic"
FONT_EN = "Calibri"

doc = Document()

# ============================================================
# ページ設定
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================
# スタイル設定
# ============================================================
style_normal = doc.styles['Normal']
style_normal.font.name = FONT_EN
style_normal.font.size = Pt(10.5)
style_normal.font.color.rgb = RC_TEXT
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_JP)

for lvl in range(1, 4):
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = FONT_EN
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_JP)
    if lvl == 1:
        hs.font.size = Pt(22)
        hs.font.color.rgb = RC_NAVY
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif lvl == 2:
        hs.font.size = Pt(16)
        hs.font.color.rgb = RC_PURPLE
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.font.color.rgb = RC_NAVY
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)


# ============================================================
# ヘルパー関数
# ============================================================
def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=Pt(10.5), color=RC_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.font.name = FONT_EN
    r.font.size = size
    r.font.bold = bold
    r.font.color.rgb = color
    r.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_JP)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_run(paragraph, text, bold=False, size=Pt(10.5), color=RC_TEXT, italic=False):
    r = paragraph.add_run(text)
    r.font.name = FONT_EN
    r.font.size = size
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_JP)
    return r


def add_body(text, bold=False, size=Pt(10.5), color=RC_TEXT, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    add_run(p, text, bold=bold, size=size, color=color)
    return p


def add_heading_decorated(text, level=1):
    """色付き左ボーダー付きの見出し"""
    h = doc.add_heading(text, level=level)
    # 左ボーダーを追加
    pPr = h._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="8" w:color="{C_PURPLE}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    return h


def add_colored_box(title, lines, bg_color=C_LIGHT_PURPLE, title_bg=C_PURPLE, title_color=RC_WHITE, border_color=C_PURPLE):
    """プロフェッショナルな囲み枠（タイトルヘッダー+コンテンツ）"""
    t = doc.add_table(rows=2, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 枠線
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'<w:left w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders_xml)

    # タイトル行
    title_cell = t.cell(0, 0)
    set_cell_shading(title_cell, title_bg)
    set_cell_text(title_cell, f"  {title}", bold=True, size=Pt(12), color=title_color)

    # コンテンツ行
    content_cell = t.cell(1, 0)
    set_cell_shading(content_cell, bg_color)
    content_cell.paragraphs[0].text = ""
    for line in lines:
        p = content_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.4
        r = p.add_run(f"  {line}")
        r.font.name = FONT_EN
        r.font.size = Pt(10)
        r.font.color.rgb = RC_TEXT
        r.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_JP)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_warning_box(title, lines):
    """警告・注意ボックス（赤系）"""
    add_colored_box(title, lines, bg_color=C_LIGHT_RED, title_bg=C_RED, title_color=RC_WHITE, border_color=C_RED)


def add_success_box(title, lines):
    """成功・ベストプラクティスボックス（緑系）"""
    add_colored_box(title, lines, bg_color=C_LIGHT_GREEN, title_bg="1B7A3D", title_color=RC_WHITE, border_color=C_GREEN)


def add_info_box(title, lines):
    """情報・ポイントボックス（紫系）"""
    add_colored_box(title, lines, bg_color=C_LIGHT_PURPLE, title_bg=C_PURPLE, title_color=RC_WHITE, border_color=C_PURPLE)


def add_image_placeholder(description, width_cm=16):
    """画像挿入箇所プレースホルダー（オレンジ枠）"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_cell_shading(cell, C_LIGHT_ORANGE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[画像挿入箇所] {description}")
    r.font.name = FONT_EN
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RC_ORANGE
    r.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_JP)
    # オレンジ枠線
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="dashed" w:sz="12" w:space="0" w:color="{C_ORANGE}"/>'
        f'<w:left w:val="dashed" w:sz="12" w:space="0" w:color="{C_ORANGE}"/>'
        f'<w:bottom w:val="dashed" w:sz="12" w:space="0" w:color="{C_ORANGE}"/>'
        f'<w:right w:val="dashed" w:sz="12" w:space="0" w:color="{C_ORANGE}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders_xml)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_checklist(items):
    """チェックボックス付きリスト"""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_run(p, "[ ] ", bold=True, size=Pt(11), color=RC_PURPLE)
        add_run(p, item, size=Pt(10.5))


def add_data_table(headers, rows, header_color=C_NAVY):
    """プロフェッショナルなデータテーブル（ヘッダー色付き・交互行色分け）"""
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    # ヘッダー
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_shading(cell, header_color)
        set_cell_text(cell, h, bold=True, size=Pt(10), color=RC_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    # データ行（交互色分け）
    for r_idx, row_data in enumerate(rows):
        bg = C_LIGHT_GRAY if r_idx % 2 == 0 else C_WHITE
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx+1].cells[c_idx]
            set_cell_shading(cell, bg)
            set_cell_text(cell, val, size=Pt(9.5))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_separator():
    """薄い区切り線"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="{C_MID_GRAY}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    p.paragraph_format.space_after = Pt(8)


# ============================================================
# フッター（ページ番号）
# ============================================================
section = doc.sections[0]
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = fp.add_run("真田孔明のAI活用講座 第1回  |  ")
r1.font.size = Pt(8)
r1.font.color.rgb = RC_LIGHT_GRAY
r1.font.name = FONT_EN
r1.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_JP)
# ページ番号フィールド
fld_xml = (
    f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
    f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="{C_DARK_GRAY}"/></w:rPr><w:t>1</w:t></w:r>'
    f'</w:fldSimple>'
)
fp._p.append(parse_xml(fld_xml))
r2 = fp.add_run(" / ")
r2.font.size = Pt(8)
r2.font.color.rgb = RC_LIGHT_GRAY
fld_xml2 = (
    f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES \\* MERGEFORMAT ">'
    f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="{C_DARK_GRAY}"/></w:rPr><w:t>1</w:t></w:r>'
    f'</w:fldSimple>'
)
fp._p.append(parse_xml(fld_xml2))


# ================================================================
# ■ 表紙
# ================================================================
for _ in range(3):
    doc.add_paragraph()

# ネイビーバー
t_bar = doc.add_table(rows=1, cols=1)
t_bar.alignment = WD_TABLE_ALIGNMENT.CENTER
bar_cell = t_bar.cell(0, 0)
set_cell_shading(bar_cell, C_NAVY)
p_bar = bar_cell.paragraphs[0]
p_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_bar = p_bar.add_run("SANADA KOMEI AI ACADEMY")
r_bar.font.size = Pt(12)
r_bar.font.color.rgb = RC_WHITE
r_bar.font.bold = True
r_bar.font.name = FONT_EN
r_bar.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_JP)
r_bar2 = p_bar.add_run("\n真田孔明のAI活用講座")
r_bar2.font.size = Pt(10)
r_bar2.font.color.rgb = RGBColor(0xBB, 0xBB, 0xDD)
r_bar2.font.name = FONT_EN
r_bar2.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_JP)

doc.add_paragraph()
doc.add_paragraph()

# メインタイトル
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(title_p, "第1回講義資料", size=Pt(14), color=RC_PURPLE, bold=True)

doc.add_paragraph()

title_p2 = doc.add_paragraph()
title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(title_p2, "AIと向き合う前に\n知っておくべきこと", size=Pt(32), color=RC_NAVY, bold=True)

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(sub_p, "〜やってはいけないこと・注意すべきこと〜", size=Pt(14), color=RC_GRAY)

doc.add_paragraph()
doc.add_paragraph()

# パープルバー
t_bar2 = doc.add_table(rows=1, cols=1)
t_bar2.alignment = WD_TABLE_ALIGNMENT.CENTER
bar_cell2 = t_bar2.cell(0, 0)
set_cell_shading(bar_cell2, C_PURPLE)
p_bar2 = bar_cell2.paragraphs[0]
p_bar2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_bar2, f"調査日：2026年3月26日（更新版）  |  AI活用講義 月額制サブスクリプション", size=Pt(10), color=RC_WHITE)

doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(meta_p, "CONFIDENTIAL  |  For Members Only", size=Pt(9), color=RC_LIGHT_GRAY, italic=True)

doc.add_page_break()

# ================================================================
# ■ 目次
# ================================================================
add_heading_decorated("目次", level=1)
doc.add_paragraph()

toc_items = [
    ("パート1", "全体像 — この講座で1年間何を学ぶか", "3"),
    ("", "Level 1〜5の定義と現在地の確認", "3"),
    ("", "年間24回の学習ロードマップ", "4"),
    ("", "1年後に手に入るもの", "5"),
    ("パート2", "AIと向き合う前に知っておくべきこと", "6"),
    ("第1章", "AIリスクの全体像", "6"),
    ("第2章", "APIキーの管理と取り扱い", "8"),
    ("第3章", "ハルシネーションへの対処法", "11"),
    ("第4章", "著作権・知的財産の基本", "13"),
    ("第5章", "プライバシー・情報漏洩リスク", "16"),
    ("第6章", "トークンコストの管理", "19"),
    ("まとめ", "AIと安全に向き合うための10の鉄則", "21"),
    ("付録", "AI安全活用チェックリスト", "22"),
]

t_toc = doc.add_table(rows=len(toc_items), cols=3)
t_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (prefix, title, page) in enumerate(toc_items):
    bg = C_LIGHT_GRAY if i % 2 == 0 else C_WHITE
    c0 = t_toc.rows[i].cells[0]
    c1 = t_toc.rows[i].cells[1]
    c2 = t_toc.rows[i].cells[2]
    set_cell_shading(c0, bg)
    set_cell_shading(c1, bg)
    set_cell_shading(c2, bg)
    is_part = prefix in ("パート1", "パート2", "まとめ", "付録")
    set_cell_text(c0, prefix, bold=is_part, size=Pt(10), color=RC_PURPLE)
    set_cell_text(c1, title, bold=is_part, size=Pt(10), color=RC_NAVY if is_part else RC_TEXT)
    set_cell_text(c2, page, size=Pt(10), color=RC_GRAY, align=WD_ALIGN_PARAGRAPH.RIGHT)

doc.add_page_break()

# ================================================================
# ■ パート1：全体像
# ================================================================
add_heading_decorated("パート1：全体像（最重要）", level=1)

# ---- この講座で何を学ぶか ----
add_heading_decorated("この講座で1年間何を学ぶか", level=2)

add_info_box("講座の目標", [
    "AIの基礎知識からAIを活用したビジネス戦略の立案、AIツールの選定と導入、",
    "そして最も重要な「AIリスクマネジメント」までを1年間で習得します。",
    "",
    "AIを企業の競争力を高める戦略的資産として活用できる人材の育成を目指します。",
    "データに基づいた意思決定と効果的なリスク回避能力の強化を重視し、",
    "実務でのAIの導入を加速させます。",
])

# ---- Level定義 ----
add_heading_decorated("Level 1〜5の定義と現在地の確認", level=2)

add_body("AI活用におけるスキルレベルを5段階で定義します。まず自分の現在地を確認してください。")

add_data_table(
    ["Level", "定義", "到達に必要なステップ"],
    [
        ("Level 1\nAI初心者", "AIの基本概念を理解し、生成AIツールの存在を知っている", "AI基礎講座受講\n基本用語の理解"),
        ("Level 2\nAI基礎ユーザー", "生成AIツールの基本操作に慣れ、業務での試行錯誤を開始", "AIツールを用いた小規模プロジェクト実施\nリスクの基本知識習得"),
        ("Level 3\nAI応用ユーザー", "特定業務でAIを効率的に活用し、生産性向上を実感", "複数AIツールの比較と選定\nリスク対策の実施"),
        ("Level 4\nAIプロモーター", "社内でのAI活用を推進し、部署横断プロジェクトをリード", "プロジェクト管理スキル\n高度なプロンプトエンジニアリング"),
        ("Level 5\nAI経営戦略家", "AI技術を企業の成長戦略に統合し、全社的なAI活用を実現", "ガバナンス体制構築\n長期的成長戦略策定"),
    ],
    header_color=C_PURPLE
)

add_success_box("本日の目標", [
    "Level 1 → Level 2 へのステップアップを目指します",
    "AIリスクの基礎を徹底的に学び、安全なAI活用の土台を築きます",
])

# ---- ロードマップ ----
add_heading_decorated("年間24回の学習ロードマップ", level=2)

add_body("本講座は全24回（月2回）。AI活用スキルとリスクマネジメント能力を段階的に向上させます。")

add_data_table(
    ["クール", "回", "テーマ", "到達レベル"],
    [
        ("第1", "1〜6回", "AIリスク基礎とマインドセット\nAPIキー管理・ハルシネーション対策・情報漏洩事例・法的リスク", "Level 1→2"),
        ("第2", "7〜12回", "主要AIツールの理解と実践\nChatGPT/Claude/Gemini比較・プロンプトエンジニアリング基礎", "Level 2→3"),
        ("第3", "13〜18回", "高度なAI活用とセキュリティ\nAPI連携・社内データ連携・AIセキュリティ対策", "Level 3→4"),
        ("第4", "19〜24回", "AI経営戦略とガバナンス\nAIガバナンス体制構築・倫理ガイドライン・最新動向と未来予測", "Level 4→5"),
    ],
    header_color=C_NAVY
)

# ---- 1年後 ----
add_heading_decorated("1年後に手に入るもの", level=2)

add_info_box("講座修了時に得られる能力と資産", [
    "1. AIリスクを回避し、安全にAIを活用する能力",
    "2. AIツールを業務に適切に選定・活用するスキル",
    "3. 効果的なプロンプト設計技術",
    "4. AI活用による業務改善や新規事業の提案力",
    "5. AI経営本部の構造を理解し、自社に展開する力",
])

doc.add_page_break()

# ================================================================
# ■ パート2：AIと向き合う前に知っておくべきこと
# ================================================================
add_heading_decorated("パート2：AIと向き合う前に知っておくべきこと", level=1)

add_body("皆さん、こんにちは！真田孔明です。", bold=True, size=Pt(12), color=RC_NAVY)
add_body("AIは今、まさに私たちの仕事や生活を劇的に変えようとしています。これは「最強の刀」を手に入れたようなもの。正しく扱えば、これまで不可能だったことさえ可能にする力があります。しかし、その「刃」の扱いを誤れば、自らを、そして周囲をも傷つける危険性も秘めているのです。")
add_body("第1回となる今回は、AIを安全かつ効果的に使いこなすために、まず「AIと向き合う前に知っておくべきこと」を徹底的に学びましょう。Level 1の皆さんがLevel 5へとステップアップするための、最も重要な基礎となる部分です。")

add_separator()

# ================================================================
# 第1章：AIリスクの全体像
# ================================================================
add_heading_decorated("第1章：AIリスクの全体像", level=2)

add_image_placeholder("AIリスク4大カテゴリの図解（ナノバナナプロで生成）")

add_info_box("結論", [
    "AI活用には情報漏洩、誤情報拡散、セキュリティ侵害、法的・倫理的課題といったリスクが伴います。",
    "これらのリスクを包括的に理解し、対策を講じることが、安全で効果的なAI導入の前提条件です。",
])

doc.add_heading("根拠となる事実", level=3)

add_body("AIリスクは以下の4つのカテゴリに分類されます。", bold=True)

add_data_table(
    ["カテゴリ", "具体的リスク", "実際の事例"],
    [
        ("1. 情報セキュリティ\nリスク", "APIキー漏洩による不正アクセス\nAIサービスへの機密情報入力", "2023年6月：大手生成AIの認証情報\n26,000件以上がダークウェブで売買\n（INSIGHT HUB, 2025年12月）"),
        ("2. ハルシネーション\n（誤情報）リスク", "事実に基づかない情報の拡散\n誤情報に基づく意思決定", "2023年6月：米国弁護士が存在しない\n判例をChatGPTから引用し制裁\n（プロナビAI, 2025年6月）"),
        ("3. 法的・コンプラ\nイアンスリスク", "生成物における著作権侵害\n個人情報保護法違反", "2025年1月：文化庁ガイドライン改訂\n「本質的特徴の直接感得」で侵害認定\n（AIsmiley, 2026年2月）"),
        ("4. コストリスク", "トークン爆発による予算超過\nAPI不正利用による高額請求", "AI経営本部の実体験：\n全履歴方式でコスト30倍に膨張"),
    ],
    header_color=C_RED
)

add_warning_box("重要な警告", [
    "これらのリスクは「知らなかった」では済まされません。",
    "特にAPIキー漏洩と情報漏洩は、発生した瞬間から金銭的損害が発生します。",
    "本講義で学ぶ対策を、今日から必ず実践してください。",
])

doc.add_page_break()

# ================================================================
# 第2章：APIキーの管理と取り扱い
# ================================================================
add_heading_decorated("第2章：APIキーの管理と取り扱い", level=2)

add_image_placeholder("APIキーの正しい管理フロー図解（ナノバナナプロで生成）")

add_info_box("結論", [
    "APIキーは、AIサービスへのアクセスを許可する「重要な鍵」です。",
    "その漏洩は、会社の資産を守る上で最も避けなければならない重大なセキュリティインシデントに直結します。",
])

add_body("これはAI経営本部を構築した実体験に基づく最重要事項です。", bold=True, color=RC_RED, size=Pt(11))

doc.add_heading("APIキーとは何か", level=3)

add_body("APIキーとは、AIサービス（Claude・Gemini・ChatGPT）にプログラムからアクセスするための「鍵」です。")

add_warning_box("APIキー漏洩の現実", [
    "・APIキーが漏洩した場合、第三者があなたの費用でAIを無制限に使用できる",
    "・気づいた時には数十万円〜数百万円の請求が来る可能性がある",
    "・GitHubのコードに誤ってAPIキーを含めて公開 → 数時間で数百万円の請求事例が世界中で報告",
    "  （GitHub Security Advisory・各社公式報告より）",
    "・2023年6月：大手生成AIの認証情報が26,000件以上ダークウェブで売買（INSIGHT HUB, 2025年12月）",
])

doc.add_heading("絶対にやってはいけないこと", level=3)

add_data_table(
    ["No.", "禁止事項", "理由"],
    [
        ("1", "コードの中にAPIキーを直接書く\n例：api_key = \"sk-proj-XXXXX\"", "コードが公開された瞬間に漏洩"),
        ("2", "GitHubのリポジトリにAPIキーを\n含むファイルをアップロード", "全世界に公開される"),
        ("3", "メモ帳・Notion・Slackなどに\nそのままAPIキーを保存", "第三者のアクセスリスク"),
        ("4", "他人に見せる・送信する", "転送・スクショで拡散の危険"),
        ("5", "同じAPIキーを複数の\nプロジェクトで使い回す", "1箇所の漏洩で全滅する"),
    ],
    header_color=C_RED
)

doc.add_heading("正しいAPIキーの管理方法", level=3)

add_success_box("方法1：環境変数として保存する（推奨）", [
    "Macのターミナルで以下を実行：",
    "  echo 'export GEMINI_API_KEY=\"あなたのキー\"' >> ~/.zshrc",
    "  echo 'export OPENAI_API_KEY=\"あなたのキー\"' >> ~/.zshrc",
    "  source ~/.zshrc",
    "",
    "コードの中では以下のように呼び出す：",
    "  import os",
    "  api_key = os.environ.get(\"GEMINI_API_KEY\")",
    "",
    "これにより：",
    "  ・コードにキーが含まれない",
    "  ・GitHubにアップロードしても安全",
    "  ・~/.zshrcはGitHubに含まれない",
])

add_success_box("方法2：.gitignoreで保護する", [
    "AI経営本部フォルダの.gitignoreに以下が含まれていることを必ず確認：",
    "  .env",
    "  *_secret.txt",
    "  *_key.txt",
    "  setup_api.sh",
])

doc.add_heading("APIキーが漏洩した時の対処法", level=3)

add_warning_box("漏洩発覚時の緊急対応（1分以内に実行）", [
    "1. 即座にAPIキーを無効化する",
    "   ・Google AI Studio → APIキー管理 → 削除",
    "   ・OpenAI Platform → API keys → Revoke",
    "   ・Anthropic Console → API Keys → Delete",
    "2. 新しいAPIキーを発行する",
    "3. Usageページで使用量の異常がないか確認する",
    "4. Billingページで予想外の課金がないか確認する",
    "",
    "対処の速度が損害を最小化します。気づいたら1分以内に無効化することが重要です。",
])

doc.add_page_break()

# ================================================================
# 第3章：ハルシネーション
# ================================================================
add_heading_decorated("第3章：ハルシネーションへの対処法", level=2)

add_image_placeholder("AI別ハルシネーション発生率比較グラフ（ナノバナナプロで生成）")

add_info_box("結論", [
    "AIは「最もらしい嘘」をつく可能性があります。",
    "その出力は、常に人間の専門家による厳格な検証とファクトチェックを経てから、",
    "初めて利用すべき情報となります。",
])

doc.add_heading("なぜハルシネーションは起きるのか", level=3)

add_body("AI（特に生成AI）は、学習データに基づいて確率的に最もらしい回答を生成します。しかし、事実と異なる情報や「存在しない情報」を作り出す「ハルシネーション（幻覚）」と呼ばれる現象が発生します。")

add_warning_box("実際の被害事例", [
    "2023年6月、米国の弁護士がChatGPTの生成した「存在しない判例」を訴訟で引用",
    "→ 裁判所から制裁を受ける事態に発展",
    "→ 専門性の高い分野でのAI誤情報は法的責任に直結する",
    "  （プロナビAI, 2025年6月30日）",
])

doc.add_heading("ハルシネーション対策の4原則", level=3)

add_data_table(
    ["No.", "原則", "具体的アクション"],
    [
        ("1", "出典の確認", "AIが提示した情報に不明点がある場合、\n公的機関・学術論文・専門書で出典を確認する\nAI自身に「この情報の出典は？」と尋ねることも有効"),
        ("2", "複数ソースで検証", "特に重要な情報は、AI以外の複数の情報源を参照し\n事実関係を照合する\n最低でも3つ以上の異なる情報源で確認する"),
        ("3", "専門知識との照合", "自身の専門分野・経験に基づき、AIの出力が\n現実的か、論理的に破綻していないかを評価する\n「本当にこれで正しいのか？」という批判的視点を持つ"),
        ("4", "重要意思決定への\n直接利用を回避", "法的判断、医療診断、財務判断など\n影響の大きい意思決定に直接利用しない\n最終的な判断は人間の責任で行う"),
    ],
    header_color=C_PURPLE
)

doc.add_page_break()

# ================================================================
# 第4章：著作権
# ================================================================
add_heading_decorated("第4章：著作権・知的財産の基本", level=2)

add_image_placeholder("AI生成コンテンツの著作権判断フロー（ナノバナナプロで生成）")

add_info_box("結論", [
    "AIの利用においては、学習データの適法性、生成物の著作物性、既存著作物との類似性など、",
    "著作権・知的財産権に関する法的リスクを常に意識し、適切な対応が必須です。",
    "AIはまだ法整備が追いついていない領域であり、慎重な姿勢が求められます。",
])

doc.add_heading("日本の著作権法とAI（文化庁ガイドライン 2025年1月改訂）", level=3)

add_data_table(
    ["項目", "ルール", "備考"],
    [
        ("AI学習データの利用", "著作権法30条の4により\n原則として許容（非享受目的利用）", "AIによるデータ分析は\n「非享受目的利用」に該当"),
        ("著作権侵害となる\n例外", "作品鑑賞目的が併存する場合\n著作権者の利益を不当に損なう場合\nは違法と判断される可能性", "AIsmiley 2026年2月報道"),
        ("生成物の著作物性", "人間による「創作意図」と\n「創作的寄与」がある場合のみ\n著作権が認められる", "AIが自律的に生成した物は\n著作物に該当しない"),
        ("具体的表現の再現", "「生成物が既存著作物の本質的特徴を\n直接感得できる場合」は著作権侵害", "単なるスタイルの類似は\n侵害にならない"),
    ],
    header_color=C_NAVY
)

doc.add_heading("海外の重要判例", level=3)

add_data_table(
    ["国", "事件", "判断", "時期"],
    [
        ("中国", "ウルトラマン画像\n生成事件", "AI生成画像の著作物性と\n著作権侵害を認定（中国初）", "2024年8月"),
        ("米国", "Zarya of the Dawn\n事件", "AI生成画像の著作権は否定\nテキスト部分と人間の選択・\n配置・修正には著作権を認定", "2025年11月"),
    ],
    header_color=C_NAVY
)

add_success_box("著作権対策のベストプラクティス", [
    "1. 学習データの適法性確認：モデル選定時に利用規約・ライセンスを詳細確認",
    "2. 生成物の著作権侵害チェック：「具体的な表現の再現」に該当しないか確認",
    "3. 商用利用のライセンス確認：不明点はサービス提供元に問い合わせ",
    "4. 独自の創作性付与：AI生成物に人間の加筆・修正・編集を加えることを推奨",
])

doc.add_page_break()

# ================================================================
# 第5章：プライバシー
# ================================================================
add_heading_decorated("第5章：プライバシー・情報漏洩リスク", level=2)

add_image_placeholder("AIに入力してよい情報・いけない情報（ナノバナナプロで生成）")

add_info_box("結論", [
    "AIサービスへの機密情報や個人情報の入力は、情報漏洩のリスクを極めて高くします。",
    "会社の重要な情報を守るため、絶対に慎むべき行為です。",
])

doc.add_heading("実際の情報漏洩事例", level=3)

add_data_table(
    ["分類", "事例", "影響", "時期"],
    [
        ("企業事例①", "韓国大手電子機器メーカー従業員が\n半導体ソースコード・議事録を\nChatGPTに入力し情報流出", "同社はAIチャットボットの\n業務利用を全面禁止", "2023年3月\n(INSIGHT HUB)"),
        ("企業事例②", "大手ECサイト運営企業の内部データと\n一致する回答をChatGPTが生成", "社内での生成AI利用を制限\n機密情報が学習データに\n含まれていた可能性", "2022年12月\n(日本通信NW)"),
        ("プロンプト\n漏洩", "対話型生成AIサービスで\nユーザーの入力・登録情報が\n第三者に閲覧・編集可能に", "ニックネーム、メールアドレス\nLINE ID等が漏洩", "2024年3月\n(インターコム)"),
        ("個人事例", "ChatGPTのバグで\n他ユーザーのチャット履歴タイトルが\n閲覧可能に", "有料会員の氏名、メール\nカード番号の一部も漏洩", "2023年3月"),
    ],
    header_color=C_RED
)

add_warning_box("絶対に入力してはいけない情報", [
    "・顧客情報（氏名・住所・電話番号・メールアドレス）",
    "・人事情報（給与・評価・採用情報）",
    "・未発表の製品情報・戦略文書",
    "・ソースコード・APIキー・認証情報",
    "・クレジットカード番号・銀行口座情報",
    "・医療情報・法的書類の詳細",
])

add_success_box("情報漏洩対策の5原則", [
    "1. 機密情報・個人情報の入力禁止を社内規定として明文化し全従業員に周知徹底",
    "2. 情報入力時のフィルタリングツールを導入（特定キーワードの検知・ブロック）",
    "3. 「入力情報を学習データとして利用しない」設定（プライベートモード）を活用",
    "4. シャドーITを排除し、企業公認のセキュアなAI環境を社内で提供",
    "5. 定期的なセキュリティ教育で最新リスク情報と対策を全社共有",
])

doc.add_page_break()

# ================================================================
# 第6章：トークンコスト
# ================================================================
add_heading_decorated("第6章：トークンコストの管理", level=2)

add_image_placeholder("トークンコスト増加の仕組みとSKILLファイル方式（ナノバナナプロで生成）")

add_info_box("結論", [
    "AIのメモリ・記憶機能の設計を誤ると、毎月数十万円の請求が来る可能性があります。",
    "これはAI経営本部を構築した実体験に基づく最重要事項です。",
])

doc.add_heading("トークンとは何か", level=3)
add_body("トークン = AIが処理する文字・単語の単位（日本語約1文字 ≒ 1〜3トークン）")
add_body("課金の仕組み：入力トークン数 × 単価 ＋ 出力トークン数 × 単価 ＝ 1回の処理コスト", bold=True)

add_data_table(
    ["モデル", "入力単価（/百万トークン）", "出力単価（/百万トークン）", "備考"],
    [
        ("Claude Sonnet", "$3", "$15", "高性能・コスト中"),
        ("Gemini 2.5 Flash", "$0.075", "（タスクにより変動）", "低コスト・高速"),
        ("GPT-4o mini", "$0.15", "（タスクにより変動）", "低コスト・汎用"),
    ],
    header_color=C_PURPLE
)
add_body("※2026年3月時点の参考単価", size=Pt(9), color=RC_GRAY)

doc.add_heading("メモリ・記憶機能が危険な理由", level=3)

add_warning_box("全履歴方式のコスト爆発シミュレーション", [
    "AIに「会話の記憶を持たせたい」→ 全ての過去履歴をプロンプトに含める設計にすると：",
    "",
    "  1日目　：100トークンの履歴 × 10回 ＝    1,000トークン",
    "  1週間後：700トークンの履歴 × 10回 ＝    7,000トークン",
    "  1ヶ月後：3,000トークンの履歴 × 10回 ＝ 30,000トークン",
    "",
    "→ 月間コストが30倍に膨れ上がる",
    "",
    "AI経営本部の実体験：全会話履歴をコンテキストに含める設計にすると",
    "月間のAPI費用が想定の10〜30倍になることが確認されている",
])

doc.add_heading("正しいメモリ管理の設計", level=3)

add_success_box("方法1：SKILLファイル方式（推奨・AI経営本部で実践中）", [
    "AIの「記憶」はファイルとして保存する。毎回の会話には含めない。",
    "",
    "実装例：",
    "  ・SKILL_メルマガ自動生成.md（文体・ルール）",
    "  ・SKILL_真田孔明学習メモ.md（人物像・哲学）",
    "  ・SKILL_品質チェック.md（チェック基準）",
    "→ 必要な時だけ読み込む。毎回全て送信しない。",
])

add_success_box("方法2：直近N件のみを履歴として使用する", [
    "全履歴ではなく直近5件のみをコンテキストに含める設計にする。",
])

add_success_box("方法3：Claude Codeの活用（最重要）", [
    "Claude Codeを使う場合はCLAUDE.mdがAIの「記憶」として機能する。",
    "  → そのフォルダのコンテキストを定義",
    "  → 毎回の会話で自動的に参照される",
    "  → 会話履歴を蓄積しない",
    "  → コストを最小化できる",
    "",
    "これが真田孔明のAI経営本部がコストを抑えながら大量処理を実行できている理由です。",
])

doc.add_heading("コスト管理の実践", level=3)

add_data_table(
    ["サービス", "使用量確認URL", "設定項目"],
    [
        ("Gemini", "console.cloud.google.com/billing", "月額予算アラート設定"),
        ("OpenAI", "platform.openai.com/usage", "Spending limit設定"),
        ("Anthropic", "console.anthropic.com/usage", "月額上限設定"),
    ],
    header_color=C_NAVY
)

doc.add_page_break()

# ================================================================
# まとめ：10の鉄則
# ================================================================
add_heading_decorated("まとめ：AIと安全に向き合うための10の鉄則", level=1)

add_image_placeholder("10の鉄則インフォグラフィック（ナノバナナプロで生成）")

add_data_table(
    ["No.", "鉄則", "カテゴリ"],
    [
        ("1", "APIキーは環境変数（~/.zshrc）で管理する", "セキュリティ"),
        ("2", "APIキーをコードに直接書かない", "セキュリティ"),
        ("3", ".gitignoreでAPIキー関連ファイルを除外する", "セキュリティ"),
        ("4", "月額上限（Spending limit）を必ず設定する", "コスト管理"),
        ("5", "AIのメモリはSKILLファイルで管理する", "コスト管理"),
        ("6", "全履歴をプロンプトに含める設計は避ける", "コスト管理"),
        ("7", "直近N件のみを使う・または要約する", "コスト管理"),
        ("8", "毎月のAPI使用量を確認する習慣をつける", "コスト管理"),
        ("9", "異常なコスト増加を検知したら即座に調査する", "コスト管理"),
        ("10", "APIキー漏洩に気づいたら1分以内に無効化する", "セキュリティ"),
    ],
    header_color=C_NAVY
)

add_separator()

# ================================================================
# チェックリスト
# ================================================================
add_heading_decorated("付録：AI安全活用チェックリスト", level=1)

add_image_placeholder("AI安全活用チェックリスト図解（ナノバナナプロで生成）")

doc.add_heading("セキュリティ", level=3)
add_checklist([
    "APIキーを環境変数で管理している",
    "コードにAPIキーが含まれていないことを確認した",
    ".gitignoreにAPIキー関連ファイルを追加した",
    "多要素認証を有効化している",
    "APIキー漏洩時の対処手順を把握している",
])

doc.add_heading("コスト管理", level=3)
add_checklist([
    "各AIサービスに月額上限（Spending limit）を設定した",
    "AIのメモリ管理にSKILLファイル方式を採用している",
    "全履歴をプロンプトに含める設計になっていない",
    "毎月のAPI使用量を確認する習慣がある",
])

doc.add_heading("法的対応", level=3)
add_checklist([
    "AI生成コンテンツの著作権リスクを理解している",
    "商用利用時のライセンス条件を確認している",
    "生成物に独自の創作性を付与している",
    "個人情報保護法の基本を理解している",
])

doc.add_heading("品質管理", level=3)
add_checklist([
    "AIの出力をファクトチェックする習慣がある",
    "重要情報は3つ以上の情報源で検証している",
    "重要な意思決定にAI出力を直接利用しない",
    "機密情報・個人情報をAIに入力していない",
])

doc.add_page_break()

# ================================================================
# 結び
# ================================================================
add_heading_decorated("結び：AIを使いこなす「知恵」と「責任」", level=1)

add_body("")
add_body("皆さん、第1回講義、お疲れ様でした。", bold=True, size=Pt(12), color=RC_NAVY)
add_body("")
add_body("AIは私たちのビジネスに計り知れない可能性をもたらしますが、同時に「知っておくべきこと」「守るべきこと」も数多く存在します。今日の講義で学んだリスクと対策は、AIを賢く、そして責任を持って使いこなすための第一歩です。")
add_body("")
add_body("この「最強の刀」を、未来を切り開くための道具として最大限に活かすためには、私たち一人ひとりがその「刃」の扱い方を熟知し、細心の注意を払う必要があります。", bold=True)
add_body("")
add_body("AI経営本部では、皆さんのAI活用を全力でサポートしていきます。次回の講義では、さらに具体的なAI活用術と、リスクを回避しながら成果を最大化する「攻めのAI活用」について深掘りしていきましょう。")
add_body("")

add_info_box("次回予告", [
    "第2回：主要AIツール（ChatGPT / Claude / Gemini）の特徴と使い分け",
    "→ 3つのAIを比較し、業務に最適なツールを選定する方法を学びます",
])

add_body("")
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_end, "ご清聴ありがとうございました！", size=Pt(14), color=RC_NAVY, bold=True)

# ============================================================
# 保存
# ============================================================
save_dir = os.path.expanduser("~/Desktop/AI経営本部/AI活用講義/講義資料/第01回")
filepath = os.path.join(save_dir, "第01回_講義資料_最高品質版.docx")
doc.save(filepath)
print(f"✅ 保存完了：{filepath}")
