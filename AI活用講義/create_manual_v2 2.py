import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SAVE_DIR = os.path.expanduser("~/Desktop/AI経営本部/AI活用講義/講義資料/第01回")

def set_table_border(table, color='1A1A2E'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

def add_section_title(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(f'■ {text}')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '534AB7')
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        run = p.add_run(f'◆ {text}')
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p

def add_warning_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_border(table, 'CC0000')
    set_cell_bg(table.rows[0].cells[0], 'FFF0F0')
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run(f'⚠️ 注意　{text}')
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    doc.add_paragraph()

def add_info_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_border(table, '534AB7')
    set_cell_bg(table.rows[0].cells[0], 'F0F0FF')
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run(f'💡 ポイント　{text}')
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    doc.add_paragraph()

def add_prompt_box(doc, title, prompt_text):
    p_title = doc.add_paragraph()
    run_title = p_title.add_run(f'📋 {title}')
    run_title.font.size = Pt(11)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    table = doc.add_table(rows=1, cols=1)
    set_table_border(table, '1A1A2E')
    set_cell_bg(table.rows[0].cells[0], 'F5F5F5')
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].clear()
    for line in prompt_text.strip().split('\n'):
        p = cell.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()

def add_step_table(doc, steps):
    table = doc.add_table(rows=len(steps)+1, cols=3)
    table.style = 'Table Grid'
    set_table_border(table)
    headers = ['手順', '操作内容', '確認ポイント']
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        set_cell_bg(header_row.cells[i], '1A1A2E')
        for para in header_row.cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)
    for i, (step, content, check) in enumerate(steps):
        row = table.rows[i+1]
        row.cells[0].text = step
        row.cells[1].text = content
        row.cells[2].text = check
        if i % 2 == 0:
            for c in row.cells: set_cell_bg(c, 'F0F0FF')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    table.columns[0].width = Cm(2)
    table.columns[1].width = Cm(10)
    table.columns[2].width = Cm(4.5)
    doc.add_paragraph()

def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run(f'☐  {item}')
        run.font.size = Pt(10.5)

def create_manual():
    today = datetime.now().strftime('%Y年%m月%d日')
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.27); section.page_height = Inches(11.69)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    add_page_number(doc)

    # ===表紙===
    for _ in range(3): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('Level 5 AI経営本部'); r.font.size=Pt(32); r.font.bold=True; r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=p2.add_run('導入マニュアル 完全版'); r2.font.size=Pt(22); r2.font.bold=True; r2.font.color.rgb=RGBColor(0x53,0x4A,0xB7)
    doc.add_paragraph()
    p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r3=p3.add_run('AIを知らないゼロから始めて\nAI組織を経営できるようになるまでの完全手順書'); r3.font.size=Pt(13); r3.font.color.rgb=RGBColor(0x53,0x4A,0xB7)
    doc.add_paragraph()
    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(info_table)
    for i, (label, value) in enumerate([('発行日', today), ('対象', '真田孔明 AI活用講座 受講生')]):
        row = info_table.rows[i]
        row.cells[0].text = label; row.cells[1].text = value
        set_cell_bg(row.cells[0], 'E8E8F0')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size = Pt(10.5)
    doc.add_page_break()

    # ===はじめに===
    add_section_title(doc, 'はじめに')
    add_body(doc, 'このマニュアルは、AIを全く知らない方が、AI経営本部を構築できるようになるための完全手順書です。')
    doc.add_paragraph()
    add_info_box(doc, '所要時間の目安：全工程48時間以内。一つ一つ順番に進めれば必ず完成します。')
    doc.add_paragraph()
    add_section_title(doc, 'このマニュアルの進め方', 2)
    add_step_table(doc, [
        ('Step 1', '第1章：アカウントを作る（約2時間）', '4つのアカウントが作れたらOK'),
        ('Step 2', '第2章：プロンプトを入力する（約3時間）', '8つのプロンプトが完了したらOK'),
        ('Step 3', '第3章：AI経営本部を理解する（約1時間）', '組織構造が理解できたらOK'),
        ('Step 4', '第4〜5章：マネージャーとパイプライン（約4時間）', '自動化が動いたらOK'),
        ('Step 5', '第6〜7章：クラウド管理と完成確認（約2時間）', '全チェックリスト完了でOK'),
    ])
    doc.add_page_break()

    # ===第1章===
    add_section_title(doc, '第1章：準備（アカウント作成）')
    add_section_title(doc, '1-1. 必要なもの一覧', 2)
    t = doc.add_table(rows=6, cols=3); set_table_border(t)
    for i, h in enumerate(['必要なもの', '用途', '費用']):
        t.rows[0].cells[i].text = h; set_cell_bg(t.rows[0].cells[i], '1A1A2E')
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs: r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.size=Pt(10.5)
    for i, (a,b,c) in enumerate([
        ('パソコン','Mac推奨・Windows可','既存のものでOK'),
        ('Claudeアカウント','AI社長・メインAI','無料〜Max月額約15,000円'),
        ('Google AI Studio','Gemini API取得','無料から'),
        ('OpenAIアカウント','ChatGPT API取得','無料から'),
        ('GitHubアカウント','クラウド保存・管理','無料'),
    ]):
        row=t.rows[i+1]; row.cells[0].text=a; row.cells[1].text=b; row.cells[2].text=c
        if i%2==0:
            for c2 in row.cells: set_cell_bg(c2, 'F0F0FF')
        for c2 in row.cells:
            for p in c2.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()
    add_section_title(doc, '1-2. アカウント作成手順', 2)
    add_step_table(doc, [
        ('①','Claude：https://claude.ai →メールで登録→Maxプラン推奨','登録完了メールが届いたらOK'),
        ('②','Google AI Studio：https://aistudio.google.com →Googleログイン→「Get API key」→コピー','APIキーが表示されたらOK'),
        ('③','OpenAI：https://platform.openai.com →登録→「API keys」→新しいキー作成','APIキーが表示されたらOK'),
        ('④','GitHub：https://github.com →「Sign up」→無料プランでOK','プロフィール画面が表示されたらOK'),
    ])
    add_warning_box(doc, 'APIキーは絶対に他人に見せないでください。スクリーンショットもNG。取得したらすぐにメモ帳にコピーして保管。')
    doc.add_page_break()

    # ===第2章===
    add_section_title(doc, '第2章：最初にやること〜プロンプトを順番に入力するだけ〜')
    add_body(doc, 'アカウントが作れたら、以下のプロンプトを①〜⑧の順番にClaude Codeにコピペして送信するだけで、全ての設定が自動で完了します。')
    doc.add_paragraph()
    add_info_box(doc, '[ ]の部分は必ずあなたの情報に書き換えてください。それ以外はそのままコピペでOKです。')
    doc.add_paragraph()
    add_section_title(doc, '2-1. Claude Codeのインストール', 2)
    add_body(doc, 'ターミナル（Macの場合：Launchpadで「ターミナル」と検索して起動）を開いて入力：')
    add_prompt_box(doc, 'インストールコマンド', 'npm install -g @anthropic-ai/claude-code')
    add_body(doc, '起動：')
    add_prompt_box(doc, '起動コマンド', 'claude')
    add_info_box(doc, '「>」が表示されたら、以下のプロンプトを順番に貼り付けてください。')
    doc.add_paragraph()

    add_section_title(doc, '2-2. プロンプト①〜⑧', 2)
    prompts = [
        ('プロンプト①：自己紹介と環境確認',
         'あなたは私の専属AIアシスタントです。\nこれから一緒にAI経営本部を構築していきます。\n\n私の状況を教えます。\n・名前：[あなたの名前]\n・職業：[あなたの職業]\n・パソコン：[MacまたはWindows]\n・AIの経験：ほぼゼロ\n\n今日からAI経営本部の構築を始めます。\n最初に何を準備すればいいか、\nステップ形式で教えてください。'),
        ('プロンプト②：AI経営本部フォルダの自動構築',
         'デスクトップに「AI経営本部」というフォルダを作成して、\n以下の構造を全て自動で作成してください。\n\nAI経営本部/\n├── CLAUDE.md\n├── 人財育成/\n│   ├── CLAUDE.md\n│   └── 資料/\n├── 物販ビジネス/\n│   ├── CLAUDE.md\n│   └── 資料/\n├── 銀行融資/\n│   ├── CLAUDE.md\n│   └── 資料/\n├── 会社法人/\n│   ├── CLAUDE.md\n│   └── 資料/\n├── 米国成長株/\n│   ├── CLAUDE.md\n│   └── 資料/\n└── API連携/\n    └── Gemini/'),
        ('プロンプト③：AI社長の設定',
         'Desktop/AI経営本部/CLAUDE.mdに以下を書き込んでください。\n\nあなたは[あなたの名前]のAI経営本部を統括するAI社長です。\n\n## 担当ビジネス\n・人財育成\n・物販ビジネス\n・銀行融資\n・会社法人\n・米国成長株\n\n## 行動原則\n・全て日本語で作成する\n・架空の事例は使わない\n・必ず根拠・データを付ける'),
        ('プロンプト④：APIキーの設定',
         '以下のAPIキーを環境変数に安全に設定してください。\n\nGEMINI_API_KEY=[取得したGeminiのAPIキーを貼り付け]\nOPENAI_API_KEY=[取得したOpenAIのAPIキーを貼り付け]\n\n~/.zshrcに追記してsource ~/.zshrcで反映。\n設定完了後に動作確認もお願いします。'),
        ('プロンプト⑤：GitHubとの連携',
         'GitHubと連携してAI経営本部をクラウドで管理できるようにしてください。\n\nGitHubのユーザー名：[あなたのGitHubユーザー名]\nリポジトリ名：ai-keiei-honbu\n\nGitの初期設定→リポジトリ作成→接続→最初のコミット＆プッシュ→.gitignore設定\n全て完了したらGitHubのURLを教えてください。'),
        ('プロンプト⑥：5つのAIマネージャーの設定',
         '以下の5つのAIマネージャーを設定してください。\n各フォルダのCLAUDE.mdに役割を書き込んでください。\n\n■人財育成マネージャー（人財育成/CLAUDE.md）\n役割：チームマネジメント・人材育成の資料作成\n\n■物販ビジネスマネージャー（物販ビジネス/CLAUDE.md）\n役割：商品リサーチ・出品文章作成・顧客対応\n\n■銀行融資マネージャー（銀行融資/CLAUDE.md）\n役割：融資申請書・事業計画書の作成サポート\n\n■会社法人マネージャー（会社法人/CLAUDE.md）\n役割：議事録・契約書ドラフト・法人経営資料\n\n■米国成長株マネージャー（米国成長株/CLAUDE.md）\n役割：米国成長株のデイリーリサーチと分析'),
        ('プロンプト⑦：SKILLファイルの作成',
         'Desktop/AI経営本部/に以下のSKILLファイルを作成してください。\n\nSKILL_自己紹介.md\n# 私のプロフィール\n・名前：[あなたの名前]\n・職業：[あなたの職業]\n・目標：[あなたの目標]\n\nSKILL_品質チェック.md\n# 品質チェック基準\n・事実と異なる内容は含まれていないか\n・根拠・データは明記されているか\n・100点満点で採点して報告すること'),
        ('プロンプト⑧：動作確認',
         'AI経営本部が正しく構築されているか全て確認してください。\n\n確認項目：\n・AI経営本部フォルダの構造\n・各CLAUDE.mdの内容\n・APIキーの設定\n・GitHubとの接続状態\n・SKILLファイルの設置\n\n全ての確認結果を表にしてまとめてください。\n問題があれば自動で修正してください。'),
    ]
    for title, prompt_text in prompts:
        add_prompt_box(doc, title, prompt_text)

    doc.add_paragraph()
    add_section_title(doc, '2-3. プロンプト入力のルール', 2)
    t2 = doc.add_table(rows=6, cols=2); set_table_border(t2)
    for i, h in enumerate(['ルール', '内容']):
        t2.rows[0].cells[i].text = h; set_cell_bg(t2.rows[0].cells[i], '1A1A2E')
        for p in t2.rows[0].cells[i].paragraphs:
            for r in p.runs: r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.size=Pt(10.5)
    for i, (a,b) in enumerate([
        ('[ ]は必ず書き換える','[あなたの名前]→自分の名前に変更'),
        ('①から順番に実行','飛ばさない・一つ完了してから次へ'),
        ('エラーが出ても慌てない','「エラーが出ました。解決してください」と入力'),
        ('わからなければ聞く','「もっと詳しく教えてください」でOK'),
        ('途中で止まったら','「続きをお願いします」と入力'),
    ]):
        row=t2.rows[i+1]; row.cells[0].text=a; row.cells[1].text=b
        if i%2==0:
            for c in row.cells: set_cell_bg(c, 'F0F0FF')
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_page_break()

    # ===第3章===
    add_section_title(doc, '第3章：AI経営本部の全体設計')
    add_section_title(doc, '3-1. 組織構造', 2)
    t3 = doc.add_table(rows=7, cols=3); set_table_border(t3)
    for i, h in enumerate(['役職', '担当', '使用AI']):
        t3.rows[0].cells[i].text = h; set_cell_bg(t3.rows[0].cells[i], '1A1A2E')
        for p in t3.rows[0].cells[i].paragraphs:
            for r in p.runs: r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.size=Pt(10.5)
    for i, (a,b,c) in enumerate([
        ('AI社長','全プロジェクトの統括','Claude'),
        ('人財育成マネージャー','チーム・人材育成','Claude＋Gemini'),
        ('物販ビジネスマネージャー','商品リサーチ・出品','Gemini＋Claude'),
        ('銀行融資マネージャー','融資資料・事業計画書','Claude＋ChatGPT'),
        ('会社法人マネージャー','議事録・契約書','Claude＋ChatGPT'),
        ('米国成長株マネージャー','株式リサーチ・決算分析','Gemini（検索）'),
    ]):
        row=t3.rows[i+1]; row.cells[0].text=a; row.cells[1].text=b; row.cells[2].text=c
        if i==0:
            for c2 in row.cells: set_cell_bg(c2, '534AB7')
            for c2 in row.cells:
                for p in c2.paragraphs:
                    for r in p.runs: r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.bold=True; r.font.size=Pt(10.5)
        else:
            if i%2==0:
                for c2 in row.cells: set_cell_bg(c2, 'F0F0FF')
            for c2 in row.cells:
                for p in c2.paragraphs:
                    for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()
    doc.add_page_break()

    # ===第4章===
    add_section_title(doc, '第4章：SKILLファイルの設定')
    add_body(doc, 'SKILLファイルはAIへの取扱説明書です。毎回同じ品質の成果物が得られます。')
    add_info_box(doc, 'SKILLファイルはAIの「記憶」の代わり。会話履歴に頼らないため、コストを大幅に削減できます。')
    doc.add_page_break()

    # ===第5章===
    add_section_title(doc, '第5章：デイリーパイプラインの構築')
    add_body(doc, 'パイプラインは工場の自動生産ラインです。一度設定すれば毎朝AIが自動で作業を完了させます。')
    add_prompt_box(doc, 'パイプライン構築プロンプト',
        '私のビジネスに合わせたデイリーパイプラインを作成してください。\n\n毎日自動で実行したい作業：\n・[作業1]\n・[作業2]\n・[作業3]\n\n保存先：Desktop/AI経営本部/API連携/pipeline_daily.sh')
    add_prompt_box(doc, '自動実行設定プロンプト',
        'pipeline_daily.shを毎朝6時に自動実行する設定をしてください。\nMac：LaunchAgent / Windows：タスクスケジューラー')
    doc.add_page_break()

    # ===第6章===
    add_section_title(doc, '第6章：GitHubでクラウド管理＋iPhoneからアクセス')
    add_step_table(doc, [
        ('①','iPhoneでSafariを開く',''),
        ('②','https://claude.ai/code にアクセスする','Claude Code Webが開く'),
        ('③','GitHubアカウントでログインする','リポジトリが表示される'),
        ('④','ai-keiei-honbuリポジトリを選択する','ファイル一覧が表示される'),
        ('⑤','Claude Codeに指示を入力する','外出先からでもAIが動く'),
    ])
    doc.add_page_break()

    # ===第7章===
    add_section_title(doc, '第7章：完成確認チェックリスト')
    add_checklist(doc, [
        'Claudeアカウントが作成できた',
        'Google AI Studio（Gemini）のAPIキーが取得できた',
        'OpenAIのAPIキーが取得できた',
        'GitHubアカウントが作成できた',
        'Claude Codeがインストールできた',
        'AI経営本部フォルダ構造が作成できた',
        'AI社長CLAUDE.mdが設定できた',
        'APIキーが環境変数に設定できた',
        'GitHubと連携できた',
        '5つのAIマネージャーが設定できた',
        'SKILLファイルが作成できた',
        'デイリーパイプラインが動作した',
        '毎朝6時の自動実行が設定できた',
        'iPhoneから操作できた',
        '品質チェックが自動で動いた',
    ])
    doc.add_paragraph()
    add_section_title(doc, 'よくあるトラブルと解決方法', 2)
    tt = doc.add_table(rows=6, cols=3); set_table_border(tt)
    for i, h in enumerate(['症状','原因','解決方法']):
        tt.rows[0].cells[i].text=h; set_cell_bg(tt.rows[0].cells[i], '1A1A2E')
        for p in tt.rows[0].cells[i].paragraphs:
            for r in p.runs: r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.size=Pt(10.5)
    for i, (a,b,c) in enumerate([
        ('claudeが起動しない','Node.js未インストール','nodejs.orgからインストールして再試行'),
        ('APIキーが認識されない','環境変数の設定ミス','「APIキーの設定を確認して」とClaudeに入力'),
        ('GitHubにプッシュできない','認証エラー','「GitHubの認証エラーを解決して」と入力'),
        ('パイプラインが動かない','スクリプトエラー','「pipeline_daily.shのエラーを修正して」と入力'),
        ('自動実行が動かない','LaunchAgent設定ミス','「自動実行の設定を確認して」と入力'),
    ]):
        row=tt.rows[i+1]; row.cells[0].text=a; row.cells[1].text=b; row.cells[2].text=c
        if i%2==0:
            for c2 in row.cells: set_cell_bg(c2, 'F0F0FF')
        for c2 in row.cells:
            for p in c2.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_page_break()

    # ===付録===
    add_section_title(doc, '付録：CLAUDE.mdテンプレート集')
    for title, template in [
        ('AI社長テンプレート', 'あなたは[名前]のAI経営本部を統括するAI社長です。\n\n## 担当ビジネス\n・人財育成\n・物販ビジネス\n・銀行融資\n・会社法人\n・米国成長株\n\n## 行動原則\n・全て日本語で作成する\n・架空の事例は使わない\n・必ず根拠・データを付ける'),
        ('マネージャーテンプレート（汎用）', 'あなたは[プロジェクト名]を担当するAIマネージャーです。\n\n## 役割\n[役割を記載]\n\n## 参照フォルダ\n./資料/\n\n## 出力形式\n[成果物の種類を記載]'),
    ]:
        add_prompt_box(doc, title, template)

    os.makedirs(SAVE_DIR, exist_ok=True)
    filepath = os.path.join(SAVE_DIR, '【受講生用】Level5_AI経営本部_導入マニュアル_v2.docx')
    doc.save(filepath)
    print(f'✅ マニュアルv2完成：{filepath}')
    return filepath

def main():
    print('=' * 60)
    print('Level5 AI経営本部 導入マニュアル v2 生成開始')
    print('=' * 60)
    filepath = create_manual()
    print(f'\n✅ 完成ファイル：{filepath}')

if __name__ == '__main__':
    main()
