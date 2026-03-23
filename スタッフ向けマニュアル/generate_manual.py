import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    return p

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0C, 0x44, 0x7C)
    run.font.name = 'Courier New'
    return p

def add_green(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x6E, 0x56)
    return p

def add_hr(doc):
    p = doc.add_paragraph('─' * 50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

# 表紙
add_title(doc, 'AI経営本部 構築マニュアル')
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('〜 エンジニア経験ゼロから4日間で構築する\nAI社長×3大AI連携×全自動化の器（フレームワーク） 〜')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x88, 0x87, 0x80)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_str = datetime.now().strftime("%Y年%m月%d日")
run = p.add_run(f'作成日：{date_str}')
run.font.size = Pt(11)
doc.add_page_break()

# はじめに
add_h1(doc, 'はじめに')
add_hr(doc)
add_body(doc, 'このマニュアルは「AI経営本部」という器（フレームワーク）を構築するためのものです。')
doc.add_paragraph()
add_body(doc, '大切なのは以下の2点です。')
doc.add_paragraph()
add_bullet(doc, '①このマニュアルで構築するのは「器（フレームワーク）」だけです')
add_bullet(doc, '②器の中に入れるスキル・業務内容は、各自のビジネスに合わせて自由にカスタマイズします')
doc.add_paragraph()
add_body(doc, 'たとえば以下のようなビジネスに活用できます。')
doc.add_paragraph()
add_bullet(doc, 'サラリーマンの業務効率化・出世昇進への活用')
add_bullet(doc, '物販ビジネスの自動化・効率化')
add_bullet(doc, '銀行融資のための資料作成自動化')
add_bullet(doc, '米国株式投資のリサーチ自動化')
add_bullet(doc, '物販会社法人の経営効率化')
add_bullet(doc, 'その他、独自のビジネス全般')
doc.add_paragraph()
add_body(doc, 'まず器を作る。その後に自分のビジネスのスキルを追加していく。この順番で進めてください。')
doc.add_page_break()

# 全体像
add_h1(doc, '全体像：構築する5本柱')
add_hr(doc)
add_bullet(doc, '①3大AI連携システム：Claude・Gemini・ChatGPTをそれぞれの強みで使い分ける')
add_bullet(doc, '②AI社長＋AIマネージャー組織：人間の組織と同じ構造をAIで再現する')
add_bullet(doc, '③基本パイプライン：よく使う作業を自動化する基本的な仕組み')
add_bullet(doc, '④GitHubクラウド管理：全データをクラウドで管理・バックアップする')
add_bullet(doc, '⑤iPhoneからの遠隔操作：スマホ1台でどこからでも指示が出せる')
doc.add_paragraph()
add_h2(doc, '所要時間の目安')
add_bullet(doc, 'Day1（2〜3時間）：環境構築＋AI組織の器を作る')
add_bullet(doc, 'Day2（1〜2時間）：3大AI連携を設定する')
add_bullet(doc, 'Day3（1〜2時間）：基本パイプラインを動かす')
add_bullet(doc, 'Day4（1時間）：GitHub連携＋iPhone設定')
add_bullet(doc, '合計：4日間・約6〜8時間で完成')
doc.add_paragraph()
add_h2(doc, '月額コストの目安')
add_bullet(doc, 'Claude Maxプラン：約15,000円/月')
add_bullet(doc, 'Gemini API：約1,000〜3,000円/月')
add_bullet(doc, 'OpenAI API：約1,000〜2,000円/月')
add_bullet(doc, '合計：約17,000〜20,000円/月')
doc.add_page_break()

# 事前準備
add_h1(doc, '事前準備：アカウントを作成する')
add_hr(doc)
add_body(doc, 'Day1を始める前に以下のアカウントを全て作成してください。')
doc.add_paragraph()
add_h2(doc, '①Claudeアカウント（必須）')
add_bullet(doc, 'URL：https://claude.ai')
add_bullet(doc, 'プラン：Maxプラン推奨（月額約15,000円）')
add_bullet(doc, '理由：大量ファイル処理・長時間作業に対応できる')
doc.add_paragraph()
add_h2(doc, '②Google AI Studioアカウント（必須）')
add_bullet(doc, 'URL：https://aistudio.google.com')
add_bullet(doc, 'プラン：無料枠でOK')
add_bullet(doc, '用途：Gemini APIキーの取得')
doc.add_paragraph()
add_h2(doc, '③OpenAIアカウント（必須）')
add_bullet(doc, 'URL：https://platform.openai.com')
add_bullet(doc, 'プラン：クレジット購入（最低10ドル）')
add_bullet(doc, '用途：ChatGPT APIキーの取得')
doc.add_paragraph()
add_h2(doc, '④GitHubアカウント（必須）')
add_bullet(doc, 'URL：https://github.com')
add_bullet(doc, 'プラン：無料プランでOK')
add_bullet(doc, '用途：AI経営本部のクラウド管理')
doc.add_page_break()

# Day1
add_h1(doc, 'Day1：環境構築＋AI組織の器を作る')
add_hr(doc)
add_h2(doc, 'Step1：Node.jsのインストール')
add_body(doc, 'Macのターミナルを開いてください。')
add_body(doc, '（Command+スペース→「terminal」と入力→Enter）')
doc.add_paragraph()
add_body(doc, '以下を実行してください。')
add_code(doc, 'node --version')
add_body(doc, 'バージョンが表示されない場合は以下を実行してください。')
add_code(doc, 'curl -o ~/Downloads/node.pkg https://nodejs.org/dist/v20.11.0/node-v20.11.0.pkg && open ~/Downloads/node.pkg')
add_body(doc, 'インストーラーが開くので「続ける」→「インストール」と進んでください。')
doc.add_paragraph()
add_h2(doc, 'Step2：Claude Codeのインストール')
add_code(doc, 'sudo npm install -g @anthropic-ai/claude-code')
add_body(doc, 'Macのパスワードを入力してください（画面には表示されませんが入力できています）。')
doc.add_paragraph()
add_h2(doc, 'Step3：Claude Codeの起動')
add_code(doc, 'claude')
add_body(doc, 'ブラウザが開くのでClaudeアカウントでログインしてください。')
add_body(doc, '「Welcome back！」と表示されれば起動成功です。')
doc.add_paragraph()
add_h2(doc, 'Step4：AI経営本部フォルダの作成')
add_body(doc, 'Claude Codeの入力欄に以下をそのまま入力してください。')
add_body(doc, '【プロジェクト名】の部分を自分のビジネスに合わせて書き換えてください。')
doc.add_paragraph()
add_green(doc, '''▼ Claude Codeへの入力（コピペしてください）

デスクトップにAI経営本部フォルダを作成して
以下のサブフォルダを全て作ってください。

【プロジェクト1名】/資料
【プロジェクト2名】/資料
【プロジェクト3名】/資料
集客販売/原稿
集客販売/日次成果物
API連携/Gemini
API連携/設定ファイル
スキル追加用
品質チェック''')
doc.add_paragraph()
add_h2(doc, 'Step5：AI社長の設定')
add_body(doc, 'Claude Codeに以下を入力してください。【】の部分を書き換えてください。')
doc.add_paragraph()
add_green(doc, '''▼ Claude Codeへの入力（コピペしてください）

Desktop/AI経営本部/CLAUDE.mdを作成してください。

内容：
あなたは【自分の名前】のビジネスを統括するAI社長です。

【管轄プロジェクト】
1.【プロジェクト1の名前と概要】
2.【プロジェクト2の名前と概要】
3.【プロジェクト3の名前と概要】

【AI社長の役割】
・全プロジェクトの進捗把握と優先順位の判断
・各AIマネージャーへの指示出し
・オーナーへの報告

【行動原則】
・常にオーナーのビジネス目標から逆算して動く
・事実と根拠に基づいて判断する
・新しいスキルが追加されたら即座に組織に反映する''')
doc.add_paragraph()
add_h2(doc, 'Step6：AIマネージャーの設定')
add_body(doc, '各プロジェクトフォルダにCLAUDE.mdを作成します。')
add_body(doc, 'Claude Codeに以下を入力してください。')
doc.add_paragraph()
add_green(doc, '''▼ Claude Codeへの入力（コピペしてください）

以下のCLAUDE.mdを各フォルダに作成してください。

Desktop/AI経営本部/【プロジェクト1】/CLAUDE.md
内容：
あなたは【プロジェクト1】の担当AIマネージャーです。
役割：【このプロジェクトで何をするか】
参照フォルダ：./資料/
出力形式：【どんな成果物を作るか】

Desktop/AI経営本部/【プロジェクト2】/CLAUDE.md
内容：
あなたは【プロジェクト2】の担当AIマネージャーです。
役割：【このプロジェクトで何をするか】
参照フォルダ：./資料/
出力形式：【どんな成果物を作るか】''')
doc.add_paragraph()
add_body(doc, '※AIマネージャーはプロジェクトの数だけ作成してください。')
add_body(doc, '※役割・出力形式は後からいつでも修正・追加できます。')
doc.add_page_break()

# Day2
add_h1(doc, 'Day2：3大AI連携を設定する')
add_hr(doc)
add_h2(doc, 'Step1：APIキーの取得')
add_body(doc, '■ Gemini APIキーの取得')
add_bullet(doc, 'https://aistudio.google.com を開く')
add_bullet(doc, '左メニュー「Get API key」→「APIキーを作成」')
add_bullet(doc, '表示されたキーをメモ帳に保存')
doc.add_paragraph()
add_body(doc, '■ OpenAI APIキーの取得')
add_bullet(doc, 'https://platform.openai.com を開く')
add_bullet(doc, '「API keys」→「Create new secret key」')
add_bullet(doc, '表示されたキーをメモ帳に保存（一度しか表示されない）')
doc.add_paragraph()
add_h2(doc, 'Step2：APIキーの設定')
add_body(doc, 'Claude Codeに以下を入力してください。【】の部分を書き換えてください。')
doc.add_paragraph()
add_green(doc, '''▼ Claude Codeへの入力（コピペしてください）

以下のAPIキーを~/.zshrcに保存して
すぐに使えるようにしてください。

Gemini APIキー：【取得したGemini APIキー】
OpenAI APIキー：【取得したOpenAI APIキー】

設定後に以下の3つのAIが全て接続できているか確認してください。
・Claude（Anthropic）
・Gemini（Google）
・ChatGPT（OpenAI）''')
doc.add_paragraph()
add_h2(doc, 'Step3：3大AIの役割分担を設定する')
add_body(doc, 'Claude Codeに以下を入力してください。')
doc.add_paragraph()
add_green(doc, '''▼ Claude Codeへの入力（コピペしてください）

Desktop/AI経営本部/API連携/SKILL_3大AI役割分担.mdを作成してください。

内容：
# 3大AI役割分担

## Claude（指揮役・司令塔）
・戦略立案・文章作成・資料構成・分析
・全AIへの指示書作成
・成果物の品質チェック

## Gemini（リサーチ・文章生成）
・ウェブ検索を活用したリサーチ
・文章の大量生成
・画像生成プロンプトの作成

## ChatGPT（補助・多様性）
・別視点からの文章チェック
・Claudeの出力に対する反論・改善提案
・英語コンテンツの作成

## 使い分けのルール
・まずClaudeに指示を出す
・リサーチが必要な場合はGeminiを活用
・品質チェックにChatGPTを活用''')
doc.add_page_break()

# Day3
add_h1(doc, 'Day3：基本パイプラインを動かす')
add_hr(doc)
add_body(doc, 'パイプラインとは「一連の作業を自動実行する仕組み」のことです。')
add_body(doc, 'まず最もシンプルな基本パイプラインを作って動かします。')
doc.add_paragraph()
add_h2(doc, 'Step1：必要なライブラリのインストール')
add_body(doc, 'Claude Codeに以下を入力してください。')
doc.add_paragraph()
add_green(doc, '''▼ Claude Codeへの入力（コピペしてください）

以下をインストールしてください。

pip3 install google-genai openai python-docx requests --break-system-packages''')
doc.add_paragraph()
add_h2(doc, 'Step2：自分のビジネス用パイプラインの作成')
add_body(doc, 'Claude Codeに以下を入力してください。【】の部分を書き換えてください。')
doc.add_paragraph()
add_green(doc, '''▼ Claude Codeへの入力（コピペしてください）

私のビジネス用の基本パイプラインを作成してください。

【私のビジネスの概要】
・ビジネスの種類：【例：物販・サラリーマン副業・投資など】
・主な作業内容：【例：商品リサーチ・資料作成・情報収集など】
・毎日やりたい自動化：【例：競合調査・ニュース収集・文章作成など】

以下のファイルを作成してください。

1. Desktop/AI経営本部/API連携/pipeline_basic.sh
   毎日の基本作業を自動実行するシェルスクリプト

2. Desktop/AI経営本部/API連携/Gemini/research_daily.py
   Geminiを使った毎日のリサーチスクリプト

テーマを入力するだけで以下を実行してください。
・指定テーマのリサーチ（Gemini使用）
・結果を日付フォルダに自動保存
・品質チェック（Claude使用）''')
doc.add_paragraph()
add_h2(doc, 'Step3：動作テスト')
add_body(doc, 'Claude Codeに以下を入力してください。')
doc.add_paragraph()
add_green(doc, '''▼ Claude Codeへの入力（コピペしてください）

基本パイプラインのテストを実行してください。
テーマ：「【自分のビジネスに関連するテーマ】」

成果物が日次成果物フォルダに保存されれば成功です。''')
doc.add_paragraph()
add_h2(doc, 'Step4：スキル追加の方法')
add_body(doc, 'パイプラインが動いたら、自分のビジネスに合わせてスキルを追加していきます。')
doc.add_paragraph()
add_body(doc, '【5ポケッツ戦略術向けスキル追加の例】')
doc.add_paragraph()
add_bullet(doc, 'サラリーマン業務効率化スキル')
add_code(doc, '「社内資料・報告書・企画書を自動生成するスキルをDesktop/AI経営本部/スキル追加用/SKILL_サラリーマン業務.mdとして作成してください」')
doc.add_paragraph()
add_bullet(doc, '物販ビジネス自動化スキル')
add_code(doc, '「商品リサーチ・出品文章・競合分析を自動化するスキルをDesktop/AI経営本部/スキル追加用/SKILL_物販ビジネス.mdとして作成してください」')
doc.add_paragraph()
add_bullet(doc, '銀行融資資料作成スキル')
add_code(doc, '「創業計画書・月別利益計画書・5カ年計画書を自動生成するスキルをDesktop/AI経営本部/スキル追加用/SKILL_銀行融資.mdとして作成してください」')
doc.add_paragraph()
add_bullet(doc, '米国株リサーチスキル')
add_code(doc, '「NVIDIA・Apple・Microsoft等の最新情報を毎日自動収集するスキルをDesktop/AI経営本部/スキル追加用/SKILL_米国株リサーチ.mdとして作成してください」')
doc.add_paragraph()
add_bullet(doc, '法人経営効率化スキル')
add_code(doc, '「経費管理・決算書分析・節税調査を自動化するスキルをDesktop/AI経営本部/スキル追加用/SKILL_法人経営.mdとして作成してください」')
doc.add_page_break()

# Day4
add_h1(doc, 'Day4：GitHub連携＋iPhone設定')
add_hr(doc)
add_h2(doc, 'Step1：GitHubリポジトリの作成')
add_bullet(doc, 'https://github.com を開く')
add_bullet(doc, '右上「+」→「New repository」をクリック')
add_bullet(doc, 'Repository name：ai-keiei-honbu')
add_bullet(doc, 'Private（必須）を選択')
add_bullet(doc, '「Initialize this repository」はチェックしない')
add_bullet(doc, '「Create repository」をクリック')
doc.add_paragraph()
add_h2(doc, 'Step2：Personal Access Tokenの取得')
add_bullet(doc, 'GitHubの右上アイコン→Settings')
add_bullet(doc, '左メニュー一番下→Developer settings')
add_bullet(doc, 'Personal access tokens→Tokens (classic)')
add_bullet(doc, 'Generate new token (classic)')
add_bullet(doc, 'Note：ai-keiei、Expiration：No expiration、repoにチェック')
add_bullet(doc, 'Generate token→表示されたトークンをコピー（一度しか表示されない）')
doc.add_paragraph()
add_h2(doc, 'Step3：AI経営本部をGitHubにアップロード')
add_body(doc, 'Claude Codeに以下を入力してください。【】の部分を書き換えてください。')
doc.add_paragraph()
add_green(doc, '''▼ Claude Codeへの入力（コピペしてください）

以下を順番に実行してください。

cd ~/Desktop/AI経営本部
git init
git add .
git commit -m "AI経営本部 初回セットアップ完了"
git remote add origin https://【GitHubユーザー名】:【トークン】@github.com/【GitHubユーザー名】/ai-keiei-honbu.git
git branch -M main
git push -u origin main''')
doc.add_paragraph()
add_h2(doc, 'Step4：iPhoneからの操作設定')
add_bullet(doc, 'iPhoneのSafariで https://claude.ai/code を開く')
add_bullet(doc, '「Claude Code Webを開始」をタップ')
add_bullet(doc, 'GitHubアカウントでログイン')
add_bullet(doc, 'ai-keiei-honbuリポジトリを選択')
add_bullet(doc, '入力欄に日本語で指示を出してみる')
doc.add_paragraph()
add_body(doc, 'iPhoneの画面下に「AI経営…」「ai-keiei…」「main」と表示されれば成功です！')
doc.add_page_break()

# 完成確認
add_h1(doc, '完成確認チェックリスト')
add_hr(doc)
add_h2(doc, '器（フレームワーク）の完成確認')
add_bullet(doc, '□ Claude Codeがターミナルで起動できる')
add_bullet(doc, '□ AI経営本部フォルダがデスクトップに存在する')
add_bullet(doc, '□ AI社長のCLAUDE.mdが設置されている')
add_bullet(doc, '□ 各プロジェクトのCLAUDE.mdが設置されている')
add_bullet(doc, '□ Gemini・OpenAI APIキーが設定されている')
add_bullet(doc, '□ 3大AIが全て接続・動作している')
add_bullet(doc, '□ 基本パイプラインが動作する')
add_bullet(doc, '□ 日次成果物フォルダに成果物が保存される')
add_bullet(doc, '□ GitHubリポジトリが作成されている')
add_bullet(doc, '□ iPhoneからClaude Code Webが操作できる')
doc.add_paragraph()
add_green(doc, '全てチェックできれば「器」の完成です！おめでとうございます。')
doc.add_paragraph()
add_h2(doc, '次のステップ：スキルの追加')
add_body(doc, '器が完成したら、自分のビジネスに合わせてスキルを一つずつ追加していきます。')
doc.add_paragraph()
add_bullet(doc, '焦って全部一度に追加しない')
add_bullet(doc, '一つのスキルを追加したら必ず動作確認する')
add_bullet(doc, 'うまくいったスキルをGitHubにプッシュして保存する')
add_bullet(doc, '使えば使うほどAIが賢くなっていく')
doc.add_page_break()

# トラブルシューティング
add_h1(doc, 'トラブルシューティング')
add_hr(doc)
add_h2(doc, 'よくあるエラーと解決策')
add_body(doc, '■ npmが見つからない')
add_bullet(doc, 'Node.jsを再インストールしてください')
doc.add_paragraph()
add_body(doc, '■ 権限エラーが出る')
add_bullet(doc, 'コマンドの前にsudoをつけて実行してください')
doc.add_paragraph()
add_body(doc, '■ APIキーが認識されない')
add_bullet(doc, 'ターミナルを再起動して再度設定してください')
doc.add_paragraph()
add_body(doc, '■ GitHubへのアップロードが失敗する')
add_bullet(doc, 'Personal Access Tokenを再発行してください')
doc.add_paragraph()
add_h2(doc, 'わからないことがあれば')
add_body(doc, 'Claude Codeに日本語でそのまま聞いてください。')
add_green(doc, '「〜がうまくいかないのですが、どうすれば解決できますか？」')
add_body(doc, 'と話しかけるだけで解決策を提示してくれます。')

# 保存
output_path = os.path.expanduser(
    '~/Desktop/AI経営本部/スタッフ向けマニュアル/AI経営本部_構築マニュアル_汎用版.docx'
)
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'✅ Word完成：{output_path}')
