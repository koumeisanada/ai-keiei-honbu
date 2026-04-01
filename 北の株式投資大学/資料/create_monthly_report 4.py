import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SAVE_DIR = os.path.expanduser("~/Desktop/AI経営本部/北の株式投資大学/資料")

def set_table_border(table, color='1A1A2E'):
    tbl=table._tbl; tblPr=tbl.tblPr; tblBorders=OxmlElement('w:tblBorders')
    for bn in ['top','left','bottom','right','insideH','insideV']:
        b=OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'0'); b.set(qn('w:color'),color); tblBorders.append(b)
    tblPr.append(tblBorders)

def set_cell_bg(cell, hc):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hc); tcPr.append(shd)

def add_page_number(doc):
    for s in doc.sections:
        f=s.footer; p=f.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run()
        fc1=OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'),'begin'); it=OxmlElement('w:instrText'); it.text='PAGE'; fc2=OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'),'end')
        r._r.append(fc1); r._r.append(it); r._r.append(fc2); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x88,0x88,0x88)

def sec_title(doc, text):
    p=doc.add_paragraph(); r=p.add_run(f'■ {text}'); r.font.size=Pt(15); r.font.bold=True; r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6'); bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'534AB7'); pBdr.append(bot); pPr.append(pBdr)

def sub_title(doc, text):
    p=doc.add_paragraph(); r=p.add_run(f'◆ {text}'); r.font.size=Pt(12); r.font.bold=True; r.font.color.rgb=RGBColor(0x53,0x4A,0xB7)

def body(doc, text):
    p=doc.add_paragraph(text)
    for r in p.runs: r.font.size=Pt(10.5)

def voice_box(doc, text):
    t=doc.add_table(rows=1,cols=1); set_table_border(t,'1A1A2E'); set_cell_bg(t.rows[0].cells[0],'E8E8F0')
    c=t.rows[0].cells[0]; c.paragraphs[0].clear()
    pl=c.add_paragraph(); rl=pl.add_run('真田孔明の視点'); rl.font.size=Pt(9); rl.font.bold=True; rl.font.color.rgb=RGBColor(0x88,0x88,0x88)
    pb=c.add_paragraph(text)
    for r in pb.runs: r.font.size=Pt(10.5); r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)
    doc.add_paragraph()

def point_box(doc, title, text):
    t=doc.add_table(rows=1,cols=1); set_table_border(t,'534AB7'); set_cell_bg(t.rows[0].cells[0],'F0F0FF')
    c=t.rows[0].cells[0]; c.paragraphs[0].clear()
    pt=c.add_paragraph(); rt=pt.add_run(f'💡 {title}'); rt.font.size=Pt(11); rt.font.bold=True; rt.font.color.rgb=RGBColor(0x53,0x4A,0xB7)
    pb=c.add_paragraph(text)
    for r in pb.runs: r.font.size=Pt(10.5)
    doc.add_paragraph()

def warn_box(doc, text):
    t=doc.add_table(rows=1,cols=1); set_table_border(t,'CC0000'); set_cell_bg(t.rows[0].cells[0],'FFF0F0')
    c=t.rows[0].cells[0]; p=c.paragraphs[0]; r=p.add_run(f'⚠️ 注意　{text}'); r.font.size=Pt(10.5); r.font.bold=True; r.font.color.rgb=RGBColor(0xCC,0x00,0x00)
    doc.add_paragraph()

def make_table(doc, headers, data, col_widths=None):
    t=doc.add_table(rows=len(data)+1, cols=len(headers)); set_table_border(t)
    hr=t.rows[0]
    for i,h in enumerate(headers):
        hr.cells[i].text=h; set_cell_bg(hr.cells[i],'1A1A2E')
        for p in hr.cells[i].paragraphs:
            for r in p.runs: r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.size=Pt(10.5)
    for i,row_data in enumerate(data):
        row=t.rows[i+1]
        for j,val in enumerate(row_data):
            row.cells[j].text=val
        if i%2==0:
            for c in row.cells: set_cell_bg(c,'F0F0FF')
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    if col_widths:
        for i,w in enumerate(col_widths): t.columns[i].width=Cm(w)
    doc.add_paragraph()

def create_report():
    today=datetime.now().strftime('%Y年%m月%d日')
    doc=Document()
    for s in doc.sections:
        s.page_width=Inches(8.27); s.page_height=Inches(11.69); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5); s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
    add_page_number(doc)

    # 表紙
    for _ in range(2): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('米国テクノロジー企業'); r.font.size=Pt(28); r.font.bold=True; r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; r2=p2.add_run('最新動向レポート'); r2.font.size=Pt(22); r2.font.bold=True; r2.font.color.rgb=RGBColor(0x53,0x4A,0xB7)
    doc.add_paragraph()
    p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER; r3=p3.add_run('2026年3月号'); r3.font.size=Pt(16); r3.font.color.rgb=RGBColor(0x53,0x4A,0xB7)
    doc.add_paragraph()
    it=doc.add_table(rows=3,cols=2); it.alignment=WD_TABLE_ALIGNMENT.CENTER; set_table_border(it)
    for i,(l,v) in enumerate([('発行日',today),('発行者','真田孔明（北の株式投資大学）'),('免責事項','本資料は情報提供目的。投資判断は自己責任で')]):
        row=it.rows[i]; row.cells[0].text=l; row.cells[1].text=v; set_cell_bg(row.cells[0],'E8E8F0')
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_page_break()

    # 今月の要点
    sec_title(doc,'今月の要点')
    doc.add_paragraph()
    make_table(doc, ['#','トピック'], [
        ('①','NVIDIAのGTC2026：売上$1兆予測・Groq統合・Alpamayo自動運転戦略を発表'),
        ('②','AIモデル著作権問題：楽天AI・Cursor炎上事件が示すオープンソース利用の倫理'),
        ('③','EVシフトの現実：欧米6社で6.7兆円の損失・トヨタの慎重戦略が評価される'),
        ('④','OpenAIがAnthropicの躍進に危機感・Claude Codeがエンジニア標準に'),
        ('⑤','電動飛行機（eVTOL）が実用化フェーズへ・ハワイで初の商用運航開始予定'),
    ], [2,14.5])
    voice_box(doc,'10年間、NVIDIAを研究し続けてきた。\nGTC2026でJensen Huangが語った$1兆予測は、アナリストの想定を7%上回る数字だった。\n株価は動かなかった。でも僕は、この数字が意味することを知っている。\n自動運転・ロボット・宇宙AIデータセンター。\nNVIDIAは次の10年の設計図を、GTC2026で見せた。')
    doc.add_page_break()

    # NVIDIA
    sec_title(doc,'NVIDIA：GTC2026の全容')
    doc.add_paragraph()
    sub_title(doc,'1. 売上$1兆予測の意味')
    body(doc,'GTC2026の基調講演でJensen Huang CEOは、BlackwellおよびVera Rubinチップの売上が2027年末までに少なくとも$1兆（約150兆円）に達すると宣言した。\n\nこれはアナリスト予想を7%上回る数字だが、株価は発表直後に5%高騰した後、すぐに元の水準に戻った。\n\nNVIDIAの株価総額は既に$4.5兆近くあり、それを大きく動かすには日本のGDPに匹敵する資金が必要だからだ。')
    doc.add_paragraph()
    sub_title(doc,'2. Groq技術の製品統合')
    body(doc,'NVIDIAはGroq LPUをVera Rubin GPUと組み合わせるハイブリッド構成を発表した。\n\n・Rubin GPU：Transformerのprefill（入力プロンプトの一括計算）を担当\n・Groq LPU：decode（KV Cacheを活用したトークン生成）を担当\n\nこの組み合わせにより全体のスループットが大幅に向上する。')
    doc.add_paragraph()
    sub_title(doc,'3. Alpamayo：自動運転戦略の核心')
    body(doc,'NVIDIAが発表した「Alpamayo」は自動運転開発向けのオープンソースAIエコシステム。\n\n特徴：\n・「Alpamayo 1」（約100億パラメータのVLAモデル）を無償公開\n・8万時間分のマルチカメラ走行データも公開\n・シミュレーター「AlphaSim」も無償提供\n\n既にメルセデス・トヨタ・GM・BYD・現代・日産・吉利が採用を表明。')
    point_box(doc,'Alpamayoの戦略的意味','NVIDIAは「タレの作り方」を無償提供することで各自動車メーカーが独自の「秘伝のタレ」を作れるようにした。\nその結果、全製造プロセスでNVIDIAのAIチップが使われる構造を作り上げた。')
    voice_box(doc,'30分の1の株価の頃からNVIDIAを持ち続けた理由がここにある。\nNVIDIAはGPUを売る会社ではない。AIが使われる全ての場所でインフラを握る会社だ。')
    doc.add_page_break()

    # AI著作権
    sec_title(doc,'AIモデル著作権問題：楽天・Cursor炎上事件')
    doc.add_paragraph()
    make_table(doc, ['企業','内容','投資への示唆'], [
        ('楽天 Rakuten AI 3.0','日本政府補助金で開発した「日本最強AI」が中国製DeepSeek-V3のリブランド版と判明。MITライセンス違反で炎上。','AI開発の透明性が今後より重要に。'),
        ('Cursor Composer 2','「独自AIモデル」として提供した製品が中国Moonshot AIのKimi K2.5ベースと判明。創業者が即謝罪し収束。','迅速な謝罪が評価された。'),
    ], [3.5,9,4])
    warn_box(doc,'DeepSeekには中国政府の影響による強いバイアスが確認されている。日本政府・トヨタ・三菱重工・ソフトバンクが利用禁止令。')
    voice_box(doc,'著作権の線引きが難しくなっている。これはAI業界が成熟する過程で必ず通る道だ。\n重要なのは、問題が起きた時に迅速に誠実に対応できるかどうか。')
    doc.add_page_break()

    # EVシフト
    sec_title(doc,'EVシフトの現実：6社で6.7兆円の損失')
    doc.add_paragraph()
    body(doc,'欧州6社に加え、GM・フォード・ホンダなど従来型自動車メーカーがEV事業で大幅赤字を計上。\n\n合計損失：約6.7兆円（欧州6社）\n主な原因：\n・Tesla・BYDなど先行勢に対抗できなかった\n・EV需要が期待ほど伸びなかった\n・トランプ政権がEV連邦税優遇（7,500ドル）を廃止')
    doc.add_paragraph()
    sub_title(doc,'トヨタの評価が変わった理由')
    body(doc,'EVシフトに大きく出遅れたトヨタが結果的に大きな損失を被らずに済んだことで評価が一変。\n\n・ハイブリッドで稼ぎながら自分のペースでEVシフト\n・NVIDIAのAlpamayo採用を表明済み←ここに注目')
    voice_box(doc,'Tesla Model YとAudi e-Tronを数年間所有している僕から言えば、EVシフトは必然だ。\nガソリンスタンドに行かなくていい。静かで臭くない。メンテが楽。一度経験したら戻れない。\n方向は変わらないと確信している。')
    doc.add_page_break()

    # OpenAI vs Anthropic
    sec_title(doc,'OpenAI vs Anthropic：市場の地殻変動')
    doc.add_paragraph()
    body(doc,'OpenAIのアプリ部門CEOのFidji Simoが「AnthropicのエンタープライズシェアがOpenAIを奪っている。ウェイクアップコールだ」と警告。\n\n現在の市場構図：\n・消費者向け：ChatGPT（OpenAI）がデファクト\n・開発者向け：Claude Code（Anthropic）が急速に標準化\n・企業向け：Claude for Workがエンタープライズ市場を侵食中')
    doc.add_paragraph()
    sub_title(doc,'AI市場のコモディティ化と差別化')
    body(doc,'AIモデルそのものはコモディティ化が進んでいる。差別化要因は「ツールの使い勝手」と「スピード」に移行。\n\n上場の可能性：OpenAI・Anthropicとも上場可能性あり。\nどちらか一方なら、Anthropicの株を選ぶ。')
    voice_box(doc,'OpenAIとAnthropicの戦いはこれからが本番だ。\nでも僕が注目しているのは、この競争を通じてNVIDIAのGPU需要がどう変わるかということだ。\nどちらが勝ってもNVIDIAにとっては喜ばしい。')
    doc.add_page_break()

    # eVTOL
    sec_title(doc,'電動飛行機（eVTOL）：実用化フェーズへ')
    doc.add_paragraph()
    body(doc,'電動飛行機が「研究段階」から「実用化フェーズ」に入った。\n\n・ハワイでSurf AirがBeta Technologiesの電動飛行機を採用決定\n・ドバイでJoby AviationによるeVTOLタクシーが2026年商用開始予定')
    make_table(doc,['企業','タイプ','状況'],[
        ('Joby Aviation','eVTOL（垂直離着陸型）','ドバイでタクシーサービス予定・上場済・赤字継続中'),
        ('Beta Technologies','通常離着陸型（CTOL）','ハワイで貨物便から開始・上場済・赤字継続中'),
    ],[3.5,5,8])
    warn_box(doc,'Joby・Betaともに上場企業だが赤字継続中。ハイリスク・ハイリターンの投資。')
    voice_box(doc,'「まだ早い」から「もうすぐ来る」に変わる瞬間を、NVIDIAで経験した人間として慎重に見ている。')
    doc.add_page_break()

    # その他
    sec_title(doc,'その他の注目トピック')
    doc.add_paragraph()
    make_table(doc,['テーマ','内容','示唆'],[
        ('NVIDIA NVFP4','4ビット精度でFP8同等性能を達成。12Bパラメータで10兆トークン学習。','NVIDIA独自チップの優位性を改めて証明。'),
        ('Microsoft BitNet','1.58ビット計算。CPUで100Bパラメータを人間の読書速度で動作。','ローカルAI実行の可能性。研究段階。'),
        ('Karpathy Autoresearch','AIが自分のトレーニングコードを改善するループ。自然言語でAIを制御。','AI自律改善の方向性。'),
        ('愛犬癌ワクチン','エンジニアがAI活用で末期癌の愛犬用mRNAワクチンを2ヶ月で開発。腫瘍75%縮小。','パーソナライズ医療の可能性。'),
    ],[3.5,9,4])
    doc.add_page_break()

    # まとめ
    sec_title(doc,'今月の投資家向けまとめ')
    doc.add_paragraph()
    make_table(doc,['企業','注目理由','時間軸'],[
        ('NVIDIA','$1兆予測・Groq統合・Alpamayo・ロボットAI・宇宙DC','長期保有継続'),
        ('Tesla','EV先行者・自動運転・Robotaxi・Optimus','中長期'),
        ('Anthropic','Claude Code急成長・エンタープライズ侵食・上場可能性','上場後に検討'),
        ('Joby Aviation','eVTOL実用化フェーズ・2026年型式証明見込み','ハイリスク・少額'),
        ('トヨタ自動車','NVIDIA Alpamayo採用・ハイブリッド強み','中長期'),
    ],[3.5,9,4])
    voice_box(doc,'10年間NVIDIAを研究してきた投資家として言えることがある。\n\n良い投資とは、10年後の世界を想像して、その世界でインフラになっている企業を今買うことだ。\n\n2016年にNVIDIAを買った時、周りは誰も見ていなかった。\n2026年の今、電動飛行機・ロボット・宇宙AIデータセンターを見ている。\n\n注意：これは投資アドバイスではない。全ての投資判断は自己責任で。')

    os.makedirs(SAVE_DIR, exist_ok=True)
    fp=os.path.join(SAVE_DIR,f'{datetime.now().strftime("%Y%m%d")}_米国テクノロジー企業_最新動向レポート_2026年3月号.docx')
    doc.save(fp)
    print(f'✅ レポート完成：{fp}')
    return fp

if __name__=='__main__':
    print('='*60); print('米国テクノロジー企業 最新動向レポート 2026年3月号'); print('='*60)
    fp=create_report()
    print(f'\n完成ファイル：{fp}')
