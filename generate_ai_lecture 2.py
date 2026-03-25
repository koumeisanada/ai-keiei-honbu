#!/usr/bin/env python3
"""AI最新情報 講義用Word資料生成スクリプト"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os

NAVY = RGBColor(0x1B, 0x4F, 0x72)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_NAVY = RGBColor(0xD4, 0xE6, 0xF1)
DARK_TEXT = RGBColor(0x1C, 0x1C, 0x1C)
ACCENT_ORANGE = RGBColor(0xE6, 0x7E, 0x22)
TODAY = datetime.now().strftime("%Y%m%d")
SAVE_DIR = os.path.expanduser("~/Desktop/AI経営本部")

doc = Document()

# --- スタイル設定 ---
style = doc.styles['Normal']
font = style.font
font.name = 'Yu Gothic'
font.size = Pt(11)
font.color.rgb = DARK_TEXT
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Yu Gothic'
    hs.font.color.rgb = NAVY
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    if level == 1:
        hs.font.size = Pt(24)
    elif level == 2:
        hs.font.size = Pt(18)
    else:
        hs.font.size = Pt(14)


def add_colored_heading(text, level=2, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def add_shaded_block(title, content_lines, shade_color="D4E6F1"):
    """色付きヘッダー付きブロック"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    # ヘッダー行
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade_color}"/>')
    # タイトル段落
    tp = cell.paragraphs[0]
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(13)
    tr.font.color.rgb = NAVY
    tr.font.name = 'Yu Gothic'
    # コンテンツ
    for line in content_lines:
        p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(10.5)
        r.font.name = 'Yu Gothic'
    # セル背景
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shading)
    doc.add_paragraph()


def add_highlight_box(text, bg_color="FFF3CD"):
    """ハイライトボックス"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT_ORANGE
    r.font.name = 'Yu Gothic'
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shading)
    doc.add_paragraph()


def add_level_table(levels):
    """レベル別活用法テーブル"""
    t = doc.add_table(rows=len(levels)+1, cols=2)
    t.style = 'Light List Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    hdr[0].text = "難易度"
    hdr[1].text = "活用方法"
    for i, (lv, desc) in enumerate(levels):
        row = t.rows[i+1].cells
        row[0].text = lv
        row[1].text = desc
    doc.add_paragraph()


# ============================================================
# 表紙
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("個人が今すぐ使える・\n知っておくべきAI最新情報")
tr.font.size = Pt(32)
tr.font.color.rgb = NAVY
tr.bold = True
tr.font.name = 'Yu Gothic'

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("〜2026年3月号 AI活用講義資料〜")
sr.font.size = Pt(16)
sr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
sr.font.name = 'Yu Gothic'

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = date_p.add_run(f"作成日：{datetime.now().strftime('%Y年%m月%d日')}")
dr.font.size = Pt(12)
dr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
dr.font.name = 'Yu Gothic'

doc.add_page_break()

# ============================================================
# エグゼクティブサマリー
# ============================================================
add_colored_heading("エグゼクティブサマリー", level=1)

doc.add_paragraph()
add_highlight_box("今月のAI界隈を一言で表すと？ → 「AIエージェント元年、ツールからチームメンバーへ」")

add_colored_heading("今月の最重要トピック TOP5", level=2)

top5 = [
    "1. Claude Opus 4.6リリース — 100万トークン対応・エージェントチーム機能で複数AIが協調作業",
    "2. Gemini 3がGmail/Chromeに本格統合 — AI Inbox・サイドパネル自動化で日常業務が激変",
    "3. Sora 2 vs Runway Gen-4 vs Veo 3.1 — 動画生成AI三つ巴、音声同期も標準装備に",
    "4. 日本「AI推進法」成立 — 国の基本計画策定・1兆円規模の国産AI開発支援が始動",
    "5. Suno v5 / ACE-Step v1.5 — 音楽生成AIがプロ品質に到達、オープンソース版も登場",
]
for item in top5:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# 第1章：画像・動画・音楽生成の最前線
# ============================================================
add_colored_heading("第1章　画像・動画・音楽生成の最前線", level=1)

# SLIDE A
add_shaded_block("SLIDE A：何が起きたか", [
    "【動画生成AI三つ巴の構図】",
    "・Sora 2（OpenAI）：最長25秒の動画生成、音声・効果音の自動同期に対応。ChatGPT Plus（$20/月〜）で利用可能",
    "・Runway Gen-4/4.5：フレーム単位の精密制御、Alephビデオエディタで実写映像のAI編集を実現",
    "・Google Veo 3.1：YouTube Shorts連携、1日2億再生規模のAI動画エコシステム構築中",
    "",
    "【音楽生成AIの急進化】",
    "・Suno v5：テキスト入力だけでプロ品質の楽曲生成。12ステム書き出し対応、日本語歌詞にも対応強化",
    "・ACE-Step v1.5（オープンソース）：4GB未満のVRAMで動作し、Suno v4.5〜v5相当の品質を無料で実現",
    "・WMG（ワーナーミュージック）がSuno/Udioとライセンス契約 → AI音楽の「合法ライセンス時代」幕開け",
    "",
    "【画像生成AI】",
    "・Bing Image Creator（DALL-Eベース）：完全無料で高品質画像生成",
    "・Adobe Firefly：商用利用に最適化。著作権クリアな学習データ使用",
    "・Google NanoBanana：OpenAIの画像生成比20倍の速度（数秒 vs 3分）",
    "",
    "【音声生成AI】",
    "・ElevenLabs v3：笑い声・ため息・ささやき等の感情表現を70以上の言語で実現",
    "・Scribe v2 Realtime：150ms未満の低遅延で90言語以上の音声認識。日本語の単語エラー率5%未満",
    "",
    "出典：WaveSpeedAI, FahmAI, Solfej, テクノエッジ, ElevenLabs公式（2026年1〜3月）",
])

# SLIDE B
add_shaded_block("SLIDE B：どう使えるか", [
    "Level 1（誰でも）：Bing Image Creatorで無料画像生成 → SNS投稿素材を即作成",
    "Level 1（誰でも）：Suno無料枠で30秒のBGM生成 → YouTube動画・ポッドキャストに活用",
    "Level 2（少し工夫）：Canva AI + ElevenLabsでナレーション付き動画を自作 → リール投稿",
    "Level 2（少し工夫）：Sora 2でプロモーション動画の草案を生成 → 撮影前のイメージ共有に",
    "Level 3以上（自動化）：ElevenLabs Agents APIでノーコード音声対話チャットボット構築",
    "Level 3以上（自動化）：ACE-Step v1.5をローカル環境で運用 → 月額コストゼロの音楽生成パイプライン",
], shade_color="D5F5E3")

# SLIDE C
add_shaded_block("SLIDE C：見落とすと損するポイント", [
    "・Sora 2の無料枠は2026年1月に廃止済み。ChatGPT Plus以上が必須",
    "・AI生成動画の著作権は「生成・利用段階」で既存著作物と同様の侵害リスクあり（文化庁見解）",
    "・音楽生成AIで生成した楽曲を商用利用する場合、各サービスの利用規約で商用ライセンスの確認が必要",
    "・ElevenLabsの無料枠は月間文字数制限あり。ナレーション量産には有料プラン（$5/月〜）が必要",
], shade_color="FADBD8")

doc.add_page_break()

# ============================================================
# 第2章：主要モデルの新機能
# ============================================================
add_colored_heading("第2章　ChatGPT / Claude / Gemini 主要モデルの新機能", level=1)

add_shaded_block("SLIDE A：何が起きたか", [
    "【ChatGPT（GPT-5.3/5.4系）】2026年3月時点",
    "・GPT-5.1系は廃止 → GPT-5.3 Instant / GPT-5.4 Thinking / GPT-5.4 Proに統合",
    "・GPT-5.4 miniがFree/Goユーザーに展開。無料でも最新モデルが使える時代に",
    "・インタラクティブ学習機能：70以上の数学・科学トピックを視覚的に学習可能",
    "・Deep Research機能でWeb上の情報を自動収集・分析・レポート化",
    "",
    "【Claude（Opus 4.6）】2026年3月リリース",
    "・最大100万トークンのコンテキスト対応（書籍数冊分を一度に処理）",
    "・Claude Codeの「エージェントチーム」機能：複数インスタンスが連携してタスクを分担",
    "・Adaptive Thinking：問題の難易度に応じて思考深度を自動調整",
    "・音声モード、Telegram/Discord連携のChannels機能、自動メモリ機能を追加",
    "",
    "【Gemini 3】2026年1〜3月の大型アップデート",
    "・Gmail AI Inbox：メール要約と能動的な受信トレイ管理を実現",
    "・Gemini in Chrome：サイドパネルでマルチタスク・タスク自動化を強化",
    "・Google Home統合：音声応答速度40%改善。スマートホーム操作をAIが最適化",
    "・Personal Intelligence：Googleアプリ横断で個人最適化情報を提供",
    "",
    "出典：OpenAI Release Notes, Anthropic公式, Google公式, Impress Watch, ITmedia（2026年1〜3月）",
])

add_shaded_block("SLIDE B：どう使えるか", [
    "Level 1（誰でも）：ChatGPT無料版でGPT-5.4 miniを使う → 日常の質問応答・文章作成が格段に向上",
    "Level 1（誰でも）：Gemini in Chromeを有効化 → ブラウジング中にサイドパネルで要約・翻訳・質問",
    "Level 2（少し工夫）：Claude無料版で長文PDF・契約書を丸ごとアップロード → 要点抽出・比較分析",
    "Level 2（少し工夫）：Gmail AI Inboxで毎朝の受信トレイ整理を自動化 → 重要メールだけに集中",
    "Level 3以上（自動化）：Claude Codeエージェントチームで開発タスクを並列自動化",
    "Level 3以上（自動化）：ChatGPT Deep Researchで市場調査→レポート生成を完全自動化",
], shade_color="D5F5E3")

add_shaded_block("SLIDE C：見落とすと損するポイント", [
    "・ChatGPT無料版は回数制限あり。ヘビーユーザーはPlus（$20/月）が実質必須",
    "・Claude Opus 4.6の100万トークンはPro（$20/月）以上のプラン。無料版はコンテキスト制限あり",
    "・Gemini in ChromeはブラウジングデータをGoogleに送信する設定に注意。プライバシー設定を確認",
    "・3大AIを「比較して使い分ける」のが2026年の正解。1つに固定するのは損",
], shade_color="FADBD8")

doc.add_page_break()

# ============================================================
# 第3章：今すぐ使えるAIツール
# ============================================================
add_colored_heading("第3章　今すぐ使えるAIツール・アプリ（無料中心）", level=1)

add_shaded_block("SLIDE A：何が起きたか", [
    "2026年3月時点で、個人が無料〜低コストで使えるAIツールが爆発的に増加。",
    "以下が「今すぐ登録して使える」主要ツール一覧：",
])

# ツール一覧テーブル
t = doc.add_table(rows=11, cols=4)
t.style = 'Light List Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["ツール名", "主な用途", "無料/有料", "難易度"]
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h

tools = [
    ("ChatGPT（Free/Go）", "会話・文章作成・学習・画像生成", "無料", "Level 1"),
    ("Claude（Free）", "長文分析・コード支援・100万トークン", "無料", "Level 1"),
    ("Gemini", "検索・Gmail連携・Chrome連携", "無料", "Level 1"),
    ("Bing Image Creator", "DALL-Eベース画像生成", "完全無料", "Level 1"),
    ("Canva AI", "デザイン・動画生成（5回無料）", "無料枠あり", "Level 1"),
    ("Suno", "テキストから音楽生成", "無料枠あり", "Level 1"),
    ("ElevenLabs", "音声・ナレーション生成", "無料枠あり", "Level 2"),
    ("Kling AI", "動画生成（中国発・高品質）", "無料枠あり", "Level 2"),
    ("Meitu / AIピカソ", "スマホで画像・アバター生成", "無料", "Level 1"),
    ("ACE-Step v1.5", "音楽生成（ローカル実行・OSS）", "完全無料", "Level 3"),
]
for i, (name, use, price, level) in enumerate(tools):
    row = t.rows[i+1].cells
    row[0].text = name
    row[1].text = use
    row[2].text = price
    row[3].text = level

doc.add_paragraph()

add_highlight_box("今すぐ使える：ChatGPT + Claude + Gemini の3つを無料登録するだけで、AIの80%の恩恵を受けられる")

add_shaded_block("SLIDE B：どう使えるか", [
    "Level 1（誰でも）：ChatGPT/Claude/Geminiの3つに無料登録 → 用途別に使い分け開始",
    "Level 1（誰でも）：Bing Image Creatorでブログ・SNS用の画像を無限生成",
    "Level 2（少し工夫）：ElevenLabsでナレーション生成 + Canva AIで動画編集 → プロ級コンテンツ制作",
    "Level 2（少し工夫）：Sunoで自社CMソング・YouTube用BGMを自作",
    "Level 3以上（自動化）：ACE-Stepをローカル環境で運用し、月額ゼロで音楽を量産",
], shade_color="D5F5E3")

add_shaded_block("SLIDE C：見落とすと損するポイント", [
    "・「無料」でもアカウント登録時に個人情報を提供している。プライバシーポリシーは確認すること",
    "・無料枠には生成回数・文字数の上限がある。業務利用なら有料プランの検討を",
    "・中国発AIツール（Kling AI等）はデータの保存先・利用規約に注意",
    "・スマホアプリの「AI○○」は玉石混交。レビューと運営元を必ず確認",
], shade_color="FADBD8")

doc.add_page_break()

# ============================================================
# 第4章：AIエージェント・自動化
# ============================================================
add_colored_heading("第4章　AIエージェント・自動化の個人活用", level=1)

add_shaded_block("SLIDE A：何が起きたか", [
    "2026年は「AIエージェント元年」。AIが「ツール」から「チームメンバー」に変化した年。",
    "",
    "【主要な動き】",
    "・Claude Code エージェントチーム：複数のAIインスタンスが協調して開発タスクを分担実行",
    "・Gemini in Chrome AutoBrowse（2026年1月〜）：価格比較・旅行計画・購入完了までAIが自律実行",
    "・ChatGPT Deep Research：Web上の情報を自動収集 → 分析 → レポート生成を一気通貫で実行",
    "・Make / Zapier / n8n：ノーコードでAI自動化ワークフローを構築。個人でも「AI秘書」が持てる時代",
    "・マルチエージェントシステムへの進化：複数エージェントが連携し、複雑な目標を自律的に達成",
    "",
    "出典：Anthropic公式, Google公式, @IT, Zenn, GBase（2026年1〜3月）",
])

add_shaded_block("SLIDE B：どう使えるか", [
    "Level 1（誰でも）：ChatGPTに「○○を調べてレポートにまとめて」と依頼するだけ → Deep Researchが自動実行",
    "Level 2（少し工夫）：Zapier無料枠でGmail→Slack通知、フォーム→スプレッドシート記録を自動化",
    "Level 2（少し工夫）：Gemini in ChromeでWebリサーチ → 比較表作成を半自動化",
    "Level 3（自動化）：n8nでAI × API連携の自動化パイプラインを構築（メール受信→AI要約→Slack投稿）",
    "Level 4（上級）：Claude Codeエージェントチームで複数タスクを並列自動実行",
], shade_color="D5F5E3")

add_shaded_block("SLIDE C：見落とすと損するポイント", [
    "・AIエージェントは「完璧」ではない。出力の事実確認（ファクトチェック）は人間の責務",
    "・自動化ツール（Zapier/Make）の無料枠は月100〜250タスクが上限。超えると有料",
    "・Chrome AutoBrowseは購入完了まで自律実行する → 意図しない決済に注意",
    "・エージェントに機密情報を渡す場合、データの送信先・保管先を必ず確認",
], shade_color="FADBD8")

doc.add_page_break()

# ============================================================
# 第5章：バズったAIニュース
# ============================================================
add_colored_heading("第5章　世界でバズったAIニュース・事件", level=1)

add_shaded_block("SLIDE A：何が起きたか", [
    "【Disney × OpenAI 10億ドル提携】",
    "・Sora 2にディズニーキャラクターのライセンス動画生成機能が搭載",
    "・エンタメ×AIの史上最大級の提携として世界中で話題に",
    "",
    "【Sora無料枠廃止（2026年1月）】",
    "・OpenAIがSoraの無料ティアを廃止。ChatGPT Plus以上が必須に",
    "・「AI民主化」の方針転換としてユーザーの反発を招く",
    "",
    "【WMG（ワーナーミュージック）× Suno/Udio ライセンス契約】",
    "・メジャーレーベルがAI音楽生成ツールと正式にライセンス契約",
    "・AI音楽の「グレーゾーン」が「合法ライセンス」に移行する転換点",
    "",
    "【推論コスト・電力問題の表面化】",
    "・AI推論の「1タスクあたりの実コスト」が業界の主要課題に浮上",
    "・GPU・電力供給の逼迫が各社の事業計画に影響",
    "・NVIDIA GTC 2026でDLSS 5（ニューラルレンダリング）と新LPUチップを発表",
    "",
    "【Apple Intelligence版Siri】",
    "・iOS 18.2でSiri × ChatGPT（GPT-4）連携が公式リリース",
    "・写真分析・創作文章などの高度タスクをSiriがChatGPTに委任する仕組み",
    "",
    "出典：各社公式発表, Impress Watch, 日経（2025年末〜2026年3月）",
])

add_shaded_block("SLIDE B：どう使えるか", [
    "Level 1（誰でも）：iPhone設定 → Siri → ChatGPT拡張を有効化 → 音声でAIに質問",
    "Level 2（少し工夫）：AI音楽生成ツールのライセンス動向を把握 → 安心して商用BGM制作",
    "Level 3以上（自動化）：推論コスト意識を持ち、APIコール設計を最適化 → コスト削減",
], shade_color="D5F5E3")

add_shaded_block("SLIDE C：見落とすと損するポイント", [
    "・Disney提携でも「キャラクター利用=自由」ではない。ライセンス範囲を確認すること",
    "・Sora無料枠廃止の流れは他サービスにも波及する可能性。無料依存は危険",
    "・AI推論コストの上昇は、最終的にサブスク料金の値上げにつながる見通し",
], shade_color="FADBD8")

doc.add_page_break()

# ============================================================
# 第6章：AIと仕事・スキル・キャリア
# ============================================================
add_colored_heading("第6章　AIと仕事・スキル・キャリアへの影響", level=1)

add_shaded_block("SLIDE A：何が起きたか", [
    "【落合陽一氏の予測（2026年）】",
    "・「2026年にはほとんどの知的作業がAIに置き換わる」",
    "・定型的な事務職・カスタマーサービス・士業のデータ処理部分は代替リスクが高い",
    "",
    "【IMF報告（2026年1月）】",
    "・新スキルを4つ以上習得した労働者は報酬が有意に高い傾向",
    "・AIを「使いこなす」スキルが労働市場での競争力に直結",
    "",
    "【AI時代の3カテゴリーの新職種】",
    "・AIを「創る」人：研究者・エンジニア（高度専門）",
    "・AIを「活かす」人：プロンプトエンジニア・AI活用コンサルタント（中級）",
    "・AIを「使いこなす」人：全職種で求められる基礎スキル（必須）",
    "",
    "【差別化のキーワード：「とげ」】",
    "・論理的飛躍・創造性・人間にしかできない「とげ」（独自の視点）が価値を持つ",
    "・AIが処理できる定型作業はAIに任せ、人間は「とげ」に集中する時代",
    "",
    "出典：エンジニアtype, IMF公式レポート, Microsoft AI Trends 2026（2026年1月）",
])

add_shaded_block("SLIDE B：どう使えるか", [
    "Level 1（誰でも）：毎日1回ChatGPTに仕事の質問をする習慣を始める → AI活用の第一歩",
    "Level 2（少し工夫）：自分の業務フローを棚卸しし、AIで代替できるタスクを特定 → 時間の再配分",
    "Level 2（少し工夫）：AIプロンプトエンジニアリングの基礎を学ぶ → 出力品質が3倍向上",
    "Level 3以上（自動化）：AI × 自分の専門分野を掛け合わせた「AIを活かす人」ポジションを確立",
], shade_color="D5F5E3")

add_shaded_block("SLIDE C：見落とすと損するポイント", [
    "・「AIに仕事を奪われる」ではなく「AIを使える人に仕事が移る」が正確な構図",
    "・プロンプト力よりも「何を聞くべきか」を知っている業務知識が本質的に重要",
    "・AI資格・検定が乱立中。実務経験 > 資格の優先順位を間違えないこと",
    "・40代以上こそ業務知識 × AI活用の掛け算で最も価値を発揮できる世代",
], shade_color="FADBD8")

doc.add_page_break()

# ============================================================
# 第7章：AI規制・著作権
# ============================================================
add_colored_heading("第7章　AI規制・著作権・社会問題", level=1)

add_shaded_block("SLIDE A：何が起きたか", [
    "【日本「AI推進法」成立】",
    "・国の基本計画策定・研究開発推進・人権侵害リスクへの調査指導が法定化",
    "・AI事業者ガイドラインの法的根拠が整備された",
    "",
    "【著作権の現状整理（文化庁見解）】",
    "・「開発・学習段階」：原則として広く許容される（著作権法30条の4）",
    "・「生成・利用段階」：既存著作物と同様の侵害リスクが存在",
    "・2026年2月時点でAI著作権侵害に関する確定判決はまだなし",
    "",
    "【日本新聞協会の動き】",
    "・生成AIによる記事無断収集への法的対応を政府に要望",
    "・メディア業界 vs AI企業の対立構造が鮮明に",
    "",
    "【知的財産推進計画2026】",
    "・2026年夏に策定予定。AI生成コンテンツの権利関係がより明確化される見込み",
    "",
    "出典：文化庁「AIと著作権」, 日経新聞, AXメディア（2025〜2026年）",
])

add_shaded_block("SLIDE B：どう使えるか", [
    "Level 1（誰でも）：AI生成コンテンツを公開する前に「既存作品との類似性」を自分でチェック",
    "Level 2（少し工夫）：商用利用時はAdobe Firefly等「著作権クリア」を明示するツールを優先",
    "Level 3以上（自動化）：自社のAI利用ガイドラインを策定 → 従業員への教育を実施",
], shade_color="D5F5E3")

add_shaded_block("SLIDE C：見落とすと損するポイント", [
    "・「AIが生成したから著作権フリー」は誤解。利用段階では通常の著作権法が適用される",
    "・確定判決がまだないため「グレーゾーン」が多い。リスクを取りすぎないこと",
    "・海外サービス利用時は、日本法だけでなく各国の規制動向にも注意",
    "・AI推進法は「規制」よりも「推進」寄りの法律。過度な萎縮は不要",
], shade_color="FADBD8")

doc.add_page_break()

# ============================================================
# 第8章：日本語対応・日本発AI
# ============================================================
add_colored_heading("第8章　日本語対応・日本発のAI最新情報", level=1)

add_shaded_block("SLIDE A：何が起きたか", [
    "【NTT「tsuzumi」】",
    "・パラメータ6億（超軽量版）と70億（軽量版）の2モデル",
    "・少ないリソースで高い日本語処理能力。テキスト・図表・画像のマルチモーダル対応",
    "・オンプレミス環境で動作可能 → 機密データを外部に出さずにAI活用できる",
    "",
    "【サイバーエージェント CyberAgentLM3】",
    "・225億パラメータ。日本語処理で世界トップクラスの性能",
    "・複雑な日本語の正確な理解と高度な文章生成が可能",
    "",
    "【Japanese StableLM（Stability AI）】",
    "・7億パラメータ、7,500億トークンのWebデータで学習",
    "・オープンソース（無料）で日本語・英語に最適化",
    "",
    "【政府の国産AI支援】",
    "・2026年度から5年間で1兆円規模の支援を決定",
    "・ソフトバンク・Preferred Networksなど十数社が設立する新会社に助成",
    "・国内最高レベルの基盤モデル開発を官民で推進",
    "",
    "出典：SHIFT AI, DX/AI研究所, 各社公式発表（2025〜2026年）",
])

add_shaded_block("SLIDE B：どう使えるか", [
    "Level 1（誰でも）：ChatGPT/Claude/Geminiの日本語対応は既に実用レベル → まずはこれらから",
    "Level 2（少し工夫）：日本語に特化したタスクはCyberAgentLM3等の国産モデルを検討",
    "Level 4（上級）：tsuzumiをオンプレミスで導入 → 機密データのAI処理をクラウド外で実現",
    "Level 4（上級）：Japanese StableLMをローカル環境で運用 → カスタマイズ自由な日本語AI",
], shade_color="D5F5E3")

add_shaded_block("SLIDE C：見落とすと損するポイント", [
    "・国産AIは性能向上中だが、2026年3月時点ではChatGPT/Claude/Geminiの日本語性能が依然トップ",
    "・「国産だから安全」とは限らない。データ管理体制は個別に確認が必要",
    "・1兆円の政府支援は「今すぐ使える」ではなく「2〜3年後に成果が出る」長期投資",
    "・日本語特化モデルはニッチ用途（法務・医療・行政文書）で真価を発揮する",
], shade_color="FADBD8")

doc.add_page_break()

# ============================================================
# まとめ
# ============================================================
add_colored_heading("まとめ", level=1)

add_colored_heading("今月の「これだけ使え」ランキング TOP5", level=2)

add_shaded_block("コスパ・使いやすさ・インパクトで総合評価", [
    "1位　ChatGPT（GPT-5.4 mini無料版）— 万能AI。まず最初に登録すべき1つ",
    "2位　Claude（Opus 4.6）— 長文分析の王者。PDF・契約書の分析なら最強",
    "3位　Gemini in Chrome — ブラウザに統合済み。Googleユーザーなら設定するだけ",
    "4位　Bing Image Creator — 完全無料の画像生成。SNS運用に必須",
    "5位　Suno — テキストだけで音楽生成。BGM・ジングル制作の革命",
])

doc.add_paragraph()
add_colored_heading("今週のアクション3問", level=2)

questions = [
    ("Q1：今日試せるツールは何か？",
     "→ ChatGPT / Claude / Gemini の3つに無料登録し、同じ質問を3つに投げて回答を比較する"),
    ("Q2：あなたの仕事で最初に自動化できるのはどこか？",
     "→ 毎日繰り返しているメール返信・データ整理・レポート作成を棚卸しし、1つをAIに委任する"),
    ("Q3：Level 2→3に上がるために今月やること1つは？",
     "→ Zapier/Make無料枠で「Gmail受信→Slack通知」の自動化ワークフローを1つ作ってみる"),
]

for q, a in questions:
    p = doc.add_paragraph()
    qr = p.add_run(q)
    qr.bold = True
    qr.font.size = Pt(12)
    qr.font.color.rgb = NAVY
    qr.font.name = 'Yu Gothic'
    p2 = doc.add_paragraph(a)
    p2.paragraph_format.space_after = Pt(8)

# ============================================================
# 保存
# ============================================================
filename = f"AI_news_lecture_{TODAY}.docx"
filepath = os.path.join(SAVE_DIR, filename)
doc.save(filepath)
print(f"✅ 保存完了：{filepath}")
